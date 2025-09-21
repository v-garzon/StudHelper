# StudHelper Backend - Complete Implementation

## 1. PROJECT CONFIGURATION

### pyproject.toml
[build-system]
requires = ["setuptools>=61.0"]
build-backend = "setuptools.build_meta"

[project]
name = "studhelper-backend"
version = "2.0.0"
description = "AI-powered study assistant backend"
authors = [{name = "StudHelper Team"}]
dependencies = [
    "fastapi>=0.104.1",
    "uvicorn[standard]>=0.24.0",
    "pydantic>=2.5.0",
    "pydantic-settings>=2.1.0",
    "sqlalchemy>=2.0.23",
    "alembic>=1.13.1",
    "asyncpg>=0.29.0",
    "python-multipart>=0.0.6",
    "openai>=1.6.1",
    "tiktoken>=0.5.2",
    "tenacity>=8.2.3",
    "pdfplumber>=0.10.3",
    "pypdf>=3.17.4",
    "python-docx>=1.1.0",
    "python-pptx>=0.6.23",
    "yt-dlp>=2023.12.30",
    "openai-whisper>=20231117",
    "chromadb>=0.4.22",
    "langchain>=0.1.0",
    "langchain-openai>=0.0.2",
    "langchain-chroma>=0.1.0",
    "sentence-transformers>=2.2.2",
    "python-jose[cryptography]>=3.3.0",
    "passlib[bcrypt]>=1.7.4",
    "httpx>=0.25.2",
    "aiofiles>=23.2.1",
]

[project.optional-dependencies]
dev = [
    "pytest>=7.4.3",
    "pytest-asyncio>=0.21.1",
    "pytest-cov>=4.1.0",
    "black>=23.12.1",
    "flake8>=6.1.0",
    "mypy>=1.8.0",
    "isort>=5.13.2",
]

[tool.black]
line-length = 88
target-version = ['py311']

[tool.isort]
profile = "black"

### requirements/base.txt
fastapi==0.104.1
uvicorn[standard]==0.24.0
pydantic==2.5.0
pydantic-settings==2.1.0
sqlalchemy==2.0.23
alembic==1.13.1
asyncpg==0.29.0
python-multipart==0.0.6
openai==1.6.1
tiktoken==0.5.2
tenacity==8.2.3
pdfplumber==0.10.3
pypdf==3.17.4
python-docx==1.1.0
python-pptx==0.6.23
yt-dlp==2023.12.30
openai-whisper==20231117
chromadb==0.4.22
langchain==0.1.0
langchain-openai==0.0.2
langchain-chroma==0.1.0
sentence-transformers==2.2.2
python-jose[cryptography]==3.3.0
passlib[bcrypt]==1.7.4
httpx==0.25.2
aiofiles==23.2.1

### requirements/dev.txt
-r base.txt
pytest==7.4.3
pytest-asyncio==0.21.1
pytest-cov==4.1.0
httpx==0.25.2
respx==0.20.2
pytest-mock==3.12.0
factory-boy==3.3.0
black==23.12.1
flake8==6.1.0
mypy==1.8.0
isort==5.13.2
python-dotenv==1.0.0
watchdog==3.0.0

### .env.example
# Environment
ENVIRONMENT=development
DEBUG=true

# Security
SECRET_KEY=your-super-secret-key-here-minimum-32-characters-long
JWT_ALGORITHM=HS256
ACCESS_TOKEN_EXPIRE_MINUTES=30

# Database
DATABASE_URL=postgresql://user:password@localhost:5432/studhelper

# OpenAI
OPENAI_API_KEY=sk-your-openai-api-key-here
OPENAI_ORG_ID=org-your-organization-id

# Pricing (per 1M tokens)
ECONOMIC_INPUT_PRICE=0.15
ECONOMIC_OUTPUT_PRICE=0.60
STANDARD_INPUT_PRICE=2.50
STANDARD_OUTPUT_PRICE=10.00

# ChromaDB
CHROMA_PERSIST_DIR=./chroma_data
CHROMA_HOST=localhost
CHROMA_PORT=8000

# File Upload
MAX_FILE_SIZE=100000000
UPLOAD_DIR=./uploads

# CORS (comma-separated)
CORS_ORIGINS=http://localhost:3000,http://localhost:8501

# Rate Limiting
RATE_LIMIT_REQUESTS=100
RATE_LIMIT_WINDOW=3600

### .gitignore
__pycache__/
*.py[cod]
*$py.class
*.so
.Python
build/
develop-eggs/
dist/
downloads/
eggs/
.eggs/
lib/
lib64/
parts/
sdist/
var/
wheels/
*.egg-info/
.installed.cfg
*.egg
MANIFEST

.env
.env.*
!.env.example

venv/
ENV/
env/
.venv/

.pytest_cache/
.coverage
htmlcov/
.tox/

.DS_Store
.vscode/
.idea/

uploads/
chroma_data/
*.log

## 2. CORE CONFIGURATION

### src/__init__.py
"""StudHelper Backend - AI-powered study assistant."""

__version__ = "2.0.0"

### src/config.py
"""Application configuration management."""

from functools import lru_cache
from typing import List, Optional

from pydantic import Field, computed_field
from pydantic_settings import BaseSettings


class Settings(BaseSettings):
    """Application settings."""
    
    # Environment
    ENVIRONMENT: str = "development"
    DEBUG: bool = False
    
    # Security
    SECRET_KEY: str = Field(..., min_length=32)
    JWT_ALGORITHM: str = "HS256"
    ACCESS_TOKEN_EXPIRE_MINUTES: int = 30
    
    # Database
    DATABASE_URL: str = Field("postgresql://localhost/studhelper")
    
    # OpenAI Configuration
    OPENAI_API_KEY: str = Field(...)
    OPENAI_ORG_ID: Optional[str] = None
    
    # Pricing per 1M tokens
    ECONOMIC_INPUT_PRICE: float = 0.15
    ECONOMIC_OUTPUT_PRICE: float = 0.60
    STANDARD_INPUT_PRICE: float = 2.50
    STANDARD_OUTPUT_PRICE: float = 10.00
    
    # ChromaDB
    CHROMA_PERSIST_DIR: str = "./chroma_data"
    CHROMA_HOST: str = "localhost"
    CHROMA_PORT: int = 8000
    
    # File Upload
    MAX_FILE_SIZE: int = 100_000_000  # 100MB
    UPLOAD_DIR: str = "./uploads"
    
    # CORS
    CORS_ORIGINS: List[str] = ["http://localhost:3000", "http://localhost:8501"]
    
    # Rate Limiting
    RATE_LIMIT_REQUESTS: int = 100
    RATE_LIMIT_WINDOW: int = 3600
    
    @computed_field
    @property
    def is_production(self) -> bool:
        """Check if running in production."""
        return self.ENVIRONMENT == "production"
    
    class Config:
        env_file = ".env"
        case_sensitive = True


@lru_cache()
def get_settings() -> Settings:
    """Get cached settings instance."""
    return Settings()

### src/core/__init__.py
"""Core infrastructure components."""

### src/core/database.py
"""Database connection and session management."""

import logging
from contextlib import asynccontextmanager
from typing import AsyncGenerator

from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.orm import DeclarativeBase

from src.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

# Database engine
engine = create_async_engine(
    settings.DATABASE_URL,
    echo=settings.DEBUG,
    pool_pre_ping=True,
    pool_size=10,
    max_overflow=20,
)

# Session factory
AsyncSessionLocal = async_sessionmaker(
    engine, class_=AsyncSession, expire_on_commit=False
)


class Base(DeclarativeBase):
    """Base class for all database models."""
    pass


async def create_tables():
    """Create all database tables."""
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
        logger.info("Database tables created")


@asynccontextmanager
async def get_db() -> AsyncGenerator[AsyncSession, None]:
    """Get database session."""
    async with AsyncSessionLocal() as session:
        try:
            yield session
        finally:
            await session.close()


async def get_db_session() -> AsyncGenerator[AsyncSession, None]:
    """FastAPI dependency for database session."""
    async with get_db() as session:
        yield session

### src/core/exceptions.py
"""Custom exception classes."""

from typing import Any, Dict, Optional


class StudHelperException(Exception):
    """Base exception for StudHelper."""
    
    def __init__(
        self,
        message: str,
        status_code: int = 500,
        details: Optional[Dict[str, Any]] = None,
    ):
        self.message = message
        self.status_code = status_code
        self.details = details or {}
        super().__init__(self.message)


class AuthenticationError(StudHelperException):
    """Authentication related errors."""
    
    def __init__(self, message: str = "Authentication failed"):
        super().__init__(message, status_code=401)


class AuthorizationError(StudHelperException):
    """Authorization related errors."""
    
    def __init__(self, message: str = "Access denied"):
        super().__init__(message, status_code=403)


class ValidationError(StudHelperException):
    """Validation related errors."""
    
    def __init__(self, message: str, details: Optional[Dict[str, Any]] = None):
        super().__init__(message, status_code=422, details=details)


class NotFoundError(StudHelperException):
    """Resource not found errors."""
    
    def __init__(self, message: str = "Resource not found"):
        super().__init__(message, status_code=404)


class ProcessingError(StudHelperException):
    """Document processing errors."""
    
    def __init__(self, message: str, details: Optional[Dict[str, Any]] = None):
        super().__init__(message, status_code=422, details=details)


class RateLimitError(StudHelperException):
    """Rate limiting errors."""
    
    def __init__(self, message: str = "Rate limit exceeded"):
        super().__init__(message, status_code=429)


class ExternalServiceError(StudHelperException):
    """External service errors."""
    
    def __init__(self, message: str, service: str):
        super().__init__(f"{service}: {message}", status_code=502)

### src/core/logging.py
"""Logging configuration."""

import logging.config
import sys
from typing import Dict, Any

from src.config import get_settings

settings = get_settings()


def setup_logging() -> None:
    """Configure application logging."""
    config: Dict[str, Any] = {
        "version": 1,
        "disable_existing_loggers": False,
        "formatters": {
            "default": {
                "format": "%(asctime)s - %(name)s - %(levelname)s - %(message)s",
            },
            "detailed": {
                "format": (
                    "%(asctime)s - %(name)s - %(levelname)s - "
                    "%(funcName)s:%(lineno)d - %(message)s"
                ),
            },
        },
        "handlers": {
            "console": {
                "class": "logging.StreamHandler",
                "level": "INFO",
                "formatter": "default",
                "stream": sys.stdout,
            },
            "file": {
                "class": "logging.FileHandler",
                "level": "DEBUG" if settings.DEBUG else "INFO",
                "formatter": "detailed",
                "filename": "studhelper.log",
            },
        },
        "loggers": {
            "studhelper": {
                "level": "DEBUG" if settings.DEBUG else "INFO",
                "handlers": ["console", "file"],
                "propagate": False,
            },
            "uvicorn": {
                "level": "INFO",
                "handlers": ["console"],
                "propagate": False,
            },
        },
        "root": {
            "level": "WARNING",
            "handlers": ["console"],
        },
    }
    
    logging.config.dictConfig(config)
    
    # Set application logger
    logger = logging.getLogger("studhelper")
    logger.info("Logging configured")

### src/core/deps.py
"""FastAPI dependency functions."""

from typing import Optional

from fastapi import Depends, HTTPException, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.jwt_handler import decode_jwt_token
from src.auth.models import User
from src.auth.service import AuthService
from src.core.database import get_db_session

security = HTTPBearer()


async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
    db: AsyncSession = Depends(get_db_session),
) -> User:
    """Get current authenticated user."""
    try:
        payload = decode_jwt_token(credentials.credentials)
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid authentication credentials",
            )
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication credentials",
        )
    
    auth_service = AuthService(db)
    user = await auth_service.get_user_by_id(int(user_id))
    if user is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="User not found",
        )
    
    return user


async def get_optional_user(
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security),
    db: AsyncSession = Depends(get_db_session),
) -> Optional[User]:
    """Get current user if authenticated, None otherwise."""
    if not credentials:
        return None
    
    try:
        return await get_current_user(credentials, db)
    except HTTPException:
        return None


async def get_default_user(
    db: AsyncSession = Depends(get_db_session),
) -> User:
    """Get or create default user for single-user mode."""
    auth_service = AuthService(db)
    user = await auth_service.get_user_by_email("default@studhelper.local")
    
    if not user:
        # Create default user
        user = await auth_service.create_user(
            email="default@studhelper.local",
            username="default_user",
            password="default_password_change_me",
        )
    
    return user

## 3. AUTHENTICATION SYSTEM

### src/auth/__init__.py
"""Authentication system."""

### src/auth/models.py
"""User authentication models."""

from datetime import datetime
from typing import Optional

from sqlalchemy import Boolean, DateTime, Integer, String, Text
from sqlalchemy.orm import Mapped, mapped_column
from sqlalchemy.sql import func

from src.core.database import Base


class User(Base):
    """User model."""
    
    __tablename__ = "users"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    email: Mapped[str] = mapped_column(String(255), unique=True, index=True)
    username: Mapped[str] = mapped_column(String(100), unique=True, index=True)
    hashed_password: Mapped[str] = mapped_column(String(255))
    
    # User profile
    full_name: Mapped[Optional[str]] = mapped_column(String(255))
    role: Mapped[str] = mapped_column(String(50), default="student")  # student, teacher, admin
    
    # Account status
    is_active: Mapped[bool] = mapped_column(Boolean, default=True)
    is_verified: Mapped[bool] = mapped_column(Boolean, default=False)
    
    # Timestamps
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )
    updated_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now(), onupdate=func.now()
    )
    last_login: Mapped[Optional[datetime]] = mapped_column(DateTime(timezone=True))
    
    # Preferences
    preferences: Mapped[Optional[str]] = mapped_column(Text)  # JSON string
    
    def __repr__(self) -> str:
        return f"<User(id={self.id}, email='{self.email}', username='{self.username}')>"

### src/auth/schemas.py
"""Authentication schemas."""

from datetime import datetime
from typing import Optional

from pydantic import BaseModel, EmailStr, Field


class UserBase(BaseModel):
    """Base user schema."""
    email: EmailStr
    username: str = Field(..., min_length=3, max_length=100)
    full_name: Optional[str] = None


class UserCreate(UserBase):
    """User creation schema."""
    password: str = Field(..., min_length=6, max_length=100)


class UserUpdate(BaseModel):
    """User update schema."""
    email: Optional[EmailStr] = None
    username: Optional[str] = Field(None, min_length=3, max_length=100)
    full_name: Optional[str] = None
    is_active: Optional[bool] = None


class UserResponse(UserBase):
    """User response schema."""
    id: int
    role: str
    is_active: bool
    is_verified: bool
    created_at: datetime
    updated_at: datetime
    last_login: Optional[datetime] = None
    
    class Config:
        from_attributes = True


class LoginRequest(BaseModel):
    """Login request schema."""
    email: EmailStr
    password: str


class LoginResponse(BaseModel):
    """Login response schema."""
    access_token: str
    token_type: str = "bearer"
    user: UserResponse


class TokenData(BaseModel):
    """Token data schema."""
    email: Optional[str] = None

### src/auth/jwt_handler.py
"""JWT token handling."""

from datetime import datetime, timedelta, timezone
from typing import Dict, Optional

from jose import JWTError, jwt
from passlib.context import CryptContext

from src.config import get_settings
from src.core.exceptions import AuthenticationError

settings = get_settings()

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")


def hash_password(password: str) -> str:
    """Hash a password."""
    return pwd_context.hash(password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify a password against its hash."""
    return pwd_context.verify(plain_password, hashed_password)


def create_access_token(data: Dict[str, str], expires_delta: Optional[timedelta] = None) -> str:
    """Create a JWT access token."""
    to_encode = data.copy()
    
    if expires_delta:
        expire = datetime.now(timezone.utc) + expires_delta
    else:
        expire = datetime.now(timezone.utc) + timedelta(minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES)
    
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, settings.SECRET_KEY, algorithm=settings.JWT_ALGORITHM)
    return encoded_jwt


def decode_jwt_token(token: str) -> Dict[str, str]:
    """Decode and validate a JWT token."""
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.JWT_ALGORITHM])
        return payload
    except JWTError as e:
        raise AuthenticationError(f"Invalid token: {str(e)}")

### src/auth/service.py
"""Authentication service."""

import logging
from datetime import datetime, timezone
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.jwt_handler import create_access_token, hash_password, verify_password
from src.auth.models import User
from src.auth.schemas import LoginResponse, UserCreate, UserResponse, UserUpdate
from src.core.exceptions import AuthenticationError, NotFoundError, ValidationError

logger = logging.getLogger(__name__)


class AuthService:
    """Authentication service."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
    
    async def create_user(self, user_data: UserCreate) -> User:
        """Create a new user."""
        # Check if user already exists
        existing_user = await self.get_user_by_email(user_data.email)
        if existing_user:
            raise ValidationError("User with this email already exists")
        
        existing_username = await self.get_user_by_username(user_data.username)
        if existing_username:
            raise ValidationError("User with this username already exists")
        
        # Create new user
        user = User(
            email=user_data.email,
            username=user_data.username,
            full_name=user_data.full_name,
            hashed_password=hash_password(user_data.password),
        )
        
        self.db.add(user)
        await self.db.commit()
        await self.db.refresh(user)
        
        logger.info(f"Created user: {user.email}")
        return user
    
    async def authenticate_user(self, email: str, password: str) -> Optional[User]:
        """Authenticate a user."""
        user = await self.get_user_by_email(email)
        if not user:
            return None
        
        if not verify_password(password, user.hashed_password):
            return None
        
        # Update last login
        user.last_login = datetime.now(timezone.utc)
        await self.db.commit()
        
        return user
    
    async def login(self, email: str, password: str) -> LoginResponse:
        """Login user and return token."""
        user = await self.authenticate_user(email, password)
        if not user:
            raise AuthenticationError("Invalid email or password")
        
        if not user.is_active:
            raise AuthenticationError("Account is disabled")
        
        # Create access token
        access_token = create_access_token(data={"sub": str(user.id)})
        
        return LoginResponse(
            access_token=access_token,
            user=UserResponse.model_validate(user),
        )
    
    async def get_user_by_id(self, user_id: int) -> Optional[User]:
        """Get user by ID."""
        result = await self.db.execute(select(User).where(User.id == user_id))
        return result.scalar_one_or_none()
    
    async def get_user_by_email(self, email: str) -> Optional[User]:
        """Get user by email."""
        result = await self.db.execute(select(User).where(User.email == email))
        return result.scalar_one_or_none()
    
    async def get_user_by_username(self, username: str) -> Optional[User]:
        """Get user by username."""
        result = await self.db.execute(select(User).where(User.username == username))
        return result.scalar_one_or_none()
    
    async def update_user(self, user_id: int, user_data: UserUpdate) -> User:
        """Update user."""
        user = await self.get_user_by_id(user_id)
        if not user:
            raise NotFoundError("User not found")
        
        # Update fields
        for field, value in user_data.model_dump(exclude_unset=True).items():
            setattr(user, field, value)
        
        user.updated_at = datetime.now(timezone.utc)
        await self.db.commit()
        await self.db.refresh(user)
        
        return user
    
    async def delete_user(self, user_id: int) -> bool:
        """Delete user."""
        user = await self.get_user_by_id(user_id)
        if not user:
            return False
        
        await self.db.delete(user)
        await self.db.commit()
        
        logger.info(f"Deleted user: {user.email}")
        return True

### src/auth/router.py
"""Authentication routes."""

from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.schemas import LoginRequest, LoginResponse, UserCreate, UserResponse
from src.auth.service import AuthService
from src.core.database import get_db_session
from src.core.exceptions import AuthenticationError, ValidationError

router = APIRouter()


@router.post("/register", response_model=UserResponse, status_code=status.HTTP_201_CREATED)
async def register(
    user_data: UserCreate,
    db: AsyncSession = Depends(get_db_session),
):
    """Register a new user."""
    try:
        auth_service = AuthService(db)
        user = await auth_service.create_user(user_data)
        return UserResponse.model_validate(user)
    except ValidationError as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=e.message)


@router.post("/login", response_model=LoginResponse)
async def login(
    login_data: LoginRequest,
    db: AsyncSession = Depends(get_db_session),
):
    """Login user."""
    try:
        auth_service = AuthService(db)
        return await auth_service.login(login_data.email, login_data.password)
    except AuthenticationError as e:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=e.message)


@router.get("/me", response_model=UserResponse)
async def get_current_user_info(
    current_user = Depends(get_current_user),
):
    """Get current user information."""
    return UserResponse.model_validate(current_user)

## 4. DOCUMENT PROCESSING SYSTEM

### src/documents/__init__.py
"""Document processing system."""

### src/documents/models.py
"""Document models."""

from datetime import datetime
from typing import Optional

from sqlalchemy import DateTime, ForeignKey, Integer, String, Text, Boolean, Float
from sqlalchemy.orm import Mapped, mapped_column, relationship
from sqlalchemy.sql import func

from src.core.database import Base


class Document(Base):
    """Document model."""
    
    __tablename__ = "documents"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    user_id: Mapped[int] = mapped_column(Integer, ForeignKey("users.id"))
    
    # File information
    filename: Mapped[str] = mapped_column(String(255))
    original_filename: Mapped[str] = mapped_column(String(255))
    file_type: Mapped[str] = mapped_column(String(50))  # pdf, docx, pptx, youtube
    file_size: Mapped[Optional[int]] = mapped_column(Integer)
    file_path: Mapped[Optional[str]] = mapped_column(String(500))
    
    # Processing information
    status: Mapped[str] = mapped_column(String(50), default="pending")  # pending, processing, completed, failed
    content: Mapped[Optional[str]] = mapped_column(Text)
    metadata: Mapped[Optional[str]] = mapped_column(Text)  # JSON string
    
    # Document statistics
    page_count: Mapped[Optional[int]] = mapped_column(Integer)
    word_count: Mapped[Optional[int]] = mapped_column(Integer)
    processing_time: Mapped[Optional[float]] = mapped_column(Float)
    
    # Timestamps
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )
    updated_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now(), onupdate=func.now()
    )
    processed_at: Mapped[Optional[datetime]] = mapped_column(DateTime(timezone=True))
    
    def __repr__(self) -> str:
        return f"<Document(id={self.id}, filename='{self.filename}', status='{self.status}')>"


class DocumentChunk(Base):
    """Document chunk model for RAG."""
    
    __tablename__ = "document_chunks"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    document_id: Mapped[int] = mapped_column(Integer, ForeignKey("documents.id"))
    
    # Chunk information
    chunk_index: Mapped[int] = mapped_column(Integer)
    content: Mapped[str] = mapped_column(Text)
    
    # Positioning
    start_char: Mapped[Optional[int]] = mapped_column(Integer)
    end_char: Mapped[Optional[int]] = mapped_column(Integer)
    page_number: Mapped[Optional[int]] = mapped_column(Integer)
    
    # Metadata
    metadata: Mapped[Optional[str]] = mapped_column(Text)  # JSON string
    
    # Vector information
    embedding_id: Mapped[Optional[str]] = mapped_column(String(255))  # ChromaDB ID
    
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )

### src/documents/schemas.py
"""Document schemas."""

from datetime import datetime
from typing import Dict, List, Optional, Any

from pydantic import BaseModel, Field


class DocumentUpload(BaseModel):
    """Document upload schema."""
    filename: str
    file_type: str
    file_size: Optional[int] = None


class DocumentResponse(BaseModel):
    """Document response schema."""
    id: int
    filename: str
    original_filename: str
    file_type: str
    file_size: Optional[int]
    status: str
    page_count: Optional[int]
    word_count: Optional[int]
    processing_time: Optional[float]
    created_at: datetime
    updated_at: datetime
    processed_at: Optional[datetime]
    
    class Config:
        from_attributes = True


class DocumentProcessingResult(BaseModel):
    """Document processing result."""
    success: bool
    document_id: int
    message: str
    chunks_created: Optional[int] = None
    processing_time: Optional[float] = None
    errors: Optional[List[str]] = None


class YouTubeUpload(BaseModel):
    """YouTube upload schema."""
    url: str = Field(..., description="YouTube video URL")
    title: Optional[str] = None


class DocumentStats(BaseModel):
    """Document statistics."""
    total_documents: int
    total_pages: int
    total_words: int
    by_type: Dict[str, int]
    processing_status: Dict[str, int]

### src/documents/processors/__init__.py
"""Document processors."""

### src/documents/processors/base.py
"""Base document processor."""

import asyncio
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentProcessor(ABC):
    """Abstract base class for document processors."""
    
    def __init__(self):
        self.supported_types: List[str] = []
    
    @abstractmethod
    async def process(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process document and return extracted content."""
        pass
    
    def can_process(self, file_type: str) -> bool:
        """Check if this processor can handle the file type."""
        return file_type.lower() in self.supported_types
    
    def _create_chunk_metadata(
        self,
        source: str,
        page: Optional[int] = None,
        chunk_index: int = 0,
        **kwargs
    ) -> Dict[str, Any]:
        """Create metadata for a content chunk."""
        metadata = {
            "source": source,
            "chunk_index": chunk_index,
            "processed_at": datetime.now().isoformat(),
            "processor": self.__class__.__name__,
        }
        
        if page is not None:
            metadata["page"] = page
        
        metadata.update(kwargs)
        return metadata
    
    async def _run_in_thread(self, func, *args, **kwargs):
        """Run CPU-bound operation in thread pool."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, func, *args, **kwargs)

### src/documents/processors/pdf_processor.py
"""PDF document processor."""

import asyncio
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional

import pdfplumber
import pypdf
from io import BytesIO

from .base import DocumentProcessor

logger = logging.getLogger(__name__)


class PDFProcessor(DocumentProcessor):
    """PDF document processor."""
    
    def __init__(self):
        super().__init__()
        self.supported_types = ["pdf", "application/pdf"]
    
    async def process(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process PDF file and extract text content."""
        try:
            # Determine processing method
            if await self._is_simple_pdf(file_path):
                result = await self._process_with_pypdf(file_path)
            else:
                result = await self._process_with_pdfplumber(file_path)
            
            result.update({
                "file_type": "pdf",
                "processor": "PDFProcessor",
                "file_path": str(file_path),
            })
            
            return result
            
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "chunks": [],
                "metadata": {"file_type": "pdf", "error": str(e)},
            }
    
    async def _is_simple_pdf(self, file_path: Path) -> bool:
        """Check if PDF is simple enough for pypdf."""
        def check():
            try:
                with open(file_path, 'rb') as file:
                    reader = pypdf.PdfReader(file)
                    if len(reader.pages) > 0:
                        page_text = reader.pages[0].extract_text()
                        # Simple heuristic: if no tables or complex formatting
                        return '\t' not in page_text and '|' not in page_text
                return True
            except:
                return False
        
        return await self._run_in_thread(check)
    
    async def _process_with_pypdf(self, file_path: Path) -> Dict[str, Any]:
        """Process simple PDFs with pypdf."""
        def extract():
            chunks = []
            metadata = {
                "extraction_method": "pypdf",
                "pages": 0,
                "word_count": 0,
            }
            
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text().strip()
                    if text:
                        chunks.append({
                            "content": text,
                            "metadata": self._create_chunk_metadata(
                                source=str(file_path),
                                page=page_num + 1,
                                chunk_index=len(chunks),
                                extraction_method="pypdf",
                            ),
                        })
                        metadata["word_count"] += len(text.split())
                
                metadata["pages"] = len(reader.pages)
            
            return {
                "success": True,
                "chunks": chunks,
                "metadata": metadata,
            }
        
        return await self._run_in_thread(extract)
    
    async def _process_with_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        """Process complex PDFs with pdfplumber."""
        def extract():
            chunks = []
            metadata = {
                "extraction_method": "pdfplumber",
                "pages": 0,
                "word_count": 0,
                "tables_found": 0,
            }
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    if text and text.strip():
                        chunks.append({
                            "content": text.strip(),
                            "metadata": self._create_chunk_metadata(
                                source=str(file_path),
                                page=page_num + 1,
                                chunk_index=len(chunks),
                                extraction_method="pdfplumber",
                            ),
                        })
                        metadata["word_count"] += len(text.split())
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = self._table_to_text(table)
                            if table_text:
                                chunks.append({
                                    "content": table_text,
                                    "metadata": self._create_chunk_metadata(
                                        source=str(file_path),
                                        page=page_num + 1,
                                        chunk_index=len(chunks),
                                        table_index=table_idx,
                                        content_type="table",
                                        extraction_method="pdfplumber",
                                    ),
                                })
                                metadata["tables_found"] += 1
                
                metadata["pages"] = len(pdf.pages)
            
            return {
                "success": True,
                "chunks": chunks,
                "metadata": metadata,
            }
        
        return await self._run_in_thread(extract)
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table to readable text."""
        if not table or not table[0]:
            return ""
        
        text_parts = []
        for row in table:
            if row:
                clean_row = [cell.strip() if cell else "" for cell in row]
                text_parts.append(" | ".join(clean_row))
        
        return "\n".join(text_parts)

### src/documents/processors/docx_processor.py
"""Word document processor."""

import logging
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document as DocxDocument
from docx.oxml.table import CT_Table
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base import DocumentProcessor

logger = logging.getLogger(__name__)


class DocxProcessor(DocumentProcessor):
    """Word document processor."""
    
    def __init__(self):
        super().__init__()
        self.supported_types = [
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
    
    async def process(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process Word document and extract text content."""
        try:
            result = await self._extract_content(file_path)
            result.update({
                "file_type": "docx",
                "processor": "DocxProcessor",
                "file_path": str(file_path),
            })
            return result
            
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "chunks": [],
                "metadata": {"file_type": "docx", "error": str(e)},
            }
    
    async def _extract_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Word document."""
        def extract():
            chunks = []
            metadata = {
                "extraction_method": "python-docx",
                "paragraphs": 0,
                "tables": 0,
                "word_count": 0,
            }
            
            doc = DocxDocument(file_path)
            
            # Process document elements in order
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    if text:
                        chunks.append({
                            "content": text,
                            "metadata": self._create_chunk_metadata(
                                source=str(file_path),
                                chunk_index=len(chunks),
                                content_type="paragraph",
                                extraction_method="python-docx",
                            ),
                        })
                        metadata["word_count"] += len(text.split())
                        metadata["paragraphs"] += 1
                
                elif isinstance(element, CT_Table):
                    # Table
                    table = Table(element, doc)
                    table_text = self._table_to_text(table)
                    if table_text:
                        chunks.append({
                            "content": table_text,
                            "metadata": self._create_chunk_metadata(
                                source=str(file_path),
                                chunk_index=len(chunks),
                                content_type="table",
                                extraction_method="python-docx",
                            ),
                        })
                        metadata["word_count"] += len(table_text.split())
                        metadata["tables"] += 1
            
            return {
                "success": True,
                "chunks": chunks,
                "metadata": metadata,
            }
        
        return await self._run_in_thread(extract)
    
    def _table_to_text(self, table: Table) -> str:
        """Convert Word table to text."""
        text_parts = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            
            if any(text for text in row_text):
                text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)

### src/documents/processors/youtube_processor.py
"""YouTube video processor."""

import asyncio
import logging
import tempfile
import whisper
import yt_dlp
from pathlib import Path
from typing import Dict, List, Any, Optional

from openai import AsyncOpenAI

from .base import DocumentProcessor
from src.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class YouTubeProcessor(DocumentProcessor):
    """YouTube video processor."""
    
    def __init__(self):
        super().__init__()
        self.supported_types = ["youtube", "video/youtube"]
        self.openai_client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
        self._whisper_model = None
    
    async def process(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process YouTube video and extract transcript."""
        try:
            # Get video metadata
            video_metadata = await self._get_video_metadata(url)
            
            # Download audio
            audio_path = await self._download_audio(url)
            
            try:
                # Transcribe audio (use local Whisper by default)
                transcript = await self._transcribe_with_local_whisper(audio_path)
                
                # Create chunks from transcript
                chunks = self._create_transcript_chunks(transcript, video_metadata)
                
                result = {
                    "success": True,
                    "chunks": chunks,
                    "metadata": {
                        **video_metadata,
                        "extraction_method": "whisper_local",
                        "word_count": len(transcript.get("text", "").split()),
                    },
                    "file_type": "youtube",
                    "processor": "YouTubeProcessor",
                }
                
                return result
                
            finally:
                # Clean up audio file
                if audio_path.exists():
                    audio_path.unlink()
            
        except Exception as e:
            logger.error(f"YouTube processing error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "chunks": [],
                "metadata": {"file_type": "youtube", "error": str(e)},
            }
    
    async def _download_audio(self, url: str) -> Path:
        """Download audio from YouTube video."""
        def download():
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp_file:
                output_path = Path(tmp_file.name)
            
            ydl_opts = {
                'format': 'bestaudio/best',
                'outtmpl': str(output_path.with_suffix('.%(ext)s')),
                'postprocessors': [{
                    'key': 'FFmpegExtractAudio',
                    'preferredcodec': 'wav',
                    'preferredquality': '192',
                }],
                'quiet': True,
                'no_warnings': True,
            }
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
            
            return output_path.with_suffix('.wav')
        
        return await self._run_in_thread(download)
    
    async def _get_video_metadata(self, url: str) -> Dict[str, Any]:
        """Extract video metadata."""
        def extract_metadata():
            ydl_opts = {
                'quiet': True,
                'no_warnings': True,
            }
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                return {
                    "title": info.get('title', ''),
                    "description": info.get('description', ''),
                    "duration": info.get('duration', 0),
                    "uploader": info.get('uploader', ''),
                    "upload_date": info.get('upload_date', ''),
                    "view_count": info.get('view_count', 0),
                    "url": url,
                }
        
        return await self._run_in_thread(extract_metadata)
    
    async def _transcribe_with_local_whisper(self, audio_path: Path) -> Dict[str, Any]:
        """Transcribe using local Whisper model."""
        def transcribe():
            if self._whisper_model is None:
                self._whisper_model = whisper.load_model("base")
            
            result = self._whisper_model.transcribe(
                str(audio_path),
                word_timestamps=True
            )
            return result
        
        return await self._run_in_thread(transcribe)
    
    async def _transcribe_with_openai(self, audio_path: Path) -> Dict[str, Any]:
        """Transcribe using OpenAI Whisper API."""
        with open(audio_path, 'rb') as audio_file:
            transcript = await self.openai_client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="verbose_json",
                timestamp_granularities=["segment"]
            )
        
        return {
            "text": transcript.text,
            "segments": getattr(transcript, 'segments', [])
        }
    
    def _create_transcript_chunks(
        self,
        transcript: Dict[str, Any],
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Create searchable chunks from transcript."""
        chunks = []
        
        # If we have segments with timestamps, create time-based chunks
        if 'segments' in transcript and transcript['segments']:
            for i, segment in enumerate(transcript['segments']):
                content = segment.get('text', '').strip()
                if content:
                    chunks.append({
                        "content": content,
                        "metadata": self._create_chunk_metadata(
                            source=metadata.get("url", "youtube"),
                            chunk_index=i,
                            start_time=segment.get('start', 0),
                            end_time=segment.get('end', 0),
                            title=metadata.get("title", ""),
                            uploader=metadata.get("uploader", ""),
                            content_type="transcript_segment",
                        ),
                    })
        else:
            # Fall back to simple text chunking
            text = transcript.get('text', '')
            if text:
                # Split into chunks of roughly 750 words (5 minutes of speech)
                words = text.split()
                chunk_size = 750
                
                for i in range(0, len(words), chunk_size):
                    chunk_words = words[i:i + chunk_size]
                    chunk_text = " ".join(chunk_words)
                    
                    chunks.append({
                        "content": chunk_text,
                        "metadata": self._create_chunk_metadata(
                            source=metadata.get("url", "youtube"),
                            chunk_index=i // chunk_size,
                            title=metadata.get("title", ""),
                            uploader=metadata.get("uploader", ""),
                            content_type="transcript_chunk",
                        ),
                    })
        
        return chunks

### src/documents/service.py
"""Document processing service."""

import asyncio
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any
import aiofiles
import time

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.documents.models import Document, DocumentChunk
from src.documents.processors.pdf_processor import PDFProcessor
from src.documents.processors.docx_processor import DocxProcessor
from src.documents.processors.youtube_processor import YouTubeProcessor
from src.documents.schemas import DocumentResponse, DocumentProcessingResult
from src.core.exceptions import NotFoundError, ProcessingError, ValidationError
from src.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentService:
    """Document processing service."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.processors = {
            "pdf": PDFProcessor(),
            "docx": DocxProcessor(),
            "youtube": YouTubeProcessor(),
        }
    
    async def upload_document(
        self,
        user_id: int,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None,
    ) -> Document:
        """Upload and store document."""
        # Validate file
        file_type = self._get_file_type(filename, content_type)
        if not self._is_supported_type(file_type):
            raise ValidationError(f"Unsupported file type: {file_type}")
        
        # Create document record
        document = Document(
            user_id=user_id,
            filename=self._sanitize_filename(filename),
            original_filename=filename,
            file_type=file_type,
            file_size=len(file_content),
            status="pending",
        )
        
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)
        
        # Save file
        file_path = await self._save_file(document.id, file_content, filename)
        document.file_path = str(file_path)
        await self.db.commit()
        
        logger.info(f"Uploaded document: {filename} (ID: {document.id})")
        return document
    
    async def process_document(self, document_id: int) -> DocumentProcessingResult:
        """Process document and extract content."""
        document = await self._get_document(document_id)
        if not document:
            raise NotFoundError("Document not found")
        
        if document.status == "completed":
            return DocumentProcessingResult(
                success=True,
                document_id=document_id,
                message="Document already processed",
            )
        
        try:
            document.status = "processing"
            await self.db.commit()
            
            start_time = time.time()
            
            # Process based on file type
            if document.file_type == "youtube":
                # For YouTube, we need the URL (stored in metadata or filename)
                result = await self.processors["youtube"].process(document.original_filename)
            else:
                # For file-based documents
                file_path = Path(document.file_path)
                if not file_path.exists():
                    raise ProcessingError(f"File not found: {file_path}")
                
                processor = self._get_processor(document.file_type)
                result = await processor.process(file_path)
            
            processing_time = time.time() - start_time
            
            if result["success"]:
                # Store extracted content
                document.content = result["chunks"][0]["content"] if result["chunks"] else ""
                document.metadata = str(result["metadata"])
                document.page_count = result["metadata"].get("pages", len(result["chunks"]))
                document.word_count = result["metadata"].get("word_count", 0)
                document.processing_time = processing_time
                document.status = "completed"
                
                # Create chunks
                chunks_created = await self._create_chunks(document_id, result["chunks"])
                
                await self.db.commit()
                
                logger.info(f"Processed document {document_id}: {chunks_created} chunks created")
                
                return DocumentProcessingResult(
                    success=True,
                    document_id=document_id,
                    message=f"Document processed successfully",
                    chunks_created=chunks_created,
                    processing_time=processing_time,
                )
            else:
                document.status = "failed"
                document.metadata = str({"error": result.get("error", "Unknown error")})
                await self.db.commit()
                
                return DocumentProcessingResult(
                    success=False,
                    document_id=document_id,
                    message=f"Processing failed: {result.get('error', 'Unknown error')}",
                    errors=[result.get("error", "Unknown error")],
                )
        
        except Exception as e:
            document.status = "failed"
            document.metadata = str({"error": str(e)})
            await self.db.commit()
            
            logger.error(f"Document processing failed: {str(e)}")
            raise ProcessingError(f"Processing failed: {str(e)}")
    
    async def add_youtube_video(self, user_id: int, url: str, title: Optional[str] = None) -> Document:
        """Add YouTube video for processing."""
        # Validate YouTube URL
        if not self._is_valid_youtube_url(url):
            raise ValidationError("Invalid YouTube URL")
        
        # Create document record
        document = Document(
            user_id=user_id,
            filename=title or f"YouTube Video - {url.split('/')[-1]}",
            original_filename=url,  # Store URL as filename
            file_type="youtube",
            status="pending",
        )
        
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)
        
        logger.info(f"Added YouTube video: {url} (ID: {document.id})")
        return document
    
    async def get_user_documents(self, user_id: int) -> List[DocumentResponse]:
        """Get all documents for a user."""
        result = await self.db.execute(
            select(Document).where(Document.user_id == user_id).order_by(Document.created_at.desc())
        )
        documents = result.scalars().all()
        return [DocumentResponse.model_validate(doc) for doc in documents]
    
    async def get_document(self, document_id: int, user_id: int) -> Optional[DocumentResponse]:
        """Get specific document."""
        result = await self.db.execute(
            select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id
            )
        )
        document = result.scalar_one_or_none()
        return DocumentResponse.model_validate(document) if document else None
    
    async def delete_document(self, document_id: int, user_id: int) -> bool:
        """Delete document and its chunks."""
        document = await self._get_document(document_id, user_id)
        if not document:
            return False
        
        # Delete file if it exists
        if document.file_path:
            file_path = Path(document.file_path)
            if file_path.exists():
                file_path.unlink()
        
        # Delete chunks
        await self.db.execute(
            delete(DocumentChunk).where(DocumentChunk.document_id == document_id)
        )
        
        # Delete document
        await self.db.delete(document)
        await self.db.commit()
        
        logger.info(f"Deleted document {document_id}")
        return True
    
    # Private methods
    
    async def _get_document(self, document_id: int, user_id: Optional[int] = None) -> Optional[Document]:
        """Get document by ID."""
        query = select(Document).where(Document.id == document_id)
        if user_id is not None:
            query = query.where(Document.user_id == user_id)
        
        result = await self.db.execute(query)
        return result.scalar_one_or_none()
    
    async def _save_file(self, document_id: int, content: bytes, filename: str) -> Path:
        """Save uploaded file to disk."""
        upload_dir = Path(settings.UPLOAD_DIR)
        upload_dir.mkdir(exist_ok=True)
        
        # Create unique filename
        file_path = upload_dir / f"{document_id}_{self._sanitize_filename(filename)}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        
        return file_path
    
    async def _create_chunks(self, document_id: int, chunks: List[Dict[str, Any]]) -> int:
        """Create document chunks in database."""
        chunk_objects = []
        
        for i, chunk in enumerate(chunks):
            chunk_obj = DocumentChunk(
                document_id=document_id,
                chunk_index=i,
                content=chunk["content"],
                metadata=str(chunk.get("metadata", {})),
                start_char=chunk.get("metadata", {}).get("start_char"),
                end_char=chunk.get("metadata", {}).get("end_char"),
                page_number=chunk.get("metadata", {}).get("page"),
            )
            chunk_objects.append(chunk_obj)
        
        self.db.add_all(chunk_objects)
        await self.db.commit()
        
        return len(chunk_objects)
    
    def _get_file_type(self, filename: str, content_type: Optional[str] = None) -> str:
        """Determine file type from filename and content type."""
        if content_type:
            return content_type
        
        # Guess from filename
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            return mime_type
        
        # Fall back to extension
        ext = Path(filename).suffix.lower()
        ext_mapping = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
        }
        
        return ext_mapping.get(ext, "application/octet-stream")
    
    def _is_supported_type(self, file_type: str) -> bool:
        """Check if file type is supported."""
        supported_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain",
            "youtube",
        ]
        return file_type in supported_types
    
    def _get_processor(self, file_type: str):
        """Get appropriate processor for file type."""
        processor_mapping = {
            "application/pdf": "pdf",
            "pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "docx": "docx",
            "youtube": "youtube",
        }
        
        processor_type = processor_mapping.get(file_type)
        if not processor_type or processor_type not in self.processors:
            raise ProcessingError(f"No processor available for file type: {file_type}")
        
        return self.processors[processor_type]
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage."""
        import re
        # Remove potentially dangerous characters
        sanitized = re.sub(r'[^\w\-_\. ]', '', filename)
        return sanitized[:100]  # Limit length
    
    def _is_valid_youtube_url(self, url: str) -> bool:
        """Validate YouTube URL."""
        import re
        youtube_patterns = [
            r'(?:https?://)?(?:www\.)?youtube\.com/watch\?v=([a-zA-Z0-9_-]{11})',
            r'(?:https?://)?(?:www\.)?youtu\.be/([a-zA-Z0-9_-]{11})',
            r'(?:https?://)?(?:www\.)?youtube\.com/embed/([a-zA-Z0-9_-]{11})',
        ]
        
        for pattern in youtube_patterns:
            if re.match(pattern, url.strip()):
                return True
        return False

### src/documents/router.py
"""Document processing routes."""

import asyncio
from typing import List

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.models import User
from src.core.database import get_db_session
from src.core.deps import get_current_user, get_default_user
from src.core.exceptions import NotFoundError, ProcessingError, ValidationError
from src.documents.schemas import (
    DocumentResponse,
    DocumentProcessingResult,
    YouTubeUpload,
    DocumentStats,
)
from src.documents.service import DocumentService

router = APIRouter()


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    user: User = Depends(get_default_user),  # Use default user for now
    db: AsyncSession = Depends(get_db_session),
):
    """Upload a document."""
    try:
        # Read file content
        content = await file.read()
        
        # Upload document
        document_service = DocumentService(db)
        document = await document_service.upload_document(
            user_id=user.id,
            file_content=content,
            filename=file.filename,
            content_type=file.content_type,
        )
        
        # Start processing in background
        asyncio.create_task(process_document_background(document.id, db))
        
        return DocumentResponse.model_validate(document)
        
    except ValidationError as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=e.message)


@router.post("/youtube", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def add_youtube_video(
    video_data: YouTubeUpload,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Add YouTube video for processing."""
    try:
        document_service = DocumentService(db)
        document = await document_service.add_youtube_video(
            user_id=user.id,
            url=video_data.url,
            title=video_data.title,
        )
        
        # Start processing in background
        asyncio.create_task(process_document_background(document.id, db))
        
        return DocumentResponse.model_validate(document)
        
    except ValidationError as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=e.message)


@router.post("/process/{document_id}", response_model=DocumentProcessingResult)
async def process_document(
    document_id: int,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Process a document."""
    try:
        document_service = DocumentService(db)
        result = await document_service.process_document(document_id)
        return result
        
    except (NotFoundError, ProcessingError) as e:
        raise HTTPException(status_code=e.status_code, detail=e.message)


@router.get("/", response_model=List[DocumentResponse])
async def get_documents(
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get all user documents."""
    document_service = DocumentService(db)
    return await document_service.get_user_documents(user.id)


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get specific document."""
    document_service = DocumentService(db)
    document = await document_service.get_document(document_id, user.id)
    
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    
    return document


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: int,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Delete document."""
    document_service = DocumentService(db)
    success = await document_service.delete_document(document_id, user.id)
    
    if not success:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")


async def process_document_background(document_id: int, db: AsyncSession):
    """Background task to process document."""
    try:
        document_service = DocumentService(db)
        await document_service.process_document(document_id)
    except Exception as e:
        # Log error but don't raise (this is a background task)
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Background processing failed for document {document_id}: {str(e)}")

## 5. RAG SYSTEM

### src/rag/__init__.py
"""Retrieval-Augmented Generation system."""

### src/rag/models.py
"""RAG-related models."""

from datetime import datetime
from typing import Optional

from sqlalchemy import DateTime, ForeignKey, Integer, String, Text, Float
from sqlalchemy.orm import Mapped, mapped_column
from sqlalchemy.sql import func

from src.core.database import Base


class VectorEmbedding(Base):
    """Vector embedding model."""
    
    __tablename__ = "vector_embeddings"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    document_id: Mapped[int] = mapped_column(Integer, ForeignKey("documents.id"))
    chunk_id: Mapped[int] = mapped_column(Integer, ForeignKey("document_chunks.id"))
    
    # Vector information
    vector_id: Mapped[str] = mapped_column(String(255), unique=True, index=True)  # ChromaDB ID
    embedding_model: Mapped[str] = mapped_column(String(100))
    
    # Metadata
    content_preview: Mapped[str] = mapped_column(Text)  # First 200 chars for reference
    
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )

### src/rag/schemas.py
"""RAG schemas."""

from typing import Dict, List, Any, Optional

from pydantic import BaseModel, Field


class SearchQuery(BaseModel):
    """Search query schema."""
    query: str = Field(..., min_length=1, max_length=1000)
    user_id: Optional[int] = None
    document_ids: Optional[List[int]] = None
    limit: int = Field(5, ge=1, le=20)
    similarity_threshold: float = Field(0.7, ge=0.0, le=1.0)


class SearchResult(BaseModel):
    """Search result schema."""
    content: str
    score: float
    metadata: Dict[str, Any]
    document_id: Optional[int] = None
    chunk_id: Optional[int] = None


class SearchResponse(BaseModel):
    """Search response schema."""
    query: str
    results: List[SearchResult]
    total_results: int
    processing_time: float


class EmbeddingRequest(BaseModel):
    """Embedding request schema."""
    texts: List[str]
    model: str = "text-embedding-3-large"


class EmbeddingResponse(BaseModel):
    """Embedding response schema."""
    embeddings: List[List[float]]
    model: str
    total_tokens: int

### src/rag/chunking.py
"""Text chunking strategies."""

import logging
from typing import Dict, List, Any
import re

logger = logging.getLogger(__name__)


class TextChunker:
    """Text chunking utility."""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks."""
        if not text or not text.strip():
            return []
        
        # Clean text
        text = self._clean_text(text)
        
        # Choose chunking strategy based on content
        if self._is_structured_content(text):
            chunks = self._chunk_by_structure(text)
        else:
            chunks = self._chunk_by_sentences(text)
        
        # Add metadata to chunks
        result = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                "chunk_index": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk),
                "word_count": len(chunk.split()),
                **(metadata or {}),
            }
            
            result.append({
                "content": chunk,
                "metadata": chunk_metadata,
            })
        
        return result
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _is_structured_content(self, text: str) -> bool:
        """Check if text has clear structure (headers, lists, etc.)."""
        structure_indicators = [
            r'^#{1,6}\s',  # Markdown headers
            r'^\d+\.\s',   # Numbered lists
            r'^[•\-\*]\s', # Bullet points
            r'^[A-Z][^.!?]*:
,  # Section headers ending with colon
        ]
        
        lines = text.split('\n')
        structured_lines = 0
        
        for line in lines:
            line = line.strip()
            if line:
                for pattern in structure_indicators:
                    if re.match(pattern, line, re.MULTILINE):
                        structured_lines += 1
                        break
        
        # If more than 10% of lines show structure, use structural chunking
        return structured_lines / max(len([l for l in lines if l.strip()]), 1) > 0.1
    
    def _chunk_by_structure(self, text: str) -> List[str]:
        """Chunk text by preserving structure."""
        chunks = []
        current_chunk = ""
        
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph would exceed chunk size
            if len(current_chunk) + len(paragraph) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                current_chunk = self._get_overlap_text(current_chunk) + paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _chunk_by_sentences(self, text: str) -> List[str]:
        """Chunk text by sentences with overlap."""
        # Split into sentences
        sentences = self._split_sentences(text)
        
        chunks = []
        current_chunk = ""
        current_sentences = []
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            test_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(test_chunk) > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(current_sentences)
                current_chunk = " ".join(overlap_sentences + [sentence])
                current_sentences = overlap_sentences + [sentence]
            else:
                current_chunk = test_chunk
                current_sentences.append(sentence)
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting (could be improved with spaCy/NLTK)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from end of chunk."""
        if len(text) <= self.overlap:
            return text
        
        # Try to find a good break point
        overlap_text = text[-self.overlap:]
        
        # Find the start of a sentence in the overlap
        sentences = self._split_sentences(overlap_text)
        if len(sentences) > 1:
            return " ".join(sentences[1:])
        
        return overlap_text
    
    def _get_overlap_sentences(self, sentences: List[str]) -> List[str]:
        """Get overlap sentences for new chunk."""
        if not sentences:
            return []
        
        # Calculate how many sentences to include for overlap
        overlap_text = ""
        overlap_sentences = []
        
        for sentence in reversed(sentences):
            test_overlap = sentence + " " + overlap_text if overlap_text else sentence
            if len(test_overlap) <= self.overlap:
                overlap_text = test_overlap
                overlap_sentences.insert(0, sentence)
            else:
                break
        
        return overlap_sentences


class AdaptiveChunker(TextChunker):
    """Adaptive chunker that adjusts strategy based on content type."""
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Adaptively chunk text based on content characteristics."""
        if not text or not text.strip():
            return []
        
        # Determine content type
        content_type = self._determine_content_type(text, metadata or {})
        
        # Adjust chunking parameters based on content type
        if content_type == "code":
            self.chunk_size = 800  # Smaller chunks for code
            self.overlap = 100
        elif content_type == "academic":
            self.chunk_size = 1200  # Larger chunks for academic content
            self.overlap = 300
        elif content_type == "transcript":
            self.chunk_size = 1500  # Large chunks for transcripts
            self.overlap = 150
        else:
            self.chunk_size = 1000  # Default
            self.overlap = 200
        
        return super().chunk_text(text, metadata)
    
    def _determine_content_type(self, text: str, metadata: Dict[str, Any]) -> str:
        """Determine content type from text and metadata."""
        # Check metadata first
        if metadata.get("content_type") == "table":
            return "table"
        if metadata.get("file_type") == "youtube":
            return "transcript"
        
        # Analyze text content
        code_indicators = len(re.findall(r'[{}\[\]();]|def |class |import |from ', text))
        academic_indicators = len(re.findall(r'\b(abstract|introduction|methodology|conclusion|references)\b', text, re.IGNORECASE))
        
        if code_indicators > 10:
            return "code"
        elif academic_indicators > 2:
            return "academic"
        elif "transcript" in text.lower() or metadata.get("extraction_method") == "whisper":
            return "transcript"
        else:
            return "general"

### src/rag/embeddings.py
"""Embedding generation service."""

import asyncio
import logging
from typing import List, Dict, Any

from openai import AsyncOpenAI
from tenacity import retry, stop_after_attempt, wait_exponential

from src.config import get_settings
from src.core.exceptions import ExternalServiceError

logger = logging.getLogger(__name__)
settings = get_settings()


class EmbeddingService:
    """Embedding generation service."""
    
    def __init__(self):
        self.client = AsyncOpenAI(
            api_key=settings.OPENAI_API_KEY,
            organization=settings.OPENAI_ORG_ID,
        )
        self.model = "text-embedding-3-large"
        self.batch_size = 100  # OpenAI's batch limit
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    async def generate_embeddings(self, texts: List[str]) -> Dict[str, Any]:
        """Generate embeddings for a list of texts."""
        if not texts:
            return {"embeddings": [], "usage": {"total_tokens": 0}}
        
        try:
            # Process in batches
            all_embeddings = []
            total_tokens = 0
            
            for i in range(0, len(texts), self.batch_size):
                batch = texts[i:i + self.batch_size]
                batch_result = await self._generate_batch_embeddings(batch)
                
                all_embeddings.extend(batch_result["embeddings"])
                total_tokens += batch_result["usage"]["total_tokens"]
            
            return {
                "embeddings": all_embeddings,
                "usage": {"total_tokens": total_tokens},
                "model": self.model,
            }
            
        except Exception as e:
            logger.error(f"Embedding generation error: {str(e)}")
            raise ExternalServiceError(f"Failed to generate embeddings: {str(e)}", "OpenAI")
    
    async def _generate_batch_embeddings(self, texts: List[str]) -> Dict[str, Any]:
        """Generate embeddings for a batch of texts."""
        response = await self.client.embeddings.create(
            model=self.model,
            input=texts,
        )
        
        embeddings = [embedding.embedding for embedding in response.data]
        
        return {
            "embeddings": embeddings,
            "usage": response.usage.model_dump(),
        }
    
    async def generate_single_embedding(self, text: str) -> List[float]:
        """Generate embedding for a single text."""
        result = await self.generate_embeddings([text])
        return result["embeddings"][0]
    
    def preprocess_text_for_embedding(self, text: str) -> str:
        """Preprocess text for optimal embedding generation."""
        # Clean text
        text = text.strip()
        
        # Truncate if too long (OpenAI has token limits)
        # Rough approximation: 1 token ≈ 4 characters
        max_chars = 32000  # Conservative estimate for 8k token limit
        if len(text) > max_chars:
            text = text[:max_chars]
            # Try to end at a sentence boundary
            last_period = text.rfind('. ')
            if last_period > max_chars * 0.8:  # If we can find a period in the last 20%
                text = text[:last_period + 1]
        
        return text

### src/rag/vector_store.py
"""ChromaDB vector store integration."""

import asyncio
import logging
from typing import Dict, List, Any, Optional

import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

from src.config import get_settings
from src.rag.embeddings import EmbeddingService

logger = logging.getLogger(__name__)
settings = get_settings()


class VectorStore:
    """ChromaDB vector store wrapper."""
    
    def __init__(self):
        self.client = self._create_client()
        self.embedding_service = EmbeddingService()
        self.collection_name = "studhelper_documents"
        self.collection = None
        self._initialize_collection()
    
    def _create_client(self):
        """Create ChromaDB client."""
        try:
            # Try HTTP client first (production)
            client = chromadb.HttpClient(
                host=settings.CHROMA_HOST,
                port=settings.CHROMA_PORT,
            )
            client.heartbeat()  # Test connection
            logger.info("Connected to ChromaDB server")
            return client
        except Exception:
            # Fall back to persistent client (development)
            logger.info("Using persistent ChromaDB client")
            return chromadb.PersistentClient(
                path=settings.CHROMA_PERSIST_DIR,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=False,
                ),
            )
    
    def _initialize_collection(self):
        """Initialize or get collection."""
        try:
            self.collection = self.client.get_collection(
                name=self.collection_name,
            )
            logger.info(f"Connected to existing collection: {self.collection_name}")
        except Exception:
            # Create new collection
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"description": "StudHelper document embeddings"},
            )
            logger.info(f"Created new collection: {self.collection_name}")
    
    async def add_documents(
        self,
        documents: List[Dict[str, Any]],
        user_id: Optional[int] = None,
    ) -> List[str]:
        """Add documents to vector store."""
        if not documents:
            return []
        
        try:
            # Prepare data for ChromaDB
            ids = []
            texts = []
            metadatas = []
            
            for i, doc in enumerate(documents):
                doc_id = f"doc_{user_id}_{i}_{hash(doc['content'])}"
                ids.append(doc_id)
                texts.append(doc["content"])
                
                metadata = doc.get("metadata", {})
                if user_id:
                    metadata["user_id"] = user_id
                metadatas.append(metadata)
            
            # Generate embeddings
            embeddings_result = await self.embedding_service.generate_embeddings(texts)
            embeddings = embeddings_result["embeddings"]
            
            # Add to ChromaDB
            await asyncio.to_thread(
                self.collection.add,
                ids=ids,
                documents=texts,
                metadatas=metadatas,
                embeddings=embeddings,
            )
            
            logger.info(f"Added {len(documents)} documents to vector store")
            return ids
            
        except Exception as e:
            logger.error(f"Error adding documents to vector store: {str(e)}")
            raise
    
    async def similarity_search(
        self,
        query: str,
        k: int = 5,
        user_id: Optional[int] = None,
        document_ids: Optional[List[int]] = None,
        similarity_threshold: float = 0.7,
    ) -> List[Dict[str, Any]]:
        """Search for similar documents."""
        try:
            # Generate query embedding
            query_embedding = await self.embedding_service.generate_single_embedding(query)
            
            # Build where clause for filtering
            where = {}
            if user_id:
                where["user_id"] = user_id
            if document_ids:
                where["document_id"] = {"$in": document_ids}
            
            # Search
            results = await asyncio.to_thread(
                self.collection.query,
                query_embeddings=[query_embedding],
                n_results=k,
                where=where if where else None,
            )
            
            # Format results
            formatted_results = []
            for i in range(len(results["ids"][0])):
                distance = results["distances"][0][i]
                # Convert distance to similarity score (0-1)
                similarity = 1 - distance
                
                if similarity >= similarity_threshold:
                    formatted_results.append({
                        "id": results["ids"][0][i],
                        "content": results["documents"][0][i],
                        "metadata": results["metadatas"][0][i],
                        "score": similarity,
                    })
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"Vector search error: {str(e)}")
            raise
    
    async def delete_documents(self, ids: List[str]) -> bool:
        """Delete documents by IDs."""
        try:
            await asyncio.to_thread(
                self.collection.delete,
                ids=ids,
            )
            logger.info(f"Deleted {len(ids)} documents from vector store")
            return True
        except Exception as e:
            logger.error(f"Error deleting documents: {str(e)}")
            return False
    
    async def delete_user_documents(self, user_id: int) -> bool:
        """Delete all documents for a user."""
        try:
            await asyncio.to_thread(
                self.collection.delete,
                where={"user_id": user_id},
            )
            logger.info(f"Deleted all documents for user {user_id}")
            return True
        except Exception as e:
            logger.error(f"Error deleting user documents: {str(e)}")
            return False
    
    async def get_collection_stats(self) -> Dict[str, Any]:
        """Get collection statistics."""
        try:
            count = await asyncio.to_thread(self.collection.count)
            return {
                "total_documents": count,
                "collection_name": self.collection_name,
                "embedding_model": self.embedding_service.model,
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {str(e)}")
            return {"error": str(e)}

### src/rag/retrieval.py
"""Document retrieval service."""

import logging
from typing import Dict, List, Any, Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.documents.models import Document, DocumentChunk
from src.rag.vector_store import VectorStore
from src.rag.schemas import SearchQuery, SearchResult, SearchResponse
import time

logger = logging.getLogger(__name__)


class RetrievalService:
    """Document retrieval service."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.vector_store = VectorStore()
    
    async def search_documents(self, search_query: SearchQuery) -> SearchResponse:
        """Search documents using vector similarity."""
        start_time = time.time()
        
        try:
            # Perform vector search
            vector_results = await self.vector_store.similarity_search(
                query=search_query.query,
                k=search_query.limit,
                user_id=search_query.user_id,
                document_ids=search_query.document_ids,
                similarity_threshold=search_query.similarity_threshold,
            )
            
            # Enrich results with database information
            enriched_results = await self._enrich_results(vector_results)
            
            processing_time = time.time() - start_time
            
            return SearchResponse(
                query=search_query.query,
                results=enriched_results,
                total_results=len(enriched_results),
                processing_time=processing_time,
            )
            
        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            raise
    
    async def _enrich_results(self, vector_results: List[Dict[str, Any]]) -> List[SearchResult]:
        """Enrich vector search results with database information."""
        enriched_results = []
        
        for result in vector_results:
            metadata = result["metadata"]
            
            # Try to get document information
            document_id = metadata.get("document_id")
            chunk_id = metadata.get("chunk_id")
            
            if document_id:
                # Get document details
                doc_result = await self.db.execute(
                    select(Document).where(Document.id == document_id)
                )
                document = doc_result.scalar_one_or_none()
                
                if document:
                    metadata.update({
                        "document_filename": document.filename,
                        "document_type": document.file_type,
                        "document_created": document.created_at.isoformat(),
                    })
            
            enriched_results.append(SearchResult(
                content=result["content"],
                score=result["score"],
                metadata=metadata,
                document_id=document_id,
                chunk_id=chunk_id,
            ))
        
        return enriched_results
    
    async def get_context_for_query(
        self,
        query: str,
        user_id: Optional[int] = None,
        max_contexts: int = 5,
    ) -> str:
        """Get relevant context for a query."""
        search_query = SearchQuery(
            query=query,
            user_id=user_id,
            limit=max_contexts,
            similarity_threshold=0.6,  # Lower threshold for context
        )
        
        search_response = await self.search_documents(search_query)
        
        if not search_response.results:
            return ""
        
        # Combine results into context
        context_parts = []
        for i, result in enumerate(search_response.results):
            context_parts.append(f"[Context {i+1}]: {result.content}")
        
        return "\n\n".join(context_parts)
    
    async def get_relevant_sources(
        self,
        query: str,
        user_id: Optional[int] = None,
    ) -> List[Dict[str, Any]]:
        """Get relevant source documents for citations."""
        search_query = SearchQuery(
            query=query,
            user_id=user_id,
            limit=3,
            similarity_threshold=0.7,
        )
        
        search_response = await self.search_documents(search_query)
        
        sources = []
        seen_documents = set()
        
        for result in search_response.results:
            doc_id = result.document_id
            if doc_id and doc_id not in seen_documents:
                seen_documents.add(doc_id)
                sources.append({
                    "document_id": doc_id,
                    "filename": result.metadata.get("document_filename", "Unknown"),
                    "type": result.metadata.get("document_type", "unknown"),
                    "relevance_score": result.score,
                    "page": result.metadata.get("page"),
                })
        
        return sources

### src/rag/service.py
"""RAG orchestration service."""

import logging
from typing import Dict, List, Any, Optional

from sqlalchemy.ext.asyncio import AsyncSession

from src.documents.models import Document, DocumentChunk
from src.rag.chunking import AdaptiveChunker
from src.rag.vector_store import VectorStore
from src.rag.retrieval import RetrievalService
from src.rag.schemas import SearchQuery, SearchResponse

logger = logging.getLogger(__name__)


class RAGService:
    """RAG orchestration service."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.vector_store = VectorStore()
        self.retrieval_service = RetrievalService(db)
        self.chunker = AdaptiveChunker()
    
    async def process_document_for_rag(
        self,
        document_id: int,
        user_id: int,
        chunks: List[Dict[str, Any]],
    ) -> bool:
        """Process document chunks for RAG."""
        try:
            # Prepare chunks for vector storage
            vector_documents = []
            
            for chunk in chunks:
                # Add document and user information to metadata
                chunk_metadata = chunk.get("metadata", {})
                chunk_metadata.update({
                    "document_id": document_id,
                    "user_id": user_id,
                })
                
                vector_documents.append({
                    "content": chunk["content"],
                    "metadata": chunk_metadata,
                })
            
            # Add to vector store
            vector_ids = await self.vector_store.add_documents(
                documents=vector_documents,
                user_id=user_id,
            )
            
            logger.info(f"Processed {len(vector_ids)} chunks for RAG (document {document_id})")
            return True
            
        except Exception as e:
            logger.error(f"RAG processing error for document {document_id}: {str(e)}")
            return False
    
    async def search_knowledge_base(
        self,
        query: str,
        user_id: Optional[int] = None,
        document_ids: Optional[List[int]] = None,
        limit: int = 5,
    ) -> SearchResponse:
        """Search user's knowledge base."""
        search_query = SearchQuery(
            query=query,
            user_id=user_id,
            document_ids=document_ids,
            limit=limit,
        )
        
        return await self.retrieval_service.search_documents(search_query)
    
    async def get_context_for_chat(
        self,
        query: str,
        user_id: Optional[int] = None,
        max_context_length: int = 4000,
    ) -> Dict[str, Any]:
        """Get context for chat AI."""
        # Get relevant context
        context = await self.retrieval_service.get_context_for_query(
            query=query,
            user_id=user_id,
        )
        
        # Truncate if too long
        if len(context) > max_context_length:
            context = context[:max_context_length]
            # Try to end at a complete context block
            last_context_end = context.rfind("[Context ")
            if last_context_end > max_context_length * 0.7:
                context = context[:last_context_end]
        
        # Get sources for citations
        sources = await self.retrieval_service.get_relevant_sources(
            query=query,
            user_id=user_id,
        )
        
        return {
            "context": context,
            "sources": sources,
            "has_context": bool(context.strip()),
        }
    
    async def remove_document_from_rag(self, document_id: int, user_id: int) -> bool:
        """Remove document from RAG system."""
        try:
            # This would require tracking vector IDs by document
            # For now, we'll just log the request
            logger.info(f"Request to remove document {document_id} from RAG")
            return True
        except Exception as e:
            logger.error(f"Error removing document from RAG: {str(e)}")
            return False

### src/rag/router.py
"""RAG API routes."""

from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.models import User
from src.core.database import get_db_session
from src.core.deps import get_default_user
from src.rag.schemas import SearchQuery, SearchResponse
from src.rag.service import RAGService

router = APIRouter()


@router.post("/search", response_model=SearchResponse)
async def search_documents(
    search_query: SearchQuery,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Search documents using semantic similarity."""
    rag_service = RAGService(db)
    
    # Use current user if not specified in query
    if search_query.user_id is None:
        search_query.user_id = user.id
    
    return await rag_service.search_knowledge_base(
        query=search_query.query,
        user_id=search_query.user_id,
        document_ids=search_query.document_ids,
        limit=search_query.limit,
    )


@router.get("/context/{query}")
async def get_context_for_query(
    query: str,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get context for a specific query."""
    rag_service = RAGService(db)
    
    context_data = await rag_service.get_context_for_chat(
        query=query,
        user_id=user.id,
    )
    
    return context_data


@router.get("/stats")
async def get_rag_stats(
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get RAG system statistics."""
    rag_service = RAGService(db)
    
    # Get vector store stats
    stats = await rag_service.vector_store.get_collection_stats()
    
    return stats

## 6. AI INTEGRATION SYSTEM

### src/ai/__init__.py
"""AI integration system."""

### src/ai/models.py
"""AI-related models."""

from datetime import datetime
from typing import Optional

from sqlalchemy import DateTime, ForeignKey, Integer, String, Text, Float, JSON
from sqlalchemy.orm import Mapped, mapped_column
from sqlalchemy.sql import func

from src.core.database import Base


class ChatSession(Base):
    """Chat session model."""
    
    __tablename__ = "chat_sessions"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    user_id: Mapped[int] = mapped_column(Integer, ForeignKey("users.id"))
    
    # Session information
    title: Mapped[Optional[str]] = mapped_column(String(255))
    mode: Mapped[str] = mapped_column(String(50), default="economic")  # economic, standard, turbo
    
    # Statistics
    message_count: Mapped[int] = mapped_column(Integer, default=0)
    total_tokens: Mapped[int] = mapped_column(Integer, default=0)
    total_cost: Mapped[float] = mapped_column(Float, default=0.0)
    
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )
    updated_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now(), onupdate=func.now()
    )


class ChatMessage(Base):
    """Chat message model."""
    
    __tablename__ = "chat_messages"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    session_id: Mapped[int] = mapped_column(Integer, ForeignKey("chat_sessions.id"))
    user_id: Mapped[int] = mapped_column(Integer, ForeignKey("users.id"))
    
    # Message content
    role: Mapped[str] = mapped_column(String(20))  # user, assistant, system
    content: Mapped[str] = mapped_column(Text)
    
    # AI response metadata
    model: Mapped[Optional[str]] = mapped_column(String(100))
    input_tokens: Mapped[Optional[int]] = mapped_column(Integer)
    output_tokens: Mapped[Optional[int]] = mapped_column(Integer)
    cost: Mapped[Optional[float]] = mapped_column(Float)
    
    # Context and sources
    context_used: Mapped[Optional[str]] = mapped_column(Text)
    sources: Mapped[Optional[str]] = mapped_column(Text)  # JSON array
    
    # Feedback
    rating: Mapped[Optional[int]] = mapped_column(Integer)  # 1-5
    feedback: Mapped[Optional[str]] = mapped_column(Text)
    
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )

### src/ai/schemas.py
"""AI schemas."""

from datetime import datetime
from typing import Dict, List, Any, Optional
from enum import Enum

from pydantic import BaseModel, Field


class ChatMode(str, Enum):
    """Chat mode enumeration."""
    ECONOMIC = "economic"
    STANDARD = "standard"
    TURBO = "turbo"


class ChatMessageRequest(BaseModel):
    """Chat message request schema."""
    message: str = Field(..., min_length=1, max_length=5000)
    mode: ChatMode = ChatMode.ECONOMIC
    session_id: Optional[int] = None
    include_context: bool = True
    max_context_length: int = Field(4000, ge=1000, le=8000)


class ChatMessageResponse(BaseModel):
    """Chat message response schema."""
    content: str
    role: str = "assistant"
    mode: str
    model: str
    session_id: int
    
    # Usage information
    input_tokens: int
    output_tokens: int
    total_tokens: int
    cost: float
    
    # Context information
    context_used: bool
    sources: List[Dict[str, Any]]
    
    # Metadata
    response_time: float
    timestamp: datetime


class ChatSessionResponse(BaseModel):
    """Chat session response schema."""
    id: int
    title: Optional[str]
    mode: str
    message_count: int
    total_tokens: int
    total_cost: float
    created_at: datetime
    updated_at: datetime
    
    class Config:
        from_attributes = True


class StreamingChunk(BaseModel):
    """Streaming response chunk."""
    delta: str
    session_id: Optional[int] = None
    finished: bool = False
    
    # Final usage info (only when finished=True)
    usage: Optional[Dict[str, int]] = None
    cost: Optional[float] = None


class MessageFeedback(BaseModel):
    """Message feedback schema."""
    rating: int = Field(..., ge=1, le=5)
    feedback: Optional[str] = None

### src/ai/cost_calculator.py
"""Token counting and cost calculation."""

import logging
from typing import Dict, List

import tiktoken

from src.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class CostCalculator:
    """Token counting and cost calculation service."""
    
    def __init__(self):
        self.encodings = {
            'gpt-4o-mini': tiktoken.encoding_for_model('gpt-4o-mini'),
            'gpt-4o': tiktoken.encoding_for_model('gpt-4o'),
        }
        
        self.pricing = {
            'economic': {
                'input': settings.ECONOMIC_INPUT_PRICE / 1_000_000,  # Per token
                'output': settings.ECONOMIC_OUTPUT_PRICE / 1_000_000,
                'model': 'gpt-4o-mini',
            },
            'standard': {
                'input': settings.STANDARD_INPUT_PRICE / 1_000_000,
                'output': settings.STANDARD_OUTPUT_PRICE / 1_000_000,
                'model': 'gpt-4o',
            },
            'turbo': {
                'input': settings.STANDARD_INPUT_PRICE / 1_000_000,
                'output': settings.STANDARD_OUTPUT_PRICE / 1_000_000,
                'model': 'gpt-4o',
            },
        }
    
    def count_tokens(self, messages: List[Dict[str, str]], model: str) -> int:
        """Count tokens in messages using tiktoken."""
        encoding = self.encodings.get(model, self.encodings['gpt-4o-mini'])
        
        tokens = 0
        for message in messages:
            tokens += 3  # Message formatting tokens
            for key, value in message.items():
                tokens += len(encoding.encode(str(value)))
                if key == "name":
                    tokens -= 1
        tokens += 3  # Assistant reply priming
        return tokens
    
    def count_string_tokens(self, text: str, model: str) -> int:
        """Count tokens in a string."""
        encoding = self.encodings.get(model, self.encodings['gpt-4o-mini'])
        return len(encoding.encode(text))
    
    def calculate_cost(
        self,
        mode: str,
        input_tokens: int,
        output_tokens: int,
    ) -> Dict[str, float]:
        """Calculate cost for token usage."""
        if mode not in self.pricing:
            raise ValueError(f"Unknown mode: {mode}")
        
        pricing = self.pricing[mode]
        
        input_cost = input_tokens * pricing['input']
        output_cost = output_tokens * pricing['output']
        total_cost = input_cost + output_cost
        
        return {
            'input_cost': input_cost,
            'output_cost': output_cost,
            'total_cost': total_cost,
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': input_tokens + output_tokens,
            'mode': mode,
            'model': pricing['model'],
        }
    
    def estimate_cost(self, mode: str, estimated_tokens: int) -> float:
        """Estimate cost for a given number of tokens."""
        if mode not in self.pricing:
            return 0.0
        
        pricing = self.pricing[mode]
        # Use average of input and output pricing for estimation
        avg_price = (pricing['input'] + pricing['output']) / 2
        return estimated_tokens * avg_price
    
    def get_mode_info(self, mode: str) -> Dict[str, Any]:
        """Get information about a specific mode."""
        if mode not in self.pricing:
            return {}
        
        pricing = self.pricing[mode]
        return {
            'mode': mode,
            'model': pricing['model'],
            'input_price_per_1m': pricing['input'] * 1_000_000,
            'output_price_per_1m': pricing['output'] * 1_000_000,
            'description': self._get_mode_description(mode),
        }
    
    def _get_mode_description(self, mode: str) -> str:
        """Get description for a mode."""
        descriptions = {
            'economic': 'Fast, cost-effective responses for basic questions',
            'standard': 'Detailed explanations with examples and context',
            'turbo': 'Advanced reasoning with step-by-step thinking',
        }
        return descriptions.get(mode, '')

### src/ai/prompt_templates.py
"""AI prompt templates."""

from typing import Dict, Any, Optional


class PromptTemplates:
    """AI prompt templates for different modes and scenarios."""
    
    @staticmethod
    def get_system_prompt(mode: str, context: Optional[str] = None) -> str:
        """Get system prompt based on mode and context."""
        base_prompt = """You are StudHelper AI, an expert study assistant. Your goal is to help students learn effectively by providing clear, accurate, and pedagogically sound explanations.

You are helpful, knowledgeable, and focused on education. Always aim to teach concepts rather than just provide answers."""
        
        if context:
            base_prompt += f"\n\nYou have access to the following relevant information from the student's materials:\n\n{context}\n\nUse this information to provide accurate, contextual answers. If the provided context doesn't contain enough information to fully answer the question, acknowledge this and provide what help you can based on your general knowledge."
        
        mode_prompts = {
            'economic': """

Provide concise, focused responses that directly address the student's question. Be efficient with your explanations while ensuring accuracy.""",
            
            'standard': """

Provide comprehensive explanations with examples and context to help students understand deeply. Include relevant background information and practical applications where appropriate.""",
            
            'turbo': """

For complex questions, use step-by-step reasoning:
1. Analyze what the question is asking
2. Break down the problem into components  
3. Work through each component systematically
4. Show your reasoning process clearly
5. Provide a comprehensive synthesis

Think through problems carefully and show your work to help students learn your reasoning process."""
        }
        
        return base_prompt + mode_prompts.get(mode, mode_prompts['economic'])
    
    @staticmethod
    def get_no_context_prompt(mode: str) -> str:
        """Get prompt when no context is available."""
        prompt = """You are StudHelper AI, an expert study assistant. The student hasn't uploaded any specific materials, so provide helpful general educational assistance based on your knowledge.

Focus on teaching concepts, providing clear explanations, and helping the student learn effectively."""
        
        if mode == 'turbo':
            prompt += " For complex topics, break down your explanation step by step."
        
        return prompt
    
    @staticmethod
    def get_context_instruction(has_context: bool) -> str:
        """Get instruction about context usage."""
        if has_context:
            return "Use the provided context from the student's materials to give accurate, specific answers. Cite relevant sections when helpful."
        else:
            return "The student hasn't uploaded materials related to this question, so provide general educational assistance."

### src/ai/openai_client.py
"""OpenAI client integration."""

import asyncio
import logging
from typing import List, Dict, Any, AsyncGenerator

from openai import AsyncOpenAI
from tenacity import retry, stop_after_attempt, wait_exponential

from src.config import get_settings
from src.core.exceptions import ExternalServiceError
from src.ai.cost_calculator import CostCalculator

logger = logging.getLogger(__name__)
settings = get_settings()


class OpenAIService:
    """OpenAI integration service."""
    
    def __init__(self):
        self.client = AsyncOpenAI(
            api_key=settings.OPENAI_API_KEY,
            organization=settings.OPENAI_ORG_ID,
        )
        self.cost_calculator = CostCalculator()
        
        self.model_mapping = {
            'economic': 'gpt-4o-mini',
            'standard': 'gpt-4o',
            'turbo': 'gpt-4o',
        }
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    async def create_completion(
        self,
        messages: List[Dict[str, str]],
        mode: str = "economic",
        stream: bool = False,
        **kwargs
    ) -> Any:
        """Create chat completion with automatic retry."""
        model = self.model_mapping.get(mode, 'gpt-4o-mini')
        
        try:
            response = await self.client.chat.completions.create(
                model=model,
                messages=messages,
                stream=stream,
                temperature=0.1 if mode == 'turbo' else 0.7,
                **kwargs
            )
            return response
        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            raise ExternalServiceError(f"AI service error: {str(e)}", "OpenAI")
    
    async def create_streaming_completion(
        self,
        messages: List[Dict[str, str]],
        mode: str = "economic",
        **kwargs
    ) -> AsyncGenerator[str, None]:
        """Create streaming chat completion."""
        model = self.model_mapping.get(mode, 'gpt-4o-mini')
        
        try:
            stream = await self.client.chat.completions.create(
                model=model,
                messages=messages,
                stream=True,
                temperature=0.1 if mode == 'turbo' else 0.7,
                **kwargs
            )
            
            async for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield chunk.choices[0].delta.content
                    
        except Exception as e:
            logger.error(f"Streaming error: {str(e)}")
            yield f"Error: {str(e)}"
    
    def count_tokens(self, messages: List[Dict[str, str]], mode: str) -> int:
        """Count tokens in messages."""
        model = self.model_mapping.get(mode, 'gpt-4o-mini')
        return self.cost_calculator.count_tokens(messages, model)
    
    def count_string_tokens(self, text: str, mode: str) -> int:
        """Count tokens in a string."""
        model = self.model_mapping.get(mode, 'gpt-4o-mini')
        return self.cost_calculator.count_string_tokens(text, model)

### src/ai/service.py
"""AI chat service."""

import asyncio
import logging
import time
from typing import List, Dict, Any, Optional, AsyncGenerator

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.ai.models import ChatSession, ChatMessage
from src.ai.schemas import ChatMode, ChatMessageRequest, ChatMessageResponse, StreamingChunk
from src.ai.openai_client import OpenAIService
from src.ai.prompt_templates import PromptTemplates
from src.ai.cost_calculator import CostCalculator
from src.rag.service import RAGService
from src.usage.tracker import UsageTracker

logger = logging.getLogger(__name__)


class ChatService:
    """AI chat service."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.openai_service = OpenAIService()
        self.cost_calculator = CostCalculator()
        self.rag_service = RAGService(db)
        self.usage_tracker = UsageTracker()
    
    async def chat(
        self,
        request: ChatMessageRequest,
        user_id: int,
        stream: bool = False,
    ) -> ChatMessageResponse:
        """Handle chat request."""
        start_time = time.time()
        
        try:
            # Get or create chat session
            session = await self._get_or_create_session(
                user_id=user_id,
                session_id=request.session_id,
                mode=request.mode.value,
            )
            
            # Get conversation history
            conversation_history = await self._get_conversation_history(session.id)
            
            # Get context from RAG if requested
            context_data = {}
            if request.include_context:
                context_data = await self.rag_service.get_context_for_chat(
                    query=request.message,
                    user_id=user_id,
                    max_context_length=request.max_context_length,
                )
            
            # Build messages for AI
            messages = self._build_messages(
                request=request,
                conversation_history=conversation_history,
                context_data=context_data,
            )
            
            # Count input tokens
            input_tokens = self.openai_service.count_tokens(messages, request.mode.value)
            
            # Generate AI response
            response = await self.openai_service.create_completion(
                messages=messages,
                mode=request.mode.value,
            )
            
            # Process response
            content = response.choices[0].message.content
            model = response.model
            
            # Count output tokens and calculate cost
            output_tokens = self.openai_service.count_string_tokens(content, request.mode.value)
            cost_info = self.cost_calculator.calculate_cost(
                mode=request.mode.value,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
            )
            
            # Save messages to database
            await self._save_messages(
                session_id=session.id,
                user_id=user_id,
                user_message=request.message,
                ai_response=content,
                model=model,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                cost=cost_info["total_cost"],
                context_data=context_data,
            )
            
            # Update session statistics
            await self._update_session_stats(
                session_id=session.id,
                tokens=input_tokens + output_tokens,
                cost=cost_info["total_cost"],
            )
            
            # Track usage
            await self.usage_tracker.track_usage(
                user_id=user_id,
                model=model,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                cost=cost_info["total_cost"],
                endpoint="chat",
            )
            
            response_time = time.time() - start_time
            
            return ChatMessageResponse(
                content=content,
                mode=request.mode.value,
                model=model,
                session_id=session.id,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                total_tokens=input_tokens + output_tokens,
                cost=cost_info["total_cost"],
                context_used=context_data.get("has_context", False),
                sources=context_data.get("sources", []),
                response_time=response_time,
                timestamp=time.time(),
            )
            
        except Exception as e:
            logger.error(f"Chat service error: {str(e)}")
            raise
    
    async def stream_chat(
        self,
        request: ChatMessageRequest,
        user_id: int,
    ) -> AsyncGenerator[StreamingChunk, None]:
        """Handle streaming chat request."""
        try:
            # Get or create session (similar to regular chat)
            session = await self._get_or_create_session(
                user_id=user_id,
                session_id=request.session_id,
                mode=request.mode.value,
            )
            
            # Get conversation history and context
            conversation_history = await self._get_conversation_history(session.id)
            context_data = {}
            if request.include_context:
                context_data = await self.rag_service.get_context_for_chat(
                    query=request.message,
                    user_id=user_id,
                    max_context_length=request.max_context_length,
                )
            
            # Build messages
            messages = self._build_messages(
                request=request,
                conversation_history=conversation_history,
                context_data=context_data,
            )
            
            # Count input tokens
            input_tokens = self.openai_service.count_tokens(messages, request.mode.value)
            
            # Stream response
            collected_content = ""
            async for chunk in self.openai_service.create_streaming_completion(
                messages=messages,
                mode=request.mode.value,
            ):
                collected_content += chunk
                yield StreamingChunk(
                    delta=chunk,
                    session_id=session.id,
                )
            
            # Calculate final costs and save
            output_tokens = self.openai_service.count_string_tokens(collected_content, request.mode.value)
            cost_info = self.cost_calculator.calculate_cost(
                mode=request.mode.value,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
            )
            
            # Save to database
            await self._save_messages(
                session_id=session.id,
                user_id=user_id,
                user_message=request.message,
                ai_response=collected_content,
                model=self.openai_service.model_mapping[request.mode.value],
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                cost=cost_info["total_cost"],
                context_data=context_data,
            )
            
            # Send final chunk with usage info
            yield StreamingChunk(
                delta="",
                session_id=session.id,
                finished=True,
                usage={
                    "input_tokens": input_tokens,
                    "output_tokens": output_tokens,
                    "total_tokens": input_tokens + output_tokens,
                },
                cost=cost_info["total_cost"],
            )
            
        except Exception as e:
            logger.error(f"Streaming chat error: {str(e)}")
            yield StreamingChunk(
                delta=f"Error: {str(e)}",
                finished=True,
            )
    
    async def get_chat_sessions(self, user_id: int) -> List[Dict[str, Any]]:
        """Get user's chat sessions."""
        result = await self.db.execute(
            select(ChatSession)
            .where(ChatSession.user_id == user_id)
            .order_by(ChatSession.updated_at.desc())
        )
        sessions = result.scalars().all()
        
        return [
            {
                "id": session.id,
                "title": session.title,
                "mode": session.mode,
                "message_count": session.message_count,
                "total_tokens": session.total_tokens,
                "total_cost": session.total_cost,
                "created_at": session.created_at,
                "updated_at": session.updated_at,
            }
            for session in sessions
        ]
    
    async def get_chat_history(self, session_id: int, user_id: int) -> List[Dict[str, Any]]:
        """Get chat history for a session."""
        result = await self.db.execute(
            select(ChatMessage)
            .where(
                ChatMessage.session_id == session_id,
                ChatMessage.user_id == user_id,
            )
            .order_by(ChatMessage.created_at.asc())
        )
        messages = result.scalars().all()
        
        return [
            {
                "id": message.id,
                "role": message.role,
                "content": message.content,
                "model": message.model,
                "input_tokens": message.input_tokens,
                "output_tokens": message.output_tokens,
                "cost": message.cost,
                "sources": message.sources,
                "created_at": message.created_at,
            }
            for message in messages
        ]
    
    # Private methods
    
    async def _get_or_create_session(
        self,
        user_id: int,
        session_id: Optional[int],
        mode: str,
    ) -> ChatSession:
        """Get existing session or create new one."""
        if session_id:
            result = await self.db.execute(
                select(ChatSession).where(
                    ChatSession.id == session_id,
                    ChatSession.user_id == user_id,
                )
            )
            session = result.scalar_one_or_none()
            if session:
                return session
        
        # Create new session
        session = ChatSession(
            user_id=user_id,
            mode=mode,
            title=f"Chat Session {mode.title()}",
        )
        
        self.db.add(session)
        await self.db.commit()
        await self.db.refresh(session)
        
        return session
    
    async def _get_conversation_history(self, session_id: int, limit: int = 10) -> List[Dict[str, str]]:
        """Get recent conversation history for context."""
        result = await self.db.execute(
            select(ChatMessage)
            .where(ChatMessage.session_id == session_id)
            .order_by(ChatMessage.created_at.desc())
            .limit(limit * 2)  # Get more to account for user/assistant pairs
        )
        messages = list(result.scalars().all())
        
        # Convert to OpenAI format and reverse to chronological order
        conversation = []
        for message in reversed(messages):
            conversation.append({
                "role": message.role,
                "content": message.content,
            })
        
        return conversation
    
    def _build_messages(
        self,
        request: ChatMessageRequest,
        conversation_history: List[Dict[str, str]],
        context_data: Dict[str, Any],
    ) -> List[Dict[str, str]]:
        """Build message list for OpenAI API."""
        # System prompt
        system_prompt = PromptTemplates.get_system_prompt(
            mode=request.mode.value,
            context=context_data.get("context"),
        )
        
        messages = [{"role": "system", "content": system_prompt}]
        
        # Add conversation history
        messages.extend(conversation_history)
        
        # Add current user message
        messages.append({"role": "user", "content": request.message})
        
        return messages
    
    async def _save_messages(
        self,
        session_id: int,
        user_id: int,
        user_message: str,
        ai_response: str,
        model: str,
        input_tokens: int,
        output_tokens: int,
        cost: float,
        context_data: Dict[str, Any],
    ):
        """Save user and AI messages to database."""
        # Save user message
        user_msg = ChatMessage(
            session_id=session_id,
            user_id=user_id,
            role="user",
            content=user_message,
        )
        
        # Save AI message
        ai_msg = ChatMessage(
            session_id=session_id,
            user_id=user_id,
            role="assistant",
            content=ai_response,
            model=model,
            input_tokens=input_tokens,
            output_tokens=output_tokens,
            cost=cost,
            context_used=context_data.get("context", ""),
            sources=str(context_data.get("sources", [])),
        )
        
        self.db.add_all([user_msg, ai_msg])
        await self.db.commit()
    
    async def _update_session_stats(self, session_id: int, tokens: int, cost: float):
        """Update session statistics."""
        result = await self.db.execute(
            select(ChatSession).where(ChatSession.id == session_id)
        )
        session = result.scalar_one()
        
        session.message_count += 2  # User + AI message
        session.total_tokens += tokens
        session.total_cost += cost
        
        await self.db.commit()

### src/ai/router.py
"""AI chat routes."""

from typing import List

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.models import User
from src.core.database import get_db_session
from src.core.deps import get_default_user
from src.ai.schemas import ChatMessageRequest, ChatMessageResponse, MessageFeedback
from src.ai.service import ChatService
from src.usage.tracker import UsageTracker

router = APIRouter()


@router.post("/chat", response_model=ChatMessageResponse)
async def chat(
    request: ChatMessageRequest,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Send a chat message."""
    # Check usage limits
    usage_tracker = UsageTracker()
    limits = await usage_tracker.check_user_limits(user.id)
    
    if not limits["within_limits"]:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="Daily usage limit exceeded",
        )
    
    chat_service = ChatService(db)
    return await chat_service.chat(request, user.id)


@router.post("/chat/stream")
async def stream_chat(
    request: ChatMessageRequest,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Send a chat message with streaming response."""
    # Check usage limits
    usage_tracker = UsageTracker()
    limits = await usage_tracker.check_user_limits(user.id)
    
    if not limits["within_limits"]:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="Daily usage limit exceeded",
        )
    
    chat_service = ChatService(db)
    
    async def generate():
        async for chunk in chat_service.stream_chat(request, user.id):
            yield f"data: {chunk.model_dump_json()}\n\n"
        yield "data: [DONE]\n\n"
    
    return StreamingResponse(generate(), media_type="text/plain")


@router.get("/sessions")
async def get_chat_sessions(
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get user's chat sessions."""
    chat_service = ChatService(db)
    return await chat_service.get_chat_sessions(user.id)


@router.get("/sessions/{session_id}/history")
async def get_chat_history(
    session_id: int,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get chat history for a session."""
    chat_service = ChatService(db)
    return await chat_service.get_chat_history(session_id, user.id)


@router.post("/messages/{message_id}/feedback")
async def submit_feedback(
    message_id: int,
    feedback: MessageFeedback,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Submit feedback for a message."""
    # Implementation would update the ChatMessage with feedback
    return {"message": "Feedback submitted successfully"}

## 7. USAGE TRACKING SYSTEM

### src/usage/__init__.py
"""Usage tracking system."""

### src/usage/models.py
"""Usage tracking models."""

from datetime import datetime, date
from typing import Optional

from sqlalchemy import DateTime, Date, ForeignKey, Integer, String, Float
from sqlalchemy.orm import Mapped, mapped_column
from sqlalchemy.sql import func

from src.core.database import Base


class UsageLog(Base):
    """Usage log model for detailed tracking."""
    
    __tablename__ = "usage_logs"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    user_id: Mapped[int] = mapped_column(Integer, ForeignKey("users.id"))
    
    # API usage details
    endpoint: Mapped[str] = mapped_column(String(100))  # chat, documents, etc.
    model: Mapped[str] = mapped_column(String(100))
    
    # Token usage
    input_tokens: Mapped[int] = mapped_column(Integer)
    output_tokens: Mapped[int] = mapped_column(Integer)
    total_tokens: Mapped[int] = mapped_column(Integer)
    
    # Cost
    cost: Mapped[float] = mapped_column(Float)
    
    # Session info
    session_id: Mapped[Optional[str]] = mapped_column(String(255))
    
    timestamp: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )


class UserQuota(Base):
    """User daily quotas and limits."""
    
    __tablename__ = "user_quotas"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    user_id: Mapped[int] = mapped_column(Integer, ForeignKey("users.id"))
    date: Mapped[date] = mapped_column(Date, index=True)
    
    # Daily usage
    tokens_used: Mapped[int] = mapped_column(Integer, default=0)
    cost_incurred: Mapped[float] = mapped_column(Float, default=0.0)
    requests_made: Mapped[int] = mapped_column(Integer, default=0)
    
    # Daily limits
    daily_token_limit: Mapped[int] = mapped_column(Integer, default=10000)
    daily_cost_limit: Mapped[float] = mapped_column(Float, default=1.00)
    daily_request_limit: Mapped[int] = mapped_column(Integer, default=100)
    
    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )
    updated_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now(), onupdate=func.now()
    )

### src/usage/schemas.py
"""Usage tracking schemas."""

from datetime import datetime, date
from typing import Dict, List, Any, Optional

from pydantic import BaseModel


class UsageLogResponse(BaseModel):
    """Usage log response schema."""
    id: int
    endpoint: str
    model: str
    input_tokens: int
    output_tokens: int
    total_tokens: int
    cost: float
    timestamp: datetime
    
    class Config:
        from_attributes = True


class DailyUsage(BaseModel):
    """Daily usage summary."""
    date: date
    tokens_used: int
    cost_incurred: float
    requests_made: int
    tokens_remaining: int
    cost_remaining: float
    requests_remaining: int
    within_limits: bool


class UsageAnalytics(BaseModel):
    """Usage analytics response."""
    period_days: int
    total_tokens: int
    total_cost: float
    total_requests: int
    avg_tokens_per_request: float
    avg_cost_per_request: float
    usage_by_model: List[Dict[str, Any]]
    usage_by_endpoint: List[Dict[str, Any]]
    daily_breakdown: List[DailyUsage]


class QuotaUpdate(BaseModel):
    """Quota update request."""
    daily_token_limit: Optional[int] = None
    daily_cost_limit: Optional[float] = None
    daily_request_limit: Optional[int] = None

### src/usage/tracker.py
"""Usage tracking service."""

import logging
from datetime import datetime, timedelta, date
from typing import Dict, Any, Optional

from sqlalchemy import select, func, and_
from sqlalchemy.ext.asyncio import AsyncSession

from src.core.database import get_db
from src.usage.models import UsageLog, UserQuota
from src.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class UsageTracker:
    """Usage tracking service."""
    
    async def track_usage(
        self,
        user_id: int,
        model: str,
        input_tokens: int,
        output_tokens: int,
        cost: float,
        endpoint: str = "chat",
        session_id: Optional[str] = None,
    ) -> bool:
        """Track API usage and update user quotas."""
        async with get_db() as db:
            try:
                # Log detailed usage
                usage_log = UsageLog(
                    user_id=user_id,
                    endpoint=endpoint,
                    model=model,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    total_tokens=input_tokens + output_tokens,
                    cost=cost,
                    session_id=session_id,
                    timestamp=datetime.utcnow(),
                )
                
                db.add(usage_log)
                
                # Update user's daily usage
                await self._update_daily_usage(
                    db=db,
                    user_id=user_id,
                    tokens=input_tokens + output_tokens,
                    cost=cost,
                    requests=1,
                )
                
                await db.commit()
                logger.info(f"Tracked usage for user {user_id}: {input_tokens + output_tokens} tokens, ${cost:.4f}")
                return True
                
            except Exception as e:
                await db.rollback()
                logger.error(f"Usage tracking error: {str(e)}")
                return False
    
    async def _update_daily_usage(
        self,
        db: AsyncSession,
        user_id: int,
        tokens: int,
        cost: float,
        requests: int,
    ):
        """Update user's daily usage totals."""
        today = datetime.utcnow().date()
        
        # Get or create daily quota record
        result = await db.execute(
            select(UserQuota).where(
                and_(
                    UserQuota.user_id == user_id,
                    UserQuota.date == today,
                )
            )
        )
        
        quota = result.scalar_one_or_none()
        
        if quota:
            quota.tokens_used += tokens
            quota.cost_incurred += cost
            quota.requests_made += requests
        else:
            quota = UserQuota(
                user_id=user_id,
                date=today,
                tokens_used=tokens,
                cost_incurred=cost,
                requests_made=requests,
                daily_token_limit=10000,  # Default limit
                daily_cost_limit=1.00,   # Default $1/day limit
                daily_request_limit=100, # Default request limit
            )
            db.add(quota)
    
    async def check_user_limits(self, user_id: int) -> Dict[str, Any]:
        """Check if user has exceeded daily limits."""
        async with get_db() as db:
            today = datetime.utcnow().date()
            
            result = await db.execute(
                select(UserQuota).where(
                    and_(
                        UserQuota.user_id == user_id,
                        UserQuota.date == today,
                    )
                )
            )
            
            quota = result.scalar_one_or_none()
            
            if not quota:
                return {
                    "within_limits": True,
                    "tokens_used": 0,
                    "cost_incurred": 0.0,
                    "requests_made": 0,
                    "tokens_remaining": 10000,
                    "cost_remaining": 1.00,
                    "requests_remaining": 100,
                }
            
            tokens_exceeded = quota.tokens_used >= quota.daily_token_limit
            cost_exceeded = quota.cost_incurred >= quota.daily_cost_limit
            requests_exceeded = quota.requests_made >= quota.daily_request_limit
            
            return {
                "within_limits": not (tokens_exceeded or cost_exceeded or requests_exceeded),
                "tokens_used": quota.tokens_used,
                "tokens_limit": quota.daily_token_limit,
                "tokens_exceeded": tokens_exceeded,
                "cost_incurred": quota.cost_incurred,
                "cost_limit": quota.daily_cost_limit,
                "cost_exceeded": cost_exceeded,
                "requests_made": quota.requests_made,
                "requests_limit": quota.daily_request_limit,
                "requests_exceeded": requests_exceeded,
                "tokens_remaining": max(0, quota.daily_token_limit - quota.tokens_used),
                "cost_remaining": max(0.0, quota.daily_cost_limit - quota.cost_incurred),
                "requests_remaining": max(0, quota.daily_request_limit - quota.requests_made),
            }
    
    async def get_usage_analytics(
        self,
        user_id: int,
        days_back: int = 30,
    ) -> Dict[str, Any]:
        """Get usage analytics for a user."""
        async with get_db() as db:
            start_date = datetime.utcnow() - timedelta(days=days_back)
            
            # Get total usage over period
            result = await db.execute(
                select(
                    func.sum(UsageLog.total_tokens).label('total_tokens'),
                    func.sum(UsageLog.cost).label('total_cost'),
                    func.count(UsageLog.id).label('request_count'),
                    func.avg(UsageLog.total_tokens).label('avg_tokens_per_request'),
                    func.avg(UsageLog.cost).label('avg_cost_per_request'),
                ).where(
                    and_(
                        UsageLog.user_id == user_id,
                        UsageLog.timestamp >= start_date,
                    )
                )
            )
            
            stats = result.one()
            
            # Get usage by model
            model_usage = await db.execute(
                select(
                    UsageLog.model,
                    func.sum(UsageLog.total_tokens).label('tokens'),
                    func.sum(UsageLog.cost).label('cost'),
                    func.count(UsageLog.id).label('requests'),
                ).where(
                    and_(
                        UsageLog.user_id == user_id,
                        UsageLog.timestamp >= start_date,
                    )
                ).group_by(UsageLog.model)
            )
            
            # Get usage by endpoint
            endpoint_usage = await db.execute(
                select(
                    UsageLog.endpoint,
                    func.sum(UsageLog.total_tokens).label('tokens'),
                    func.sum(UsageLog.cost).label('cost'),
                    func.count(UsageLog.id).label('requests'),
                ).where(
                    and_(
                        UsageLog.user_id == user_id,
                        UsageLog.timestamp >= start_date,
                    )
                ).group_by(UsageLog.endpoint)
            )
            
            # Get daily breakdown
            daily_usage = await db.execute(
                select(UserQuota).where(
                    and_(
                        UserQuota.user_id == user_id,
                        UserQuota.date >= start_date.date(),
                    )
                ).order_by(UserQuota.date.desc())
            )
            
            return {
                "period_days": days_back,
                "total_tokens": int(stats.total_tokens or 0),
                "total_cost": float(stats.total_cost or 0),
                "total_requests": int(stats.request_count or 0),
                "avg_tokens_per_request": float(stats.avg_tokens_per_request or 0),
                "avg_cost_per_request": float(stats.avg_cost_per_request or 0),
                "usage_by_model": [
                    {
                        "model": row.model,
                        "tokens": int(row.tokens),
                        "cost": float(row.cost),
                        "requests": int(row.requests),
                    }
                    for row in model_usage
                ],
                "usage_by_endpoint": [
                    {
                        "endpoint": row.endpoint,
                        "tokens": int(row.tokens),
                        "cost": float(row.cost),
                        "requests": int(row.requests),
                    }
                    for row in endpoint_usage
                ],
                "daily_breakdown": [
                    {
                        "date": quota.date,
                        "tokens_used": quota.tokens_used,
                        "cost_incurred": quota.cost_incurred,
                        "requests_made": quota.requests_made,
                        "tokens_remaining": max(0, quota.daily_token_limit - quota.tokens_used),
                        "cost_remaining": max(0.0, quota.daily_cost_limit - quota.cost_incurred),
                        "requests_remaining": max(0, quota.daily_request_limit - quota.requests_made),
                        "within_limits": (
                            quota.tokens_used < quota.daily_token_limit and
                            quota.cost_incurred < quota.daily_cost_limit and
                            quota.requests_made < quota.daily_request_limit
                        ),
                    }
                    for quota in daily_usage.scalars().all()
                ],
            }
    
    async def update_user_quotas(
        self,
        user_id: int,
        daily_token_limit: Optional[int] = None,
        daily_cost_limit: Optional[float] = None,
        daily_request_limit: Optional[int] = None,
    ) -> bool:
        """Update user's daily quotas."""
        async with get_db() as db:
            try:
                today = datetime.utcnow().date()
                
                # Get or create quota record
                result = await db.execute(
                    select(UserQuota).where(
                        and_(
                            UserQuota.user_id == user_id,
                            UserQuota.date == today,
                        )
                    )
                )
                
                quota = result.scalar_one_or_none()
                
                if not quota:
                    quota = UserQuota(
                        user_id=user_id,
                        date=today,
                        daily_token_limit=daily_token_limit or 10000,
                        daily_cost_limit=daily_cost_limit or 1.00,
                        daily_request_limit=daily_request_limit or 100,
                    )
                    db.add(quota)
                else:
                    if daily_token_limit is not None:
                        quota.daily_token_limit = daily_token_limit
                    if daily_cost_limit is not None:
                        quota.daily_cost_limit = daily_cost_limit
                    if daily_request_limit is not None:
                        quota.daily_request_limit = daily_request_limit
                
                await db.commit()
                return True
                
            except Exception as e:
                await db.rollback()
                logger.error(f"Error updating quotas: {str(e)}")
                return False

### src/usage/service.py
"""Usage service."""

import logging
from typing import Dict, List, Any

from sqlalchemy.ext.asyncio import AsyncSession

from src.usage.tracker import UsageTracker
from src.usage.schemas import UsageAnalytics, DailyUsage, QuotaUpdate

logger = logging.getLogger(__name__)


class UsageService:
    """Usage service."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.tracker = UsageTracker()
    
    async def get_user_analytics(self, user_id: int, days: int = 30) -> UsageAnalytics:
        """Get comprehensive usage analytics."""
        data = await self.tracker.get_usage_analytics(user_id, days)
        
        return UsageAnalytics(
            period_days=data["period_days"],
            total_tokens=data["total_tokens"],
            total_cost=data["total_cost"],
            total_requests=data["total_requests"],
            avg_tokens_per_request=data["avg_tokens_per_request"],
            avg_cost_per_request=data["avg_cost_per_request"],
            usage_by_model=data["usage_by_model"],
            usage_by_endpoint=data["usage_by_endpoint"],
            daily_breakdown=[
                DailyUsage(**daily) for daily in data["daily_breakdown"]
            ],
        )
    
    async def get_current_limits(self, user_id: int) -> Dict[str, Any]:
        """Get current usage limits and status."""
        return await self.tracker.check_user_limits(user_id)
    
    async def update_quotas(self, user_id: int, quota_update: QuotaUpdate) -> bool:
        """Update user quotas."""
        return await self.tracker.update_user_quotas(
            user_id=user_id,
            daily_token_limit=quota_update.daily_token_limit,
            daily_cost_limit=quota_update.daily_cost_limit,
            daily_request_limit=quota_update.daily_request_limit,
        )

### src/usage/router.py
"""Usage tracking routes."""

from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession

from src.auth.models import User
from src.core.database import get_db_session
from src.core.deps import get_default_user
from src.usage.schemas import UsageAnalytics, QuotaUpdate
from src.usage.service import UsageService

router = APIRouter()


@router.get("/analytics", response_model=UsageAnalytics)
async def get_usage_analytics(
    days: int = 30,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get usage analytics for the user."""
    usage_service = UsageService(db)
    return await usage_service.get_user_analytics(user.id, days)


@router.get("/limits")
async def get_current_limits(
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Get current usage limits and status."""
    usage_service = UsageService(db)
    return await usage_service.get_current_limits(user.id)


@router.put("/quotas")
async def update_quotas(
    quota_update: QuotaUpdate,
    user: User = Depends(get_default_user),
    db: AsyncSession = Depends(get_db_session),
):
    """Update user quotas."""
    usage_service = UsageService(db)
    success = await usage_service.update_quotas(user.id, quota_update)
    
    if not success:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to update quotas",
        )
    
    return {"message": "Quotas updated successfully"}

## 8. UTILITY MODULES

### src/utils/__init__.py
"""Utility modules."""

### src/utils/file_utils.py
"""File handling utilities."""

import hashlib
import mimetypes
import os
import tempfile
from pathlib import Path
from typing import Optional, Tuple

import aiofiles


async def save_uploaded_file(content: bytes, filename: str, upload_dir: str) -> Path:
    """Save uploaded file to disk."""
    # Create upload directory if it doesn't exist
    upload_path = Path(upload_dir)
    upload_path.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename to avoid conflicts
    file_hash = hashlib.md5(content).hexdigest()[:8]
    name, ext = os.path.splitext(filename)
    unique_filename = f"{name}_{file_hash}{ext}"
    
    file_path = upload_path / unique_filename
    
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)
    
    return file_path


def get_file_type(filename: str, content_type: Optional[str] = None) -> str:
    """Determine file type from filename and content type."""
    if content_type:
        return content_type
    
    # Guess from filename
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type:
        return mime_type
    
    # Fall back to extension
    ext = Path(filename).suffix.lower()
    ext_mapping = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }
    
    return ext_mapping.get(ext, "application/octet-stream")


def validate_file_size(file_size: int, max_size: int = 100_000_000) -> bool:
    """Validate file size."""
    return file_size <= max_size


def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe storage."""
    import re
    
    # Remove potentially dangerous characters
    sanitized = re.sub(r'[^\w\-_\. ]', '', filename)
    
    # Replace spaces with underscores
    sanitized = re.sub(r'\s+', '_', sanitized)
    
    # Limit length
    if len(sanitized) > 100:
        name, ext = os.path.splitext(sanitized)
        sanitized = name[:90] + ext
    
    return sanitized


def get_file_hash(file_path: Path) -> str:
    """Get MD5 hash of file."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


async def cleanup_temp_files(temp_dir: str, max_age_hours: int = 24):
    """Clean up temporary files older than max_age_hours."""
    import time
    
    temp_path = Path(temp_dir)
    if not temp_path.exists():
        return
    
    current_time = time.time()
    max_age_seconds = max_age_hours * 3600
    
    for file_path in temp_path.glob("*"):
        if file_path.is_file():
            file_age = current_time - file_path.stat().st_mtime
            if file_age > max_age_seconds:
                try:
                    file_path.unlink()
                except Exception:
                    pass  # Ignore errors, file might be in use

### src/utils/validation.py
"""Input validation utilities."""

import re
from typing import List, Optional
from urllib.parse import urlparse

from pydantic import BaseModel, Field, validator


class FileValidation:
    """File validation utilities."""
    
    ALLOWED_EXTENSIONS = {
        ".pdf", ".docx", ".doc", ".pptx", ".ppt", ".txt", ".md"
    }
    
    ALLOWED_MIME_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "application/vnd.ms-powerpoint",
        "text/plain",
        "text/markdown",
    }
    
    @classmethod
    def validate_file_extension(cls, filename: str) -> bool:
        """Validate file extension."""
        from pathlib import Path
        return Path(filename).suffix.lower() in cls.ALLOWED_EXTENSIONS
    
    @classmethod
    def validate_mime_type(cls, mime_type: str) -> bool:
        """Validate MIME type."""
        return mime_type in cls.ALLOWED_MIME_TYPES
    
    @classmethod
    def validate_file_size(cls, size: int, max_size: int = 100_000_000) -> bool:
        """Validate file size (default 100MB)."""
        return 0 < size <= max_size


class YouTubeValidation:
    """YouTube URL validation."""
    
    YOUTUBE_PATTERNS = [
        r'(?:https?://)?(?:www\.)?youtube\.com/watch\?v=([a-zA-Z0-9_-]{11})',
        r'(?:https?://)?(?:www\.)?youtu\.be/([a-zA-Z0-9_-]{11})',
        r'(?:https?://)?(?:www\.)?youtube\.com/embed/([a-zA-Z0-9_-]{11})',
        r'(?:https?://)?(?:www\.)?youtube\.com/v/([a-zA-Z0-9_-]{11})',
    ]
    
    @classmethod
    def validate_youtube_url(cls, url: str) -> bool:
        """Validate YouTube URL."""
        if not url:
            return False
        
        for pattern in cls.YOUTUBE_PATTERNS:
            if re.match(pattern, url.strip()):
                return True
        return False
    
    @classmethod
    def extract_video_id(cls, url: str) -> Optional[str]:
        """Extract YouTube video ID from URL."""
        for pattern in cls.YOUTUBE_PATTERNS:
            match = re.search(pattern, url.strip())
            if match:
                return match.group(1)
        return None


class TextValidation:
    """Text content validation."""
    
    @staticmethod
    def validate_text_length(text: str, min_length: int = 1, max_length: int = 10000) -> bool:
        """Validate text length."""
        return min_length <= len(text.strip()) <= max_length
    
    @staticmethod
    def sanitize_text(text: str) -> str:
        """Sanitize text input."""
        # Remove null bytes and control characters
        sanitized = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normalize whitespace
        sanitized = re.sub(r'\s+', ' ', sanitized)
        
        return sanitized.strip()
    
    @staticmethod
    def validate_email(email: str) -> bool:
        """Basic email validation."""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}

        return bool(re.match(pattern, email))
    
    @staticmethod
    def validate_username(username: str) -> bool:
        """Validate username format."""
        # Alphanumeric, underscores, hyphens, 3-30 characters
        pattern = r'^[a-zA-Z0-9_-]{3,30}

        return bool(re.match(pattern, username))


def validate_pagination(page: int = 1, page_size: int = 20, max_page_size: int = 100) -> tuple[int, int]:
    """Validate and normalize pagination parameters."""
    page = max(1, page)
    page_size = max(1, min(page_size, max_page_size))
    return page, page_size

## 9. MAIN APPLICATION

### src/main.py
"""FastAPI application entry point."""

import logging
from contextlib import asynccontextmanager

from fastapi import FastAPI, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.exceptions import RequestValidationError
from starlette.exceptions import HTTPException as StarletteHTTPException

from src.config import get_settings
from src.core.database import create_tables
from src.core.logging import setup_logging
from src.core.exceptions import StudHelperException

# Import routers
from src.auth.router import router as auth_router
from src.documents.router import router as documents_router
from src.rag.router import router as rag_router
from src.ai.router import router as ai_router
from src.usage.router import router as usage_router

# Setup
settings = get_settings()
setup_logging()
logger = logging.getLogger(__name__)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan handler."""
    # Startup
    logger.info("Starting StudHelper AI Backend")
    await create_tables()
    logger.info("Database tables created/verified")
    
    # Ensure default user exists for single-user mode
    try:
        from src.core.database import get_db
        from src.auth.service import AuthService
        
        async with get_db() as db:
            auth_service = AuthService(db)
            user = await auth_service.get_user_by_email("default@studhelper.local")
            if not user:
                await auth_service.create_user(
                    email="default@studhelper.local",
                    username="default_user",
                    password="default_password_change_me",
                )
                logger.info("Created default user for single-user mode")
    except Exception as e:
        logger.error(f"Failed to create default user: {e}")
    
    logger.info("StudHelper AI Backend startup complete")
    
    yield
    
    # Shutdown
    logger.info("StudHelper AI Backend shutting down")


# Create FastAPI app
app = FastAPI(
    title="StudHelper AI Backend",
    description="AI-powered study assistant with document processing and RAG",
    version="2.0.0",
    docs_url="/docs" if settings.DEBUG else None,
    redoc_url="/redoc" if settings.DEBUG else None,
    lifespan=lifespan,
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)


# Exception handlers
@app.exception_handler(StudHelperException)
async def studhelper_exception_handler(request: Request, exc: StudHelperException):
    """Handle custom StudHelper exceptions."""
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "detail": exc.message,
            "status_code": exc.status_code,
            "details": exc.details,
        },
    )


@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Handle validation errors."""
    return JSONResponse(
        status_code=422,
        content={
            "detail": "Validation error",
            "errors": exc.errors(),
        },
    )


@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request: Request, exc: StarletteHTTPException):
    """Handle HTTP exceptions."""
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail, "status_code": exc.status_code},
    )


@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """Handle unexpected exceptions."""
    logger.error(f"Unexpected error: {str(exc)}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "detail": "Internal server error" if settings.is_production else str(exc),
            "status_code": 500,
        },
    )


# Health check
@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "service": "StudHelper AI Backend",
        "version": "2.0.0",
        "environment": settings.ENVIRONMENT,
    }


@app.get("/")
async def root():
    """Root endpoint."""
    return {
        "message": "StudHelper AI Backend",
        "version": "2.0.0",
        "docs": "/docs" if settings.DEBUG else "Documentation disabled in production",
    }


# Include routers
app.include_router(
    auth_router,
    prefix="/api/v1/auth",
    tags=["Authentication"],
)

app.include_router(
    documents_router,
    prefix="/api/v1/documents",
    tags=["Documents"],
)

app.include_router(
    rag_router,
    prefix="/api/v1/rag",
    tags=["RAG"],
)

app.include_router(
    ai_router,
    prefix="/api/v1/ai",
    tags=["AI"],
)

app.include_router(
    usage_router,
    prefix="/api/v1/usage",
    tags=["Usage"],
)


# Development middleware
if settings.DEBUG:
    @app.middleware("http")
    async def log_requests(request: Request, call_next):
        """Log all requests in debug mode."""
        logger.debug(f"{request.method} {request.url}")
        response = await call_next(request)
        logger.debug(f"Response: {response.status_code}")
        return response


if __name__ == "__main__":
    import uvicorn
    
    uvicorn.run(
        "src.main:app",
        host="0.0.0.0",
        port=8000,
        reload=settings.DEBUG,
        log_level="info",
    )

## 10. TESTING FRAMEWORK

### tests/__init__.py
"""Test package."""

### tests/conftest.py
"""Test configuration and fixtures."""

import pytest
import asyncio
import tempfile
import os
from pathlib import Path
from typing import AsyncGenerator

from httpx import AsyncClient
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy.orm import sessionmaker
import respx

from src.main import app
from src.core.database import Base, get_db_session
from src.auth.service import AuthService
from src.auth.models import User
from src.config import get_settings

# Test settings
settings = get_settings()


@pytest.fixture(scope="session")
def event_loop():
    """Create event loop for async tests."""
    loop = asyncio.new_event_loop()
    yield loop
    loop.close()


@pytest.fixture(scope="session")
async def test_engine():
    """Create test database engine."""
    # Use in-memory SQLite for fast tests
    engine = create_async_engine(
        "sqlite+aiosqlite:///./test.db",
        echo=False,
        pool_pre_ping=True,
    )
    
    # Create tables
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    
    yield engine
    
    # Cleanup
    await engine.dispose()
    if os.path.exists("./test.db"):
        os.unlink("./test.db")


@pytest.fixture
async def db_session(test_engine) -> AsyncGenerator[AsyncSession, None]:
    """Create database session for tests."""
    async_session = sessionmaker(
        test_engine, class_=AsyncSession, expire_on_commit=False
    )
    
    async with async_session() as session:
        yield session


@pytest.fixture
async def client(db_session) -> AsyncGenerator[AsyncClient, None]:
    """Create test client with database override."""
    async def override_get_db():
        yield db_session
    
    app.dependency_overrides[get_db_session] = override_get_db
    
    async with AsyncClient(app=app, base_url="http://test") as client:
        yield client
    
    app.dependency_overrides.clear()


@pytest.fixture
async def test_user(db_session: AsyncSession) -> User:
    """Create test user."""
    auth_service = AuthService(db_session)
    user = await auth_service.create_user(
        email="test@example.com",
        username="testuser",
        password="testpass123",
    )
    return user


@pytest.fixture
async def auth_headers(test_user: User) -> dict:
    """Create authentication headers."""
    from src.auth.jwt_handler import create_access_token
    
    token = create_access_token(data={"sub": str(test_user.id)})
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def mock_openai():
    """Mock OpenAI API responses."""
    with respx.mock:
        # Mock chat completions
        respx.post("https://api.openai.com/v1/chat/completions").mock(
            return_value=httpx.Response(200, json={
                "choices": [{
                    "message": {"content": "Mock AI response"},
                    "finish_reason": "stop"
                }],
                "usage": {
                    "prompt_tokens": 50,
                    "completion_tokens": 20,
                    "total_tokens": 70
                },
                "model": "gpt-4o-mini"
            })
        )
        
        # Mock embeddings
        respx.post("https://api.openai.com/v1/embeddings").mock(
            return_value=httpx.Response(200, json={
                "data": [{"embedding": [0.1] * 1536}],
                "usage": {"total_tokens": 10}
            })
        )
        
        yield


@pytest.fixture
def sample_pdf_content() -> bytes:
    """Create sample PDF content."""
    # This would ideally be a real minimal PDF
    # For testing, we'll use a simple bytes representation
    return b"%PDF-1.4\n1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\nxref\n0 3\n0000000000 65535 f \ntrailer\n<</Size 3/Root 1 0 R>>\nstartxref\n9\n%%EOF"


@pytest.fixture
def temp_upload_dir():
    """Create temporary upload directory."""
    with tempfile.TemporaryDirectory() as temp_dir:
        yield temp_dir

### tests/unit/test_auth.py
"""Authentication tests."""

import pytest
from httpx import AsyncClient

from src.auth.service import AuthService
from src.auth.schemas import UserCreate
from src.core.exceptions import ValidationError, AuthenticationError


class TestAuthService:
    """Test authentication service."""
    
    async def test_create_user(self, db_session):
        """Test user creation."""
        auth_service = AuthService(db_session)
        
        user_data = UserCreate(
            email="new@example.com",
            username="newuser",
            password="password123",
        )
        
        user = await auth_service.create_user(user_data)
        
        assert user.email == "new@example.com"
        assert user.username == "newuser"
        assert user.hashed_password != "password123"  # Should be hashed
        assert user.is_active is True
    
    async def test_create_duplicate_user(self, db_session, test_user):
        """Test creating user with duplicate email."""
        auth_service = AuthService(db_session)
        
        user_data = UserCreate(
            email=test_user.email,
            username="different",
            password="password123",
        )
        
        with pytest.raises(ValidationError, match="already exists"):
            await auth_service.create_user(user_data)
    
    async def test_authenticate_user(self, db_session, test_user):
        """Test user authentication."""
        auth_service = AuthService(db_session)
        
        # Valid credentials
        user = await auth_service.authenticate_user(test_user.email, "testpass123")
        assert user is not None
        assert user.id == test_user.id
        
        # Invalid password
        user = await auth_service.authenticate_user(test_user.email, "wrongpass")
        assert user is None
        
        # Invalid email
        user = await auth_service.authenticate_user("wrong@example.com", "testpass123")
        assert user is None


class TestAuthRoutes:
    """Test authentication routes."""
    
    async def test_register_user(self, client: AsyncClient):
        """Test user registration endpoint."""
        response = await client.post("/api/v1/auth/register", json={
            "email": "register@example.com",
            "username": "registeruser",
            "password": "password123",
        })
        
        assert response.status_code == 201
        data = response.json()
        assert data["email"] == "register@example.com"
        assert data["username"] == "registeruser"
        assert "id" in data
    
    async def test_login_user(self, client: AsyncClient, test_user):
        """Test user login endpoint."""
        response = await client.post("/api/v1/auth/login", json={
            "email": test_user.email,
            "password": "testpass123",
        })
        
        assert response.status_code == 200
        data = response.json()
        assert "access_token" in data
        assert data["token_type"] == "bearer"
        assert data["user"]["email"] == test_user.email
    
    async def test_login_invalid_credentials(self, client: AsyncClient):
        """Test login with invalid credentials."""
        response = await client.post("/api/v1/auth/login", json={
            "email": "wrong@example.com",
            "password": "wrongpass",
        })
        
        assert response.status_code == 401
    
    async def test_get_current_user(self, client: AsyncClient, auth_headers):
        """Test getting current user info."""
        response = await client.get("/api/v1/auth/me", headers=auth_headers)
        
        assert response.status_code == 200
        data = response.json()
        assert data["email"] == "test@example.com"

### tests/unit/test_document_processing.py
"""Document processing tests."""

import pytest
from pathlib import Path
import tempfile

from src.documents.service import DocumentService
from src.documents.processors.pdf_processor import PDFProcessor


class TestDocumentService:
    """Test document service."""
    
    async def test_upload_document(self, db_session, test_user, sample_pdf_content):
        """Test document upload."""
        service = DocumentService(db_session)
        
        document = await service.upload_document(
            user_id=test_user.id,
            file_content=sample_pdf_content,
            filename="test.pdf",
            content_type="application/pdf",
        )
        
        assert document.filename == "test.pdf"
        assert document.file_type == "application/pdf"
        assert document.file_size == len(sample_pdf_content)
        assert document.status == "pending"
    
    async def test_add_youtube_video(self, db_session, test_user):
        """Test adding YouTube video."""
        service = DocumentService(db_session)
        
        document = await service.add_youtube_video(
            user_id=test_user.id,
            url="https://www.youtube.com/watch?v=dQw4w9WgXcQ",
            title="Test Video",
        )
        
        assert document.filename == "Test Video"
        assert document.original_filename == "https://www.youtube.com/watch?v=dQw4w9WgXcQ"
        assert document.file_type == "youtube"


class TestPDFProcessor:
    """Test PDF processor."""
    
    def test_can_process(self):
        """Test processor type checking."""
        processor = PDFProcessor()
        
        assert processor.can_process("pdf")
        assert processor.can_process("application/pdf")
        assert not processor.can_process("docx")

### tests/unit/test_ai_service.py
"""AI service tests."""

import pytest
from unittest.mock import AsyncMock, MagicMock

from src.ai.service import ChatService
from src.ai.schemas import ChatMessageRequest, ChatMode
from src.ai.cost_calculator import CostCalculator


class TestCostCalculator:
    """Test cost calculation."""
    
    def test_calculate_cost(self):
        """Test cost calculation."""
        calculator = CostCalculator()
        
        result = calculator.calculate_cost(
            mode="economic",
            input_tokens=100,
            output_tokens=50,
        )
        
        assert result["mode"] == "economic"
        assert result["input_tokens"] == 100
        assert result["output_tokens"] == 50
        assert result["total_tokens"] == 150
        assert result["total_cost"] > 0
    
    def test_count_tokens(self):
        """Test token counting."""
        calculator = CostCalculator()
        
        messages = [
            {"role": "user", "content": "Hello, how are you?"}
        ]
        
        tokens = calculator.count_tokens(messages, "gpt-4o-mini")
        assert tokens > 0


class TestChatService:
    """Test chat service."""
    
    async def test_chat_without_context(self, db_session, test_user, mock_openai):
        """Test basic chat without context."""
        service = ChatService(db_session)
        
        request = ChatMessageRequest(
            message="Hello, how are you?",
            mode=ChatMode.ECONOMIC,
            include_context=False,
        )
        
        response = await service.chat(request, test_user.id)
        
        assert response.content == "Mock AI response"
        assert response.mode == "economic"
        assert response.model == "gpt-4o-mini"
        assert response.input_tokens > 0
        assert response.output_tokens > 0

### tests/unit/test_rag.py
"""RAG system tests."""

import pytest
from unittest.mock import AsyncMock

from src.rag.chunking import TextChunker, AdaptiveChunker
from src.rag.service import RAGService


class TestTextChunker:
    """Test text chunking."""
    
    def test_basic_chunking(self):
        """Test basic text chunking."""
        chunker = TextChunker(chunk_size=100, overlap=20)
        
        text = "This is a test. " * 20  # Long text
        chunks = chunker.chunk_text(text)
        
        assert len(chunks) > 1
        assert all(len(chunk["content"]) <= 120 for chunk in chunks)  # Allow some flexibility
        assert all("chunk_index" in chunk["metadata"] for chunk in chunks)
    
    def test_empty_text(self):
        """Test chunking empty text."""
        chunker = TextChunker()
        
        chunks = chunker.chunk_text("")
        assert chunks == []
        
        chunks = chunker.chunk_text("   ")
        assert chunks == []


class TestAdaptiveChunker:
    """Test adaptive chunking."""
    
    def test_content_type_detection(self):
        """Test content type detection."""
        chunker = AdaptiveChunker()
        
        # Code content
        code_text = "def function():\n    return True\nclass MyClass:\n    pass"
        code_type = chunker._determine_content_type(code_text, {})
        assert code_type == "code"
        
        # Academic content
        academic_text = "Abstract: This paper presents a methodology for testing..."
        academic_type = chunker._determine_content_type(academic_text, {})
        assert academic_type == "academic"

### tests/integration/test_api_endpoints.py
"""Integration tests for API endpoints."""

import pytest
from httpx import AsyncClient


class TestDocumentEndpoints:
    """Test document API endpoints."""
    
    async def test_upload_and_process_document(
        self, 
        client: AsyncClient, 
        auth_headers, 
        sample_pdf_content,
        mock_openai
    ):
        """Test complete document upload and processing flow."""
        # Upload document
        files = {"file": ("test.pdf", sample_pdf_content, "application/pdf")}
        response = await client.post(
            "/api/v1/documents/upload",
            files=files,
            headers=auth_headers,
        )
        
        assert response.status_code == 201
        document = response.json()
        document_id = document["id"]
        
        # Check document was created
        response = await client.get(
            f"/api/v1/documents/{document_id}",
            headers=auth_headers,
        )
        
        assert response.status_code == 200
        assert response.json()["filename"] == "test.pdf"
    
    async def test_youtube_upload(self, client: AsyncClient, auth_headers):
        """Test YouTube video upload."""
        response = await client.post(
            "/api/v1/documents/youtube",
            json={
                "url": "https://www.youtube.com/watch?v=dQw4w9WgXcQ",
                "title": "Test Video"
            },
            headers=auth_headers,
        )
        
        assert response.status_code == 201
        document = response.json()
        assert document["file_type"] == "youtube"


class TestChatEndpoints:
    """Test chat API endpoints."""
    
    async def test_chat_message(self, client: AsyncClient, auth_headers, mock_openai):
        """Test sending chat message."""
        response = await client.post(
            "/api/v1/ai/chat",
            json={
                "message": "Hello, how are you?",
                "mode": "economic",
                "include_context": False,
            },
            headers=auth_headers,
        )
        
        assert response.status_code == 200
        data = response.json()
        assert data["content"] == "Mock AI response"
        assert data["mode"] == "economic"
        assert "session_id" in data
    
    async def test_get_chat_sessions(self, client: AsyncClient, auth_headers):
        """Test getting chat sessions."""
        response = await client.get("/api/v1/ai/sessions", headers=auth_headers)
        
        assert response.status_code == 200
        assert isinstance(response.json(), list)

### tests/integration/test_workflows.py
"""End-to-end workflow tests."""

import pytest
from httpx import AsyncClient


class TestCompleteWorkflow:
    """Test complete application workflows."""
    
    async def test_study_session_workflow(
        self,
        client: AsyncClient,
        auth_headers,
        sample_pdf_content,
        mock_openai,
    ):
        """Test complete study session workflow."""
        # 1. Upload document
        files = {"file": ("study_material.pdf", sample_pdf_content, "application/pdf")}
        upload_response = await client.post(
            "/api/v1/documents/upload",
            files=files,
            headers=auth_headers,
        )
        assert upload_response.status_code == 201
        
        # 2. Wait for processing (in real scenario)
        # For tests, we'll assume processing completes
        
        # 3. Search documents
        search_response = await client.post(
            "/api/v1/rag/search",
            json={
                "query": "test content",
                "limit": 5,
            },
            headers=auth_headers,
        )
        assert search_response.status_code == 200
        
        # 4. Chat about the document
        chat_response = await client.post(
            "/api/v1/ai/chat",
            json={
                "message": "What can you tell me about this document?",
                "mode": "standard",
                "include_context": True,
            },
            headers=auth_headers,
        )
        assert chat_response.status_code == 200
        
        # 5. Check usage statistics
        usage_response = await client.get(
            "/api/v1/usage/analytics?days=1",
            headers=auth_headers,
        )
        assert usage_response.status_code == 200
        usage_data = usage_response.json()
        assert usage_data["total_requests"] > 0

## 11. DOCKER AND DEPLOYMENT

### docker/Dockerfile
# Multi-stage build for production
FROM python:3.11-slim as builder

WORKDIR /app

# Install system dependencies
RUN apt-get update && apt-get install -y \
    build-essential \
    curl \
    ffmpeg \
    && rm -rf /var/lib/apt/lists/*

# Copy requirements and install Python dependencies
COPY requirements/prod.txt .
RUN pip install --no-cache-dir --user -r prod.txt

# Production stage
FROM python:3.11-slim

WORKDIR /app

# Install runtime dependencies
RUN apt-get update && apt-get install -y \
    ffmpeg \
    curl \
    && rm -rf /var/lib/apt/lists/*

# Create non-root user
RUN groupadd -r studhelper && useradd -r -g studhelper studhelper

# Copy installed packages from builder
COPY --from=builder /root/.local /home/studhelper/.local

# Copy application code
COPY src/ ./src/
COPY migrations/ ./migrations/
COPY alembic.ini ./

# Create necessary directories
RUN mkdir -p /app/uploads /app/chroma_data /app/logs && \
    chown -R studhelper:studhelper /app

USER studhelper

# Add local bin to PATH
ENV PATH=/home/studhelper/.local/bin:$PATH

# Health check
HEALTHCHECK --interval=30s --timeout=10s --start-period=30s --retries=3 \
    CMD curl -f http://localhost:8000/health || exit 1

EXPOSE 8000

CMD ["uvicorn", "src.main:app", "--host", "0.0.0.0", "--port", "8000", "--workers", "4"]

### docker/docker-compose.yml
version: '3.8'

services:
  api:
    build:
      context: ..
      dockerfile: docker/Dockerfile
    ports:
      - "8000:8000"
    environment:
      - ENVIRONMENT=development
      - DEBUG=true
      - DATABASE_URL=postgresql://postgres:postgres@db:5432/studhelper
      - CHROMA_HOST=chromadb
      - CHROMA_PORT=8000
    env_file:
      - ../.env.dev
    depends_on:
      - db
      - chromadb
      - redis
    volumes:
      - ../uploads:/app/uploads
      - ../chroma_data:/app/chroma_data
    restart: unless-stopped

  db:
    image: postgres:15-alpine
    environment:
      POSTGRES_DB: studhelper
      POSTGRES_USER: postgres
      POSTGRES_PASSWORD: postgres
    ports:
      - "5432:5432"
    volumes:
      - postgres_data:/var/lib/postgresql/data
      - ./init.sql:/docker-entrypoint-initdb.d/init.sql
    restart: unless-stopped

  chromadb:
    image: chromadb/chroma:latest
    ports:
      - "8001:8000"
    volumes:
      - chroma_data:/chroma/chroma
    environment:
      - CHROMA_SERVER_HOST=0.0.0.0
    restart: unless-stopped

  redis:
    image: redis:7-alpine
    ports:
      - "6379:6379"
    volumes:
      - redis_data:/data
    restart: unless-stopped

volumes:
  postgres_data:
  chroma_data:
  redis_data:

### docker/docker-compose.prod.yml
version: '3.8'

services:
  api:
    build:
      context: ..
      dockerfile: docker/Dockerfile
    ports:
      - "8000:8000"
    environment:
      - ENVIRONMENT=production
      - DEBUG=false
      - DATABASE_URL=postgresql://postgres:${DB_PASSWORD}@db:5432/studhelper
      - CHROMA_HOST=chromadb
      - CHROMA_PORT=8000
    env_file:
      - ../.env.prod
    depends_on:
      - db
      - chromadb
      - redis
    volumes:
      - uploads_data:/app/uploads
      - chroma_data:/app/chroma_data
      - logs_data:/app/logs
    restart: unless-stopped
    deploy:
      replicas: 2
      resources:
        limits:
          memory: 2G
        reservations:
          memory: 1G

  nginx:
    image: nginx:alpine
    ports:
      - "80:80"
      - "443:443"
    volumes:
      - ./nginx.conf:/etc/nginx/nginx.conf
      - ./ssl:/etc/nginx/ssl
    depends_on:
      - api
    restart: unless-stopped

  db:
    image: postgres:15-alpine
    environment:
      POSTGRES_DB: studhelper
      POSTGRES_USER: postgres
      POSTGRES_PASSWORD: ${DB_PASSWORD}
    volumes:
      - postgres_data:/var/lib/postgresql/data
    restart: unless-stopped
    deploy:
      resources:
        limits:
          memory: 1G
        reservations:
          memory: 512M

  chromadb:
    image: chromadb/chroma:latest
    volumes:
      - chroma_data:/chroma/chroma
    environment:
      - CHROMA_SERVER_HOST=0.0.0.0
    restart: unless-stopped

  redis:
    image: redis:7-alpine
    volumes:
      - redis_data:/data
    restart: unless-stopped

volumes:
  postgres_data:
  chroma_data:
  redis_data:
  uploads_data:
  logs_data:

## 12. DATABASE MIGRATIONS

### alembic.ini
# A generic, single database configuration.

[alembic]
# path to migration scripts
script_location = migrations

# template used to generate migration file names
file_template = %%(year)d_%%(month).2d_%%(day).2d_%%(hour).2d%%(minute).2d_%%(rev)s_%%(slug)s

# sys.path path, will be prepended to sys.path if present.
prepend_sys_path = .

# timezone to use when rendering the date within the migration file
# as well as the filename.
timezone =

# max length of characters to apply to the
# "slug" field
truncate_slug_length = 40

# set to 'true' to run the environment during
# the 'revision' command, regardless of autogenerate
revision_environment = false

# set to 'true' to allow .pyc and .pyo files without
# a source .py file to be detected as revisions in the
# versions/ directory
sourceless = false

# version path separator; As mentioned above, this is the character used to split
# version_locations. The default within new alembic.ini files is "os", which uses
# os.pathsep. If this key is omitted entirely, it falls back to the legacy
# behavior of splitting on spaces and/or commas.
version_path_separator = :

# set to 'true' to search source files recursively
# in each "version_locations" directory
recursive_version_locations = false

# the output encoding used when revision files
# are written from script.py.mako
output_encoding = utf-8

sqlalchemy.url = postgresql://localhost/studhelper

[post_write_hooks]
# post_write_hooks defines scripts or Python functions that are run
# on newly generated revision scripts.  See the documentation for further
# detail and examples

# format using "black" - use the console_scripts runner, against the "black" entrypoint
hooks = black
black.type = console_scripts
black.entrypoint = black
black.options = -l 79 REVISION_SCRIPT_FILENAME

# Logging configuration
[loggers]
keys = root,sqlalchemy,alembic

[handlers]
keys = console

[formatters]
keys = generic

[logger_root]
level = WARN
handlers = console
qualname =

[logger_sqlalchemy]
level = WARN
handlers =
qualname = sqlalchemy.engine

[logger_alembic]
level = INFO
handlers =
qualname = alembic

[handler_console]
class = StreamHandler
args = (sys.stderr,)
level = NOTSET
formatter = generic

[formatter_generic]
format = %(levelname)-5.5s [%(name)s] %(message)s
datefmt = %H:%M:%S

### migrations/env.py
import asyncio
from logging.config import fileConfig
import sys
from pathlib import Path

from sqlalchemy import pool
from sqlalchemy.engine import Connection
from sqlalchemy.ext.asyncio import async_engine_from_config

from alembic import context

# Add src to path
sys.path.append(str(Path(__file__).parent.parent))

from src.core.database import Base
from src.config import get_settings

# this is the Alembic Config object
config = context.config

# Interpret the config file for Python logging
if config.config_file_name is not None:
    fileConfig(config.config_file_name)

# Set SQLAlchemy URL from settings
settings = get_settings()
config.set_main_option("sqlalchemy.url", settings.DATABASE_URL)

# add your model's MetaData object here for 'autogenerate' support
target_metadata = Base.metadata


def run_migrations_offline() -> None:
    """Run migrations in 'offline' mode."""
    url = config.get_main_option("sqlalchemy.url")
    context.configure(
        url=url,
        target_metadata=target_metadata,
        literal_binds=True,
        dialect_opts={"paramstyle": "named"},
    )

    with context.begin_transaction():
        context.run_migrations()


def do_run_migrations(connection: Connection) -> None:
    context.configure(connection=connection, target_metadata=target_metadata)

    with context.begin_transaction():
        context.run_migrations()


async def run_async_migrations() -> None:
    """In this scenario we need to create an Engine and associate a connection with the context."""
    connectable = async_engine_from_config(
        config.get_section(config.config_ini_section, {}),
        prefix="sqlalchemy.",
        poolclass=pool.NullPool,
    )

    async with connectable.connect() as connection:
        await connection.run_sync(do_run_migrations)

    await connectable.dispose()


def run_migrations_online() -> None:
    """Run migrations in 'online' mode."""
    asyncio.run(run_async_migrations())


if context.is_offline_mode():
    run_migrations_offline()
else:
    run_migrations_online()

### migrations/script.py.mako
"""${message}

Revision ID: ${up_revision}
Revises: ${down_revision | comma,n}
Create Date: ${create_date}

"""
from typing import Sequence, Union

from alembic import op
import sqlalchemy as sa
${imports if imports else ""}

# revision identifiers, used by Alembic.
revision: str = ${repr(up_revision)}
down_revision: Union[str, None] = ${repr(down_revision)}
branch_labels: Union[str, Sequence[str], None] = ${repr(branch_labels)}
depends_on: Union[str, Sequence[str], None] = ${repr(depends_on)}


def upgrade() -> None:
    ${upgrades if upgrades else "pass"}


def downgrade() -> None:
    ${downgrades if downgrades else "pass"}

## 13. SETUP SCRIPTS

### scripts/setup_dev.py
#!/usr/bin/env python3
"""Development environment setup script."""

import asyncio
import os
import sys
from pathlib import Path

# Add src to path
sys.path.append(str(Path(__file__).parent.parent))

from src.core.database import create_tables
from src.auth.service import AuthService
from src.core.database import get_db
from src.config import get_settings


async def setup_database():
    """Setup database tables."""
    print("Creating database tables...")
    await create_tables()
    print("✅ Database tables created")


async def create_default_user():
    """Create default user for development."""
    print("Creating default user...")
    
    async with get_db() as db:
        auth_service = AuthService(db)
        
        # Check if default user exists
        user = await auth_service.get_user_by_email("default@studhelper.local")
        if user:
            print("✅ Default user already exists")
            return
        
        # Create default user
        await auth_service.create_user(
            email="default@studhelper.local",
            username="default_user",
            password="default_password_change_me",
        )
        print("✅ Default user created")


async def setup_directories():
    """Create necessary directories."""
    settings = get_settings()
    
    directories = [
        settings.UPLOAD_DIR,
        settings.CHROMA_PERSIST_DIR,
        "logs",
    ]
    
    for directory in directories:
        Path(directory).mkdir(parents=True, exist_ok=True)
        print(f"✅ Created directory: {directory}")


async def main():
    """Main setup function."""
    print("🚀 Setting up StudHelper development environment...")
    
    try:
        await setup_directories()
        await setup_database()
        await create_default_user()
        
        print("\n🎉 Development environment setup complete!")
        print("\nNext steps:")
        print("1. Copy .env.example to .env and configure your settings")
        print("2. Start the development server: uvicorn src.main:app --reload")
        print("3. Visit http://localhost:8000/docs to see the API documentation")
        
    except Exception as e:
        print(f"❌ Setup failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())

### scripts/migrate_db.py
#!/usr/bin/env python3
"""Database migration script."""

import subprocess
import sys
from pathlib import Path


def run_command(command: str) -> bool:
    """Run a shell command."""
    print(f"Running: {command}")
    result = subprocess.run(command, shell=True)
    return result.returncode == 0


def main():
    """Main migration function."""
    if len(sys.argv) < 2:
        print("Usage: python migrate_db.py <command>")
        print("Commands:")
        print("  init     - Initialize migrations")
        print("  migrate  - Create new migration")
        print("  upgrade  - Apply migrations")
        print("  downgrade - Downgrade one migration")
        sys.exit(1)
    
    command = sys.argv[1]
    
    if command == "init":
        success = run_command("alembic init migrations")
        if success:
            print("✅ Migrations initialized")
    
    elif command == "migrate":
        message = sys.argv[2] if len(sys.argv) > 2 else "auto migration"
        success = run_command(f'alembic revision --autogenerate -m "{message}"')
        if success:
            print("✅ Migration created")
    
    elif command == "upgrade":
        success = run_command("alembic upgrade head")
        if success:
            print("✅ Migrations applied")
    
    elif command == "downgrade":
        success = run_command("alembic downgrade -1")
        if success:
            print("✅ Migration downgraded")
    
    else:
        print(f"Unknown command: {command}")
        sys.exit(1)


if __name__ == "__main__":
    main()

## 14. README AND DOCUMENTATION

### README.md
# StudHelper Backend

AI-powered study assistant backend with document processing, RAG, and OpenAI integration.

## Features

- **Document Processing**: PDF, Word, PowerPoint, YouTube transcription
- **RAG System**: ChromaDB vector storage with intelligent retrieval
- **AI Integration**: OpenAI GPT-4o with three modes (Economic/Standard/Turbo)
- **Cost Tracking**: Real-time token usage and cost monitoring
- **Authentication**: JWT-based user authentication
- **Scalable**: Async FastAPI with proper database management

## Quick Start

### 1. Environment Setup

```bash
# Clone repository
git clone <repository-url>
cd studhelper-backend

# Create virtual environment
python -m venv venv
source venv/bin/activate  # On Windows: venv\Scripts\activate

# Install dependencies
pip install -r requirements/dev.txt
```

### 2. Configuration

```bash
# Copy environment template
cp .env.example .env

# Edit .env with your settings
nano .env
```

Required environment variables:
- `OPENAI_API_KEY`: Your OpenAI API key
- `DATABASE_URL`: PostgreSQL connection string
- `SECRET_KEY`: JWT secret key (32+ characters)

### 3. Database Setup

```bash
# Setup development environment
python scripts/setup_dev.py

# Run migrations
python scripts/migrate_db.py upgrade
```

### 4. Start Development Server

```bash
uvicorn src.main:app --reload --host 0.0.0.0 --port 8000
```

Visit http://localhost:8000/docs for API documentation.

## Architecture

### Core Components

- **FastAPI**: Async web framework
- **SQLAlchemy**: Database ORM with async support
- **ChromaDB**: Vector database for RAG
- **OpenAI**: GPT-4o integration with cost tracking
- **Alembic**: Database migrations

### Project Structure

```
src/
├── main.py              # FastAPI application
├── config.py            # Configuration management
├── auth/                # Authentication system
├── documents/           # Document processing
├── rag/                # RAG system
├── ai/                 # OpenAI integration
├── usage/              # Usage tracking
├── core/               # Core infrastructure
└── utils/              # Utilities
```

## API Endpoints

### Authentication
- `POST /api/v1/auth/register` - Register user
- `POST /api/v1/auth/login` - Login user
- `GET /api/v1/auth/me` - Get current user

### Documents
- `POST /api/v1/documents/upload` - Upload document
- `POST /api/v1/documents/youtube` - Add YouTube video
- `GET /api/v1/documents/` - List documents
- `DELETE /api/v1/documents/{id}` - Delete document

### AI Chat
- `POST /api/v1/ai/chat` - Send chat message
- `POST /api/v1/ai/chat/stream` - Stream chat response
- `GET /api/v1/ai/sessions` - List chat sessions

### RAG Search
- `POST /api/v1/rag/search` - Search documents
- `GET /api/v1/rag/context/{query}` - Get context

### Usage Tracking
- `GET /api/v1/usage/analytics` - Usage analytics
- `GET /api/v1/usage/limits` - Current limits

## Development

### Running Tests

```bash
# Run all tests
pytest

# Run with coverage
pytest --cov=src

# Run specific test file
pytest tests/unit/test_auth.py -v
```

### Code Quality

```bash
# Format code
black src tests
isort src tests

# Lint code
flake8 src tests

# Type checking
mypy src
```

### Database Migrations

```bash
# Create migration
python scripts/migrate_db.py migrate "add new table"

# Apply migrations
python scripts/migrate_db.py upgrade

# Rollback migration
python scripts/migrate_db.py downgrade
```

## Deployment

### Docker Development

```bash
# Start development environment
docker-compose -f docker/docker-compose.yml up -d

# View logs
docker-compose logs -f api
```

### Docker Production

```bash
# Build and start production environment
docker-compose -f docker/docker-compose.prod.yml up -d

# Scale API service
docker-compose -f docker/docker-compose.prod.yml up -d --scale api=3
```

### Environment Variables

Production settings:
- `ENVIRONMENT=production`
- `DEBUG=false`
- `DATABASE_URL`: Production database
- `CORS_ORIGINS`: Allowed origins
- `OPENAI_API_KEY`: Production API key

## Cost Management

### Default Limits (Per Day)
- **Tokens**: 10,000
- **Cost**: $1.00
- **Requests**: 100

### Pricing per 1M tokens
- **Economic Mode**: $0.15 input, $0.60 output (gpt-4o-mini)
- **Standard Mode**: $2.50 input, $10.00 output (gpt-4o)
- **Turbo Mode**: $2.50 input, $10.00 output (gpt-4o + reasoning)

## Troubleshooting

### Common Issues

**Database connection errors:**
```bash
# Check PostgreSQL is running
docker-compose ps db

# Reset database
docker-compose down -v
docker-compose up -d db
python scripts/migrate_db.py upgrade
```

**ChromaDB connection errors:**
```bash
# Check ChromaDB service
docker-compose ps chromadb

# Reset ChromaDB
docker-compose down chromadb
docker volume rm studhelper_chroma_data
docker-compose up -d chromadb
```

**OpenAI API errors:**
- Verify API key in .env
- Check account billing status
- Monitor rate limits

### Performance Optimization

**Database:**
- Add indexes for frequently queried columns
- Use connection pooling
- Optimize query patterns

**Vector Store:**
- Batch document uploads
- Use appropriate chunk sizes
- Monitor embedding costs

**API:**
- Enable response caching
- Use connection keep-alive
- Monitor memory usage

## Contributing

1. Fork the repository
2. Create feature branch
3. Add tests for new functionality
4. Ensure all tests pass
5. Submit pull request

## License

MIT License - see LICENSE file for details.
