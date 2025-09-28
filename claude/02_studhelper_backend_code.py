### NEW FILE: app/main.py
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
import logging
from contextlib import asynccontextmanager

from app.database import create_tables
from app.routes import auth, classes, permissions, chat, documents, usage
from app.config import get_settings

settings = get_settings()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("logs/app.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("Starting StudHelper Backend...")
    create_tables()
    logger.info("Database tables created successfully")
    yield
    # Shutdown
    logger.info("Shutting down StudHelper Backend...")

app = FastAPI(
    title="StudHelper Backend",
    description="AI-powered class-based learning platform with flexible permissions",
    version="1.0.0",
    lifespan=lifespan
)

# Security middleware
app.add_middleware(TrustedHostMiddleware, allowed_hosts=settings.ALLOWED_HOSTS)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include routers
app.include_router(auth.router, prefix="/api/v1/auth", tags=["Authentication"])
app.include_router(classes.router, prefix="/api/v1/classes", tags=["Classes"])
app.include_router(permissions.router, prefix="/api/v1/classes", tags=["Permissions"])
app.include_router(chat.router, prefix="/api/v1/chat", tags=["Chat"])
app.include_router(documents.router, prefix="/api/v1/documents", tags=["Documents"])
app.include_router(usage.router, prefix="/api/v1/usage", tags=["Usage"])

@app.get("/")
async def root():
    return {
        "message": "StudHelper Backend API",
        "version": "1.0.0",
        "docs": "/docs"
    }

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

### NEW FILE: app/database.py
from sqlalchemy import create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from app.config import get_settings
import logging

settings = get_settings()
logger = logging.getLogger(__name__)

engine = create_engine(
    settings.DATABASE_URL,
    pool_pre_ping=True,
    pool_recycle=300,
    echo=settings.DEBUG
)

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def create_tables():
    try:
        Base.metadata.create_all(bind=engine)
        logger.info("Database tables created successfully")
    except Exception as e:
        logger.error(f"Error creating database tables: {e}")
        raise

### NEW FILE: app/config.py
from pydantic_settings import BaseSettings
from typing import List
import os

class Settings(BaseSettings):
    # Database
    DATABASE_URL: str = "postgresql://postgres:password@localhost/studhelper"
    
    # Security
    SECRET_KEY: str = "your-super-secret-key-change-in-production"
    ALGORITHM: str = "HS256"
    ACCESS_TOKEN_EXPIRE_MINUTES: int = 30
    
    # OpenAI
    OPENAI_API_KEY: str = ""
    OPENAI_MODEL: str = "gpt-4o-mini"
    EMBEDDING_MODEL: str = "text-embedding-3-small"
    
    # File upload
    UPLOAD_DIR: str = "uploads"
    MAX_FILE_SIZE: int = 10 * 1024 * 1024  # 10MB
    ALLOWED_FILE_TYPES: List[str] = ["pdf", "txt", "docx"]
    
    # CORS and security
    CORS_ORIGINS: List[str] = ["http://localhost:3000", "http://localhost:8080"]
    ALLOWED_HOSTS: List[str] = ["localhost", "127.0.0.1"]
    
    # Application
    DEBUG: bool = False
    
    # OpenAI Pricing (per million tokens)
    OPENAI_PRICING: dict = {
        "gpt-4o-mini": {"input": 0.15, "output": 0.60},
        "text-embedding-3-small": {"input": 0.02, "output": 0.0}
    }

    class Config:
        env_file = ".env"
        case_sensitive = True

_settings = None

def get_settings() -> Settings:
    global _settings
    if _settings is None:
        _settings = Settings()
    return _settings

### NEW FILE: app/models/__init__.py
from sqlalchemy import Column, Integer, String, DateTime, ForeignKey, Text, Float, Boolean, Date, UniqueConstraint
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship
from datetime import datetime, date
import enum
from sqlalchemy import Enum

Base = declarative_base()

class ProcessingStatus(enum.Enum):
    PENDING = "pending"
    PROCESSING = "processing" 
    COMPLETED = "completed"
    FAILED = "failed"

class DocumentScope(enum.Enum):
    CLASS = "class"    # Available to all chats in the class
    CHAT = "chat"      # Only available to specific chat session

class User(Base):
    __tablename__ = "users"
    
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String, unique=True, index=True, nullable=False)
    username = Column(String, unique=True, index=True, nullable=False)
    hashed_password = Column(String, nullable=False)
    full_name = Column(String)
    created_at = Column(DateTime, default=datetime.utcnow)
    is_active = Column(Boolean, default=True)
    
    # Relationships
    owned_classes = relationship("Class", back_populates="owner")
    class_memberships = relationship("ClassMembership", back_populates="user")
    chat_sessions = relationship("ChatSession", back_populates="user")
    usage_records = relationship("UsageRecord", back_populates="user")

class Class(Base):
    __tablename__ = "classes"
    
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String, nullable=False)
    description = Column(Text)
    class_code = Column(String, unique=True, index=True, nullable=False)
    owner_id = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)
    is_active = Column(Boolean, default=True)
    
    # Relationships
    owner = relationship("User", back_populates="owned_classes")
    memberships = relationship("ClassMembership", back_populates="class_obj")
    documents = relationship("Document", back_populates="class_obj")
    chat_sessions = relationship("ChatSession", back_populates="class_obj")

class ClassMembership(Base):
    """Class membership with granular permissions"""
    __tablename__ = "class_memberships"
    
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    class_id = Column(Integer, ForeignKey("classes.id"))
    joined_at = Column(DateTime, default=datetime.utcnow)
    
    # Permission flags
    is_manager = Column(Boolean, default=False)
    can_read = Column(Boolean, default=True)
    can_chat = Column(Boolean, default=True)
    max_concurrent_chats = Column(Integer, default=3)
    can_share_class = Column(Boolean, default=False)
    can_upload_documents = Column(Boolean, default=True)
    
    # Billing and limits
    is_sponsored = Column(Boolean, default=False)
    daily_token_limit = Column(Integer, default=1_000_000)    # 1M tokens
    weekly_token_limit = Column(Integer, default=5_000_000)   # 5M tokens
    monthly_token_limit = Column(Integer, default=15_000_000) # 15M tokens
    
    # Relationships
    user = relationship("User", back_populates="class_memberships")
    class_obj = relationship("Class", back_populates="memberships")
    
    __table_args__ = (
        UniqueConstraint('user_id', 'class_id', name='_user_class_membership_uc'),
    )

class ClassUsageTracker(Base):
    """Per-class, per-user usage tracking"""
    __tablename__ = "class_usage_trackers"
    
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    class_id = Column(Integer, ForeignKey("classes.id"))
    
    # Current usage counters
    daily_tokens_used = Column(Integer, default=0)
    weekly_tokens_used = Column(Integer, default=0)
    monthly_tokens_used = Column(Integer, default=0)
    
    # Reset tracking (Madrid timezone)
    last_daily_reset = Column(Date, default=date.today)
    last_weekly_reset = Column(Date, default=date.today)
    last_monthly_reset = Column(Date, default=date.today)
    
    __table_args__ = (
        UniqueConstraint('user_id', 'class_id', name='_user_class_usage_uc'),
    )

class Document(Base):
    __tablename__ = "documents"
    
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, nullable=False)
    original_filename = Column(String, nullable=False)
    file_path = Column(String, nullable=False)
    file_type = Column(String, nullable=False)
    file_size = Column(Integer, nullable=False)
    
    # Dual-level document system
    scope = Column(Enum(DocumentScope), nullable=False)
    class_id = Column(Integer, ForeignKey("classes.id"))
    session_id = Column(Integer, ForeignKey("chat_sessions.id"), nullable=True)
    
    uploaded_by = Column(Integer, ForeignKey("users.id"))
    uploaded_at = Column(DateTime, default=datetime.utcnow)
    processing_status = Column(Enum(ProcessingStatus), default=ProcessingStatus.PENDING)
    processing_error = Column(Text, nullable=True)
    
    # Relationships
    class_obj = relationship("Class", back_populates="documents")
    session = relationship("ChatSession", back_populates="documents")
    chunks = relationship("DocumentChunk", back_populates="document")

class DocumentChunk(Base):
    __tablename__ = "document_chunks"
    
    id = Column(Integer, primary_key=True, index=True)
    document_id = Column(Integer, ForeignKey("documents.id"))
    content = Column(Text, nullable=False)
    chunk_index = Column(Integer, nullable=False)
    char_start = Column(Integer, nullable=False)
    char_end = Column(Integer, nullable=False)
    vector_id = Column(String, nullable=True)  # Vector store ID
    created_at = Column(DateTime, default=datetime.utcnow)
    
    # Relationships
    document = relationship("Document", back_populates="chunks")

class ChatSession(Base):
    __tablename__ = "chat_sessions"
    
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String, nullable=False)
    user_id = Column(Integer, ForeignKey("users.id"))
    class_id = Column(Integer, ForeignKey("classes.id"))
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    is_active = Column(Boolean, default=True)
    
    # Relationships
    user = relationship("User", back_populates="chat_sessions")
    class_obj = relationship("Class", back_populates="chat_sessions")
    documents = relationship("Document", back_populates="session")
    messages = relationship("ChatMessage", back_populates="session")

class ChatMessage(Base):
    __tablename__ = "chat_messages"
    
    id = Column(Integer, primary_key=True, index=True)
    session_id = Column(Integer, ForeignKey("chat_sessions.id"))
    content = Column(Text, nullable=False)
    is_user = Column(Boolean, nullable=False)
    timestamp = Column(DateTime, default=datetime.utcnow)
    response_time_ms = Column(Integer, nullable=True)
    context_used = Column(Text, nullable=True)
    tokens_used = Column(Integer, default=0)
    
    # Relationships
    session = relationship("ChatSession", back_populates="messages")

class UsageRecord(Base):
    __tablename__ = "usage_records"
    
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    model_name = Column(String, nullable=False)
    operation_type = Column(String, nullable=False)  # 'chat', 'embedding'
    input_tokens = Column(Integer, default=0)
    output_tokens = Column(Integer, default=0)
    cost = Column(Float, nullable=False)
    timestamp = Column(DateTime, default=datetime.utcnow)
    session_id = Column(Integer, ForeignKey("chat_sessions.id"), nullable=True)
    
    # Billing attribution
    billed_to_user_id = Column(Integer, ForeignKey("users.id"))
    is_sponsored = Column(Boolean, default=False)
    is_overflow = Column(Boolean, default=False)
    
    # Relationships
    user = relationship("User", back_populates="usage_records")

### NEW FILE: app/schemas/__init__.py
from pydantic import BaseModel, EmailStr
from typing import Optional, List, Dict, Any
from datetime import datetime
from app.models import ProcessingStatus, DocumentScope

# User schemas
class UserBase(BaseModel):
    email: EmailStr
    username: str
    full_name: Optional[str] = None

class UserCreate(UserBase):
    password: str

class UserUpdate(BaseModel):
    email: Optional[EmailStr] = None
    full_name: Optional[str] = None

class UserResponse(UserBase):
    id: int
    created_at: datetime
    is_active: bool
    
    class Config:
        from_attributes = True

# Auth schemas
class Token(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse

class LoginRequest(BaseModel):
    username: str
    password: str

# Class schemas
class ClassBase(BaseModel):
    name: str
    description: Optional[str] = None

class ClassCreate(ClassBase):
    pass

class ClassResponse(ClassBase):
    id: int
    class_code: str
    owner_id: int
    created_at: datetime
    is_active: bool
    member_count: Optional[int] = 0
    
    class Config:
        from_attributes = True

class JoinClassRequest(BaseModel):
    class_code: str

# Permission schemas
class PermissionUpdate(BaseModel):
    can_read: Optional[bool] = None
    can_chat: Optional[bool] = None
    max_concurrent_chats: Optional[int] = None
    can_share_class: Optional[bool] = None
    can_upload_documents: Optional[bool] = None
    daily_token_limit: Optional[int] = None
    weekly_token_limit: Optional[int] = None
    monthly_token_limit: Optional[int] = None

class MembershipResponse(BaseModel):
    id: int
    user_id: int
    username: str
    full_name: Optional[str]
    joined_at: datetime
    is_manager: bool
    can_read: bool
    can_chat: bool
    max_concurrent_chats: int
    can_share_class: bool
    can_upload_documents: bool
    is_sponsored: bool
    daily_token_limit: int
    weekly_token_limit: int
    monthly_token_limit: int
    
    class Config:
        from_attributes = True

class SponsorshipUpdate(BaseModel):
    is_sponsored: bool

# Document schemas
class DocumentResponse(BaseModel):
    id: int
    filename: str
    original_filename: str
    file_type: str
    file_size: int
    scope: DocumentScope
    class_id: int
    session_id: Optional[int]
    uploaded_by: int
    uploaded_at: datetime
    processing_status: ProcessingStatus
    processing_error: Optional[str]
    
    class Config:
        from_attributes = True

# Chat schemas
class ChatSessionCreate(BaseModel):
    title: str
    class_id: int

class ChatSessionResponse(BaseModel):
    id: int
    title: str
    user_id: int
    class_id: int
    created_at: datetime
    updated_at: datetime
    is_active: bool
    message_count: Optional[int] = 0
    
    class Config:
        from_attributes = True

class MessageCreate(BaseModel):
    content: str

class MessageResponse(BaseModel):
    id: int
    session_id: int
    content: str
    is_user: bool
    timestamp: datetime
    response_time_ms: Optional[int]
    context_used: Optional[str]
    tokens_used: int
    
    class Config:
        from_attributes = True

class ChatResponse(BaseModel):
    user_message: MessageResponse
    ai_response: MessageResponse
    cost: float
    response_time_ms: int
    context_provided: bool

# Usage schemas
class UsageStats(BaseModel):
    daily_tokens_used: int
    weekly_tokens_used: int
    monthly_tokens_used: int
    daily_limit: int
    weekly_limit: int
    monthly_limit: int
    daily_remaining: int
    weekly_remaining: int
    monthly_remaining: int

class ClassUsageOverview(BaseModel):
    user_id: int
    username: str
    usage_stats: UsageStats
    is_sponsored: bool
    last_activity: Optional[datetime]

class UsageRecord(BaseModel):
    id: int
    model_name: str
    operation_type: str
    input_tokens: int
    output_tokens: int
    cost: float
    timestamp: datetime
    is_sponsored: bool
    is_overflow: bool

### NEW FILE: app/routes/__init__.py
# Routes package initialization

### NEW FILE: app/routes/auth.py
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.orm import Session
from app.database import get_db
from app.schemas import UserCreate, UserResponse, UserUpdate, LoginRequest, Token
from app.services.auth_service import AuthService
from app.utils.security import get_current_user
import logging

router = APIRouter()
security = HTTPBearer()
logger = logging.getLogger(__name__)

@router.post("/register", response_model=UserResponse, status_code=status.HTTP_201_CREATED)
async def register(user_data: UserCreate, db: Session = Depends(get_db)):
    """Register a new user"""
    try:
        auth_service = AuthService()
        user = await auth_service.create_user(db, user_data)
        logger.info(f"New user registered: {user.email}")
        return user
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error registering user: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.post("/login", response_model=Token)
async def login(login_data: LoginRequest, db: Session = Depends(get_db)):
    """Authenticate user and return JWT token"""
    try:
        auth_service = AuthService()
        token_data = await auth_service.authenticate_user(db, login_data.username, login_data.password)
        logger.info(f"User logged in: {login_data.username}")
        return token_data
    except ValueError as e:
        raise HTTPException(status_code=401, detail=str(e))
    except Exception as e:
        logger.error(f"Error during login: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.post("/logout")
async def logout(current_user: UserResponse = Depends(get_current_user)):
    """Logout user (client should discard token)"""
    logger.info(f"User logged out: {current_user.username}")
    return {"message": "Successfully logged out"}

@router.get("/me", response_model=UserResponse)
async def get_current_user_profile(current_user: UserResponse = Depends(get_current_user)):
    """Get current user profile"""
    return current_user

@router.put("/profile", response_model=UserResponse)
async def update_profile(
    user_update: UserUpdate,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update user profile"""
    try:
        auth_service = AuthService()
        updated_user = await auth_service.update_user(db, current_user.id, user_update)
        logger.info(f"User profile updated: {current_user.username}")
        return updated_user
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error updating profile: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.delete("/account")
async def delete_account(
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete user account"""
    try:
        auth_service = AuthService()
        await auth_service.delete_user(db, current_user.id)
        logger.info(f"User account deleted: {current_user.username}")
        return {"message": "Account successfully deleted"}
    except Exception as e:
        logger.error(f"Error deleting account: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

### NEW FILE: app/routes/classes.py
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from typing import List
from app.database import get_db
from app.schemas import ClassCreate, ClassResponse, JoinClassRequest, UserResponse
from app.services.permission_service import PermissionService
from app.utils.security import get_current_user
import logging
import string
import secrets

router = APIRouter()
logger = logging.getLogger(__name__)

def generate_class_code() -> str:
    """Generate a unique 8-character class code"""
    return ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(8))

@router.post("/", response_model=ClassResponse, status_code=status.HTTP_201_CREATED)
async def create_class(
    class_data: ClassCreate,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new class"""
    try:
        from app.models import Class, ClassMembership
        
        # Generate unique class code
        class_code = generate_class_code()
        while db.query(Class).filter(Class.class_code == class_code).first():
            class_code = generate_class_code()
        
        # Create class
        new_class = Class(
            name=class_data.name,
            description=class_data.description,
            class_code=class_code,
            owner_id=current_user.id
        )
        db.add(new_class)
        db.flush()
        
        # Add owner as manager
        owner_membership = ClassMembership(
            user_id=current_user.id,
            class_id=new_class.id,
            is_manager=True,
            can_read=True,
            can_chat=True,
            can_share_class=True,
            can_upload_documents=True,
            max_concurrent_chats=10  # Managers get more chats
        )
        db.add(owner_membership)
        db.commit()
        db.refresh(new_class)
        
        # Add member count
        result = ClassResponse.model_validate(new_class)
        result.member_count = 1
        
        logger.info(f"Class created: {new_class.name} by {current_user.username}")
        return result
        
    except Exception as e:
        logger.error(f"Error creating class: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/", response_model=List[ClassResponse])
async def get_user_classes(
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get classes where user is owner or member"""
    try:
        from app.models import Class, ClassMembership
        
        # Get classes where user is a member
        memberships = db.query(ClassMembership).filter(
            ClassMembership.user_id == current_user.id
        ).all()
        
        classes = []
        for membership in memberships:
            class_obj = membership.class_obj
            if class_obj.is_active:
                # Count members
                member_count = db.query(ClassMembership).filter(
                    ClassMembership.class_id == class_obj.id
                ).count()
                
                result = ClassResponse.model_validate(class_obj)
                result.member_count = member_count
                classes.append(result)
        
        return classes
        
    except Exception as e:
        logger.error(f"Error getting user classes: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.post("/join", response_model=ClassResponse)
async def join_class(
    join_data: JoinClassRequest,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Join a class using class code"""
    try:
        from app.models import Class, ClassMembership
        
        # Find class by code
        class_obj = db.query(Class).filter(
            Class.class_code == join_data.class_code,
            Class.is_active == True
        ).first()
        
        if not class_obj:
            raise HTTPException(status_code=404, detail="Class not found")
        
        # Check if already a member
        existing_membership = db.query(ClassMembership).filter(
            ClassMembership.user_id == current_user.id,
            ClassMembership.class_id == class_obj.id
        ).first()
        
        if existing_membership:
            raise HTTPException(status_code=400, detail="Already a member of this class")
        
        # Create membership with default permissions
        membership = ClassMembership(
            user_id=current_user.id,
            class_id=class_obj.id,
            is_manager=False,
            can_read=True,
            can_chat=True,
            can_share_class=False,
            can_upload_documents=True,
            max_concurrent_chats=3
        )
        db.add(membership)
        db.commit()
        
        # Get member count
        member_count = db.query(ClassMembership).filter(
            ClassMembership.class_id == class_obj.id
        ).count()
        
        result = ClassResponse.model_validate(class_obj)
        result.member_count = member_count
        
        logger.info(f"User {current_user.username} joined class {class_obj.name}")
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error joining class: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/{class_id}", response_model=ClassResponse)
async def get_class_details(
    class_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get class details"""
    try:
        permission_service = PermissionService()
        
        # Check if user has access to this class
        membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not membership:
            raise HTTPException(status_code=403, detail="Access denied")
        
        class_obj = membership.class_obj
        if not class_obj.is_active:
            raise HTTPException(status_code=404, detail="Class not found")
        
        # Get member count
        from app.models import ClassMembership
        member_count = db.query(ClassMembership).filter(
            ClassMembership.class_id == class_id
        ).count()
        
        result = ClassResponse.model_validate(class_obj)
        result.member_count = member_count
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting class details: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.delete("/{class_id}")
async def delete_class(
    class_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a class (owner only)"""
    try:
        from app.models import Class
        
        class_obj = db.query(Class).filter(
            Class.id == class_id,
            Class.owner_id == current_user.id
        ).first()
        
        if not class_obj:
            raise HTTPException(status_code=404, detail="Class not found or access denied")
        
        # Soft delete
        class_obj.is_active = False
        db.commit()
        
        logger.info(f"Class deleted: {class_obj.name} by {current_user.username}")
        return {"message": "Class successfully deleted"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting class: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Internal server error")

### NEW FILE: app/routes/permissions.py
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from typing import List
from app.database import get_db
from app.schemas import MembershipResponse, PermissionUpdate, SponsorshipUpdate, UserResponse
from app.services.permission_service import PermissionService
from app.utils.security import get_current_user
import logging

router = APIRouter()
logger = logging.getLogger(__name__)

@router.get("/{class_id}/members", response_model=List[MembershipResponse])
async def get_class_members(
    class_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all members of a class (managers only)"""
    try:
        permission_service = PermissionService()
        
        # Check if user is manager of this class
        user_membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not user_membership or not user_membership.is_manager:
            raise HTTPException(status_code=403, detail="Only class managers can view members")
        
        from app.models import ClassMembership, User
        
        # Get all memberships with user details
        memberships = db.query(ClassMembership).join(User).filter(
            ClassMembership.class_id == class_id
        ).all()
        
        result = []
        for membership in memberships:
            member_data = MembershipResponse(
                id=membership.id,
                user_id=membership.user_id,
                username=membership.user.username,
                full_name=membership.user.full_name,
                joined_at=membership.joined_at,
                is_manager=membership.is_manager,
                can_read=membership.can_read,
                can_chat=membership.can_chat,
                max_concurrent_chats=membership.max_concurrent_chats,
                can_share_class=membership.can_share_class,
                can_upload_documents=membership.can_upload_documents,
                is_sponsored=membership.is_sponsored,
                daily_token_limit=membership.daily_token_limit,
                weekly_token_limit=membership.weekly_token_limit,
                monthly_token_limit=membership.monthly_token_limit
            )
            result.append(member_data)
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting class members: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.put("/{class_id}/members/{user_id}", response_model=MembershipResponse)
async def update_member_permissions(
    class_id: int,
    user_id: int,
    permission_update: PermissionUpdate,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update member permissions (managers only)"""
    try:
        permission_service = PermissionService()
        
        # Check if current user is manager of this class
        user_membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not user_membership or not user_membership.is_manager:
            raise HTTPException(status_code=403, detail="Only class managers can update permissions")
        
        # Get target membership
        from app.models import ClassMembership
        target_membership = db.query(ClassMembership).filter(
            ClassMembership.class_id == class_id,
            ClassMembership.user_id == user_id
        ).first()
        
        if not target_membership:
            raise HTTPException(status_code=404, detail="Member not found")
        
        # Prevent manager from modifying their own manager status
        if user_id == current_user.id and hasattr(permission_update, 'is_manager'):
            raise HTTPException(status_code=400, detail="Cannot modify your own manager status")
        
        # Update permissions
        update_data = permission_update.model_dump(exclude_unset=True)
        for field, value in update_data.items():
            if hasattr(target_membership, field):
                setattr(target_membership, field, value)
        
        db.commit()
        db.refresh(target_membership)
        
        # Return updated membership
        result = MembershipResponse(
            id=target_membership.id,
            user_id=target_membership.user_id,
            username=target_membership.user.username,
            full_name=target_membership.user.full_name,
            joined_at=target_membership.joined_at,
            is_manager=target_membership.is_manager,
            can_read=target_membership.can_read,
            can_chat=target_membership.can_chat,
            max_concurrent_chats=target_membership.max_concurrent_chats,
            can_share_class=target_membership.can_share_class,
            can_upload_documents=target_membership.can_upload_documents,
            is_sponsored=target_membership.is_sponsored,
            daily_token_limit=target_membership.daily_token_limit,
            weekly_token_limit=target_membership.weekly_token_limit,
            monthly_token_limit=target_membership.monthly_token_limit
        )
        
        logger.info(f"Member permissions updated for user {user_id} in class {class_id}")
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating member permissions: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Internal server error")

@router.put("/{class_id}/sponsorship")
async def update_class_sponsorship(
    class_id: int,
    sponsorship_update: SponsorshipUpdate,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update class sponsorship settings (managers only)"""
    try:
        permission_service = PermissionService()
        
        # Check if current user is manager of this class
        user_membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not user_membership or not user_membership.is_manager:
            raise HTTPException(status_code=403, detail="Only class managers can update sponsorship")
        
        from app.models import ClassMembership
        
        # Update all non-manager memberships
        db.query(ClassMembership).filter(
            ClassMembership.class_id == class_id,
            ClassMembership.is_manager == False
        ).update({
            "is_sponsored": sponsorship_update.is_sponsored
        })
        
        db.commit()
        
        action = "enabled" if sponsorship_update.is_sponsored else "disabled"
        logger.info(f"Class sponsorship {action} for class {class_id}")
        return {"message": f"Class sponsorship {action} successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating class sponsorship: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Internal server error")

### NEW FILE: app/routes/chat.py
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from typing import List
from app.database import get_db
from app.schemas import ChatSessionCreate, ChatSessionResponse, MessageCreate, MessageResponse, ChatResponse, UserResponse
from app.services.permission_service import PermissionService
from app.services.chat_service import ChatService
from app.utils.security import get_current_user
import logging

router = APIRouter()
logger = logging.getLogger(__name__)

@router.post("/sessions", response_model=ChatSessionResponse, status_code=status.HTTP_201_CREATED)
async def create_chat_session(
    session_data: ChatSessionCreate,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new chat session"""
    try:
        permission_service = PermissionService()
        
        # Check if user can chat in this class
        can_chat, reason = await permission_service.can_user_chat(db, current_user.id, session_data.class_id)
        if not can_chat:
            if "limit reached" in reason.lower():
                raise HTTPException(status_code=429, detail=reason)
            else:
                raise HTTPException(status_code=403, detail=reason)
        
        chat_service = ChatService()
        session = await chat_service.create_session(db, current_user.id, session_data)
        
        logger.info(f"Chat session created: {session.title} by {current_user.username}")
        return session
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating chat session: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/sessions", response_model=List[ChatSessionResponse])
async def get_user_sessions(
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get user's chat sessions"""
    try:
        from app.models import ChatSession, ChatMessage
        
        sessions = db.query(ChatSession).filter(
            ChatSession.user_id == current_user.id,
            ChatSession.is_active == True
        ).order_by(ChatSession.updated_at.desc()).all()
        
        result = []
        for session in sessions:
            # Count messages
            message_count = db.query(ChatMessage).filter(
                ChatMessage.session_id == session.id
            ).count()
            
            session_response = ChatSessionResponse.model_validate(session)
            session_response.message_count = message_count
            result.append(session_response)
        
        return result
        
    except Exception as e:
        logger.error(f"Error getting user sessions: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/sessions/{session_id}", response_model=ChatSessionResponse)
async def get_session_details(
    session_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get chat session details"""
    try:
        from app.models import ChatSession, ChatMessage
        
        session = db.query(ChatSession).filter(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
            ChatSession.is_active == True
        ).first()
        
        if not session:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        # Count messages
        message_count = db.query(ChatMessage).filter(
            ChatMessage.session_id == session.id
        ).count()
        
        result = ChatSessionResponse.model_validate(session)
        result.message_count = message_count
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting session details: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.post("/sessions/{session_id}/messages", response_model=ChatResponse)
async def send_message(
    session_id: int,
    message_data: MessageCreate,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Send a message and get AI response"""
    try:
        from app.models import ChatSession
        
        # Get session and verify ownership
        session = db.query(ChatSession).filter(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
            ChatSession.is_active == True
        ).first()
        
        if not session:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        # Check permissions
        permission_service = PermissionService()
        can_chat, reason = await permission_service.can_user_chat(db, current_user.id, session.class_id)
        if not can_chat:
            if "limit reached" in reason.lower():
                raise HTTPException(status_code=429, detail=reason)
            else:
                raise HTTPException(status_code=403, detail=reason)
        
        # Send message and get AI response
        chat_service = ChatService()
        chat_response = await chat_service.send_message(db, session, message_data.content, current_user.id)
        
        logger.info(f"Message sent in session {session_id} by {current_user.username}")
        return chat_response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending message: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/sessions/{session_id}/messages", response_model=List[MessageResponse])
async def get_session_messages(
    session_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db),
    limit: int = 50,
    offset: int = 0
):
    """Get messages from a chat session"""
    try:
        from app.models import ChatSession, ChatMessage
        
        # Verify session ownership
        session = db.query(ChatSession).filter(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id
        ).first()
        
        if not session:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        # Get messages
        messages = db.query(ChatMessage).filter(
            ChatMessage.session_id == session_id
        ).order_by(ChatMessage.timestamp.asc()).offset(offset).limit(limit).all()
        
        return [MessageResponse.model_validate(msg) for msg in messages]
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting session messages: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

### NEW FILE: app/routes/documents.py
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from sqlalchemy.orm import Session
from typing import List
from app.database import get_db
from app.schemas import DocumentResponse, UserResponse
from app.services.permission_service import PermissionService
from app.services.document_service import DocumentService
from app.utils.security import get_current_user
from app.config import get_settings
import logging

router = APIRouter()
settings = get_settings()
logger = logging.getLogger(__name__)

@router.post("/classes/{class_id}/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_class_document(
    class_id: int,
    file: UploadFile = File(...),
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a document to a class (available to all class chats)"""
    try:
        permission_service = PermissionService()
        
        # Check if user can upload documents to this class
        membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not membership:
            raise HTTPException(status_code=403, detail="Access denied")
        
        if not membership.can_upload_documents:
            raise HTTPException(status_code=403, detail="Document upload permission denied")
        
        # Upload and process document
        document_service = DocumentService()
        document = await document_service.upload_class_document(
            db, file, class_id, current_user.id
        )
        
        logger.info(f"Class document uploaded: {document.original_filename} by {current_user.username}")
        return document
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading class document: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.post("/sessions/{session_id}/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_session_document(
    session_id: int,
    file: UploadFile = File(...),
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a document to a specific chat session"""
    try:
        from app.models import ChatSession
        
        # Verify session ownership
        session = db.query(ChatSession).filter(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
            ChatSession.is_active == True
        ).first()
        
        if not session:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        # Check upload permissions
        permission_service = PermissionService()
        membership = await permission_service.get_user_membership(db, current_user.id, session.class_id)
        if not membership.can_upload_documents:
            raise HTTPException(status_code=403, detail="Document upload permission denied")
        
        # Upload and process document
        document_service = DocumentService()
        document = await document_service.upload_session_document(
            db, file, session_id, session.class_id, current_user.id
        )
        
        logger.info(f"Session document uploaded: {document.original_filename} by {current_user.username}")
        return document
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading session document: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/classes/{class_id}", response_model=List[DocumentResponse])
async def get_class_documents(
    class_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all documents uploaded to a class"""
    try:
        permission_service = PermissionService()
        
        # Check if user has read access to this class
        membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not membership:
            raise HTTPException(status_code=403, detail="Access denied")
        
        if not membership.can_read:
            raise HTTPException(status_code=403, detail="Read permission denied")
        
        from app.models import Document, DocumentScope
        
        # Get class documents
        documents = db.query(Document).filter(
            Document.class_id == class_id,
            Document.scope == DocumentScope.CLASS
        ).order_by(Document.uploaded_at.desc()).all()
        
        return [DocumentResponse.model_validate(doc) for doc in documents]
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting class documents: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/sessions/{session_id}", response_model=List[DocumentResponse])
async def get_session_documents(
    session_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all documents uploaded to a specific chat session"""
    try:
        from app.models import ChatSession, Document, DocumentScope
        
        # Verify session ownership
        session = db.query(ChatSession).filter(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id
        ).first()
        
        if not session:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        # Get session documents
        documents = db.query(Document).filter(
            Document.session_id == session_id,
            Document.scope == DocumentScope.CHAT
        ).order_by(Document.uploaded_at.desc()).all()
        
        return [DocumentResponse.model_validate(doc) for doc in documents]
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting session documents: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a document (uploader or class manager only)"""
    try:
        from app.models import Document
        
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check if user can delete this document
        can_delete = False
        
        # Uploader can always delete their own documents
        if document.uploaded_by == current_user.id:
            can_delete = True
        else:
            # Class managers can delete class documents
            permission_service = PermissionService()
            membership = await permission_service.get_user_membership(db, current_user.id, document.class_id)
            if membership and membership.is_manager:
                can_delete = True
        
        if not can_delete:
            raise HTTPException(status_code=403, detail="Permission denied")
        
        # Delete document and its chunks
        document_service = DocumentService()
        await document_service.delete_document(db, document_id)
        
        logger.info(f"Document deleted: {document.original_filename} by {current_user.username}")
        return {"message": "Document deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

### NEW FILE: app/routes/usage.py
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from typing import List
from app.database import get_db
from app.schemas import UsageStats, ClassUsageOverview, UsageRecord, UserResponse
from app.services.usage_service import UsageService
from app.services.permission_service import PermissionService
from app.utils.security import get_current_user
import logging

router = APIRouter()
logger = logging.getLogger(__name__)

@router.get("/my-usage", response_model=List[UsageStats])
async def get_my_usage(
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get current user's usage statistics across all classes"""
    try:
        usage_service = UsageService()
        usage_stats = await usage_service.get_user_usage_by_class(db, current_user.id)
        return usage_stats
        
    except Exception as e:
        logger.error(f"Error getting user usage: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/classes/{class_id}/members", response_model=List[ClassUsageOverview])
async def get_class_usage_overview(
    class_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get usage overview for all members of a class (managers only)"""
    try:
        permission_service = PermissionService()
        
        # Check if user is manager of this class
        membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not membership or not membership.is_manager:
            raise HTTPException(status_code=403, detail="Only class managers can view usage statistics")
        
        usage_service = UsageService()
        usage_overview = await usage_service.get_class_usage_overview(db, class_id)
        return usage_overview
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting class usage overview: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.get("/classes/{class_id}/limits")
async def get_class_limits(
    class_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get current user's limits and usage for a specific class"""
    try:
        permission_service = PermissionService()
        usage_service = UsageService()
        
        # Check if user is member of this class
        membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not membership:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Get usage statistics
        usage_stats = await usage_service.get_user_class_usage(db, current_user.id, class_id)
        return usage_stats
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting class limits: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@router.put("/classes/{class_id}/limits/{user_id}")
async def update_user_limits(
    class_id: int,
    user_id: int,
    daily_limit: int,
    weekly_limit: int,
    monthly_limit: int,
    current_user: UserResponse = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update token limits for a specific user in a class (managers only)"""
    try:
        permission_service = PermissionService()
        
        # Check if current user is manager of this class
        membership = await permission_service.get_user_membership(db, current_user.id, class_id)
        if not membership or not membership.is_manager:
            raise HTTPException(status_code=403, detail="Only class managers can update limits")
        
        # Validate limits
        if daily_limit <= 0 or weekly_limit <= 0 or monthly_limit <= 0:
            raise HTTPException(status_code=400, detail="Limits must be positive")
        
        if daily_limit > weekly_limit or weekly_limit > monthly_limit:
            raise HTTPException(status_code=400, detail="Daily ≤ Weekly ≤ Monthly limits")
        
        # Update user's membership limits
        from app.models import ClassMembership
        target_membership = db.query(ClassMembership).filter(
            ClassMembership.class_id == class_id,
            ClassMembership.user_id == user_id
        ).first()
        
        if not target_membership:
            raise HTTPException(status_code=404, detail="User membership not found")
        
        target_membership.daily_token_limit = daily_limit
        target_membership.weekly_token_limit = weekly_limit
        target_membership.monthly_token_limit = monthly_limit
        
        db.commit()
        
        logger.info(f"Token limits updated for user {user_id} in class {class_id}")
        return {
            "message": "Limits updated successfully",
            "daily_limit": daily_limit,
            "weekly_limit": weekly_limit,
            "monthly_limit": monthly_limit
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating user limits: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Internal server error")

### NEW FILE: app/services/auth_service.py
from sqlalchemy.orm import Session
from app.models import User
from app.schemas import UserCreate, UserUpdate, UserResponse, Token
from app.utils.security import get_password_hash, verify_password, create_access_token
from datetime import timedelta
from app.config import get_settings
import logging

settings = get_settings()
logger = logging.getLogger(__name__)

class AuthService:
    async def create_user(self, db: Session, user_data: UserCreate) -> UserResponse:
        """Create a new user"""
        # Check if email already exists
        existing_user = db.query(User).filter(User.email == user_data.email).first()
        if existing_user:
            raise ValueError("Email already registered")
        
        # Check if username already exists
        existing_username = db.query(User).filter(User.username == user_data.username).first()
        if existing_username:
            raise ValueError("Username already taken")
        
        # Create new user
        hashed_password = get_password_hash(user_data.password)
        new_user = User(
            email=user_data.email,
            username=user_data.username,
            full_name=user_data.full_name,
            hashed_password=hashed_password
        )
        
        db.add(new_user)
        db.commit()
        db.refresh(new_user)
        
        return UserResponse.model_validate(new_user)
    
    async def authenticate_user(self, db: Session, username: str, password: str) -> Token:
        """Authenticate user and return JWT token"""
        user = db.query(User).filter(User.username == username).first()
        
        if not user:
            raise ValueError("Invalid credentials")
        
        if not user.is_active:
            raise ValueError("Account is deactivated")
        
        if not verify_password(password, user.hashed_password):
            raise ValueError("Invalid credentials")
        
        # Create JWT token
        access_token_expires = timedelta(minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": user.username}, expires_delta=access_token_expires
        )
        
        user_response = UserResponse.model_validate(user)
        
        return Token(
            access_token=access_token,
            token_type="bearer",
            user=user_response
        )
    
    async def update_user(self, db: Session, user_id: int, user_update: UserUpdate) -> UserResponse:
        """Update user profile"""
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise ValueError("User not found")
        
        update_data = user_update.model_dump(exclude_unset=True)
        
        # Check email uniqueness if email is being updated
        if "email" in update_data:
            existing_user = db.query(User).filter(
                User.email == update_data["email"],
                User.id != user_id
            ).first()
            if existing_user:
                raise ValueError("Email already taken")
        
        # Update user
        for field, value in update_data.items():
            setattr(user, field, value)
        
        db.commit()
        db.refresh(user)
        
        return UserResponse.model_validate(user)
    
    async def delete_user(self, db: Session, user_id: int):
        """Delete user account (soft delete)"""
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise ValueError("User not found")
        
        # Soft delete by deactivating
        user.is_active = False
        db.commit()

### NEW FILE: app/services/permission_service.py
from sqlalchemy.orm import Session
from datetime import date, timedelta
from typing import Tuple, Optional
from app.models import ClassMembership, ClassUsageTracker, ChatSession
import pytz
import logging

logger = logging.getLogger(__name__)

class PermissionService:
    
    async def get_user_membership(self, db: Session, user_id: int, class_id: int) -> Optional[ClassMembership]:
        """Get user's membership in a specific class"""
        return db.query(ClassMembership).filter(
            ClassMembership.user_id == user_id,
            ClassMembership.class_id == class_id
        ).first()
    
    async def can_user_chat(self, db: Session, user_id: int, class_id: int) -> Tuple[bool, str]:
        """Check if user can chat in this class"""
        membership = await self.get_user_membership(db, user_id, class_id)
        
        if not membership:
            return False, "User not enrolled in this class"
        
        if not membership.can_chat:
            return False, "Chat permission disabled by class manager"
        
        # Check token limits
        usage_check = await self.check_token_limits(db, user_id, class_id, membership)
        if not usage_check[0]:
            return usage_check
        
        # Check concurrent chat limit
        active_chats = await self.count_active_chats(db, user_id, class_id)
        if active_chats >= membership.max_concurrent_chats:
            return False, f"Maximum concurrent chats reached ({membership.max_concurrent_chats}). Please close some chats first."
        
        return True, "OK"
    
    async def check_token_limits(self, db: Session, user_id: int, class_id: int, membership: ClassMembership) -> Tuple[bool, str]:
        """Check if user is within token limits"""
        # Get or create usage tracker
        tracker = await self.get_usage_tracker(db, user_id, class_id)
        
        # Reset counters if needed
        await self.reset_usage_if_needed(db, tracker)
        
        # Check limits
        if tracker.daily_tokens_used >= membership.daily_token_limit:
            return False, "Daily token limit reached. Upgrade to continue chatting or wait for daily reset."
        
        if tracker.weekly_tokens_used >= membership.weekly_token_limit:
            return False, "Weekly token limit reached. Upgrade to continue chatting or wait for weekly reset."
        
        if tracker.monthly_tokens_used >= membership.monthly_token_limit:
            return False, "Monthly token limit reached. Upgrade to continue chatting or wait for monthly reset."
        
        return True, "OK"
    
    async def get_usage_tracker(self, db: Session, user_id: int, class_id: int) -> ClassUsageTracker:
        """Get or create usage tracker for user in class"""
        tracker = db.query(ClassUsageTracker).filter(
            ClassUsageTracker.user_id == user_id,
            ClassUsageTracker.class_id == class_id
        ).first()
        
        if not tracker:
            tracker = ClassUsageTracker(
                user_id=user_id,
                class_id=class_id,
                last_daily_reset=date.today(),
                last_weekly_reset=date.today(),
                last_monthly_reset=date.today()
            )
            db.add(tracker)
            db.commit()
            db.refresh(tracker)
        
        return tracker
    
    async def reset_usage_if_needed(self, db: Session, tracker: ClassUsageTracker):
        """Reset usage counters based on Madrid timezone (00:00)"""
        madrid_tz = pytz.timezone('Europe/Madrid')
        today = date.today()
        
        reset_needed = False
        
        # Daily reset
        if tracker.last_daily_reset != today:
            tracker.daily_tokens_used = 0
            tracker.last_daily_reset = today
            reset_needed = True
        
        # Weekly reset (every Monday)
        if self.is_new_week(tracker.last_weekly_reset, today):
            tracker.weekly_tokens_used = 0
            tracker.last_weekly_reset = today
            reset_needed = True
        
        # Monthly reset (1st of month)
        if self.is_new_month(tracker.last_monthly_reset, today):
            tracker.monthly_tokens_used = 0
            tracker.last_monthly_reset = today
            reset_needed = True
        
        if reset_needed:
            db.commit()
    
    def is_new_week(self, last_reset: date, today: date) -> bool:
        """Check if we've entered a new week (Monday start)"""
        # Get Monday of last reset week
        last_monday = last_reset - timedelta(days=last_reset.weekday())
        # Get Monday of current week
        current_monday = today - timedelta(days=today.weekday())
        return current_monday > last_monday
    
    def is_new_month(self, last_reset: date, today: date) -> bool:
        """Check if we've entered a new month"""
        return (today.year, today.month) != (last_reset.year, last_reset.month)
    
    async def count_active_chats(self, db: Session, user_id: int, class_id: int) -> int:
        """Count active chat sessions for user in class"""
        return db.query(ChatSession).filter(
            ChatSession.user_id == user_id,
            ChatSession.class_id == class_id,
            ChatSession.is_active == True
        ).count()
    
    async def determine_billing(self, db: Session, user_id: int, class_id: int) -> Tuple[int, bool, bool]:
        """
        Determine billing for usage
        Returns: (billed_user_id, is_sponsored, is_overflow)
        """
        membership = await self.get_user_membership(db, user_id, class_id)
        
        if membership and membership.is_sponsored:
            # Find class owner (manager who sponsors)
            from app.models import Class
            class_obj = db.query(Class).filter(Class.id == class_id).first()
            return class_obj.owner_id, True, False  # Manager pays, sponsored, not overflow
        else:
            return user_id, False, True  # User pays, not sponsored, is overflow
    
    async def record_token_usage(self, db: Session, user_id: int, class_id: int, tokens_used: int):
        """Record token usage in the tracker"""
        tracker = await self.get_usage_tracker(db, user_id, class_id)
        await self.reset_usage_if_needed(db, tracker)
        
        tracker.daily_tokens_used += tokens_used
        tracker.weekly_tokens_used += tokens_used
        tracker.monthly_tokens_used += tokens_used
        
        db.commit()

### NEW FILE: app/services/chat_service.py
from sqlalchemy.orm import Session
from app.models import ChatSession, ChatMessage, UsageRecord
from app.schemas import ChatSessionCreate, ChatSessionResponse, MessageResponse, ChatResponse
from app.services.openai_service import OpenAIService
from app.services.permission_service import PermissionService
from app.services.usage_service import UsageService
from datetime import datetime
import time
import logging

logger = logging.getLogger(__name__)

class ChatService:
    def __init__(self):
        self.openai_service = OpenAIService()
    
    async def create_session(self, db: Session, user_id: int, session_data: ChatSessionCreate) -> ChatSessionResponse:
        """Create a new chat session"""
        new_session = ChatSession(
            title=session_data.title,
            user_id=user_id,
            class_id=session_data.class_id
        )
        
        db.add(new_session)
        db.commit()
        db.refresh(new_session)
        
        result = ChatSessionResponse.model_validate(new_session)
        result.message_count = 0
        
        return result
    
    async def send_message(self, db: Session, session: ChatSession, content: str, user_id: int) -> ChatResponse:
        """Send a message and get AI response"""
        start_time = time.time()
        
        try:
            # Create user message
            user_message = ChatMessage(
                session_id=session.id,
                content=content,
                is_user=True,
                timestamp=datetime.utcnow()
            )
            db.add(user_message)
            db.flush()
            
            # Get context from documents
            context = await self._get_context_for_session(db, session)
            
            # Get AI response
            ai_content, tokens_used = await self.openai_service.generate_response(content, context)
            
            # Calculate response time
            response_time_ms = int((time.time() - start_time) * 1000)
            
            # Create AI message
            ai_message = ChatMessage(
                session_id=session.id,
                content=ai_content,
                is_user=False,
                timestamp=datetime.utcnow(),
                response_time_ms=response_time_ms,
                context_used=context[:500] if context else None,  # Store first 500 chars
                tokens_used=tokens_used
            )
            db.add(ai_message)
            
            # Update session timestamp
            session.updated_at = datetime.utcnow()
            
            db.commit()
            db.refresh(user_message)
            db.refresh(ai_message)
            
            # Record usage and billing
            await self._record_usage(db, session, user_id, tokens_used)
            
            # Record token usage for limits
            permission_service = PermissionService()
            await permission_service.record_token_usage(db, user_id, session.class_id, tokens_used)
            
            return ChatResponse(
                user_message=MessageResponse.model_validate(user_message),
                ai_response=MessageResponse.model_validate(ai_message),
                cost=self._calculate_cost(tokens_used),
                response_time_ms=response_time_ms,
                context_provided=bool(context)
            )
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error in send_message: {e}")
            raise
    
    async def _get_context_for_session(self, db: Session, session: ChatSession) -> str:
        """Get relevant document context for the session"""
        try:
            from app.models import Document, DocumentChunk, DocumentScope
            
            context_chunks = []
            
            # Get class-level documents
            class_documents = db.query(Document).filter(
                Document.class_id == session.class_id,
                Document.scope == DocumentScope.CLASS,
                Document.processing_status == "completed"
            ).all()
            
            for doc in class_documents:
                chunks = db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == doc.id
                ).limit(3).all()  # Get top 3 chunks per document
                
                for chunk in chunks:
                    context_chunks.append(f"[{doc.original_filename}]: {chunk.content}")
            
            # Get session-specific documents
            session_documents = db.query(Document).filter(
                Document.session_id == session.id,
                Document.scope == DocumentScope.CHAT,
                Document.processing_status == "completed"
            ).all()
            
            for doc in session_documents:
                chunks = db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == doc.id
                ).limit(5).all()  # More chunks for session-specific docs
                
                for chunk in chunks:
                    context_chunks.append(f"[{doc.original_filename}]: {chunk.content}")
            
            # Combine context (limit to ~4000 chars to leave room for message)
            context = "\n\n".join(context_chunks)
            if len(context) > 4000:
                context = context[:4000] + "..."
            
            return context
            
        except Exception as e:
            logger.error(f"Error getting context: {e}")
            return ""
    
    async def _record_usage(self, db: Session, session: ChatSession, user_id: int, tokens_used: int):
        """Record usage for billing purposes"""
        try:
            # Determine billing
            permission_service = PermissionService()
            billed_user_id, is_sponsored, is_overflow = await permission_service.determine_billing(
                db, user_id, session.class_id
            )
            
            # Calculate cost
            cost = self._calculate_cost(tokens_used)
            
            # Create usage record
            usage_record = UsageRecord(
                user_id=user_id,
                model_name="gpt-4o-mini",
                operation_type="chat",
                input_tokens=int(tokens_used * 0.7),  # Rough estimate
                output_tokens=int(tokens_used * 0.3),
                cost=cost,
                session_id=session.id,
                billed_to_user_id=billed_user_id,
                is_sponsored=is_sponsored,
                is_overflow=is_overflow
            )
            
            db.add(usage_record)
            db.commit()
            
        except Exception as e:
            logger.error(f"Error recording usage: {e}")
    
    def _calculate_cost(self, tokens_used: int) -> float:
        """Calculate cost based on tokens used"""
        from app.config import get_settings
        settings = get_settings()
        
        # GPT-4o-mini pricing: $0.15/$0.60 per million tokens
        input_tokens = int(tokens_used * 0.7)
        output_tokens = int(tokens_used * 0.3)
        
        pricing = settings.OPENAI_PRICING["gpt-4o-mini"]
        input_cost = (input_tokens / 1_000_000) * pricing["input"]
        output_cost = (output_tokens / 1_000_000) * pricing["output"]
        
        return round(input_cost + output_cost, 6)

### NEW FILE: app/services/document_service.py
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException
from app.models import Document, DocumentChunk, DocumentScope, ProcessingStatus
from app.schemas import DocumentResponse
from app.utils.file_processing import FileProcessor
from app.utils.vector_operations import VectorOperations
from app.config import get_settings
import os
import uuid
import logging

settings = get_settings()
logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        self.file_processor = FileProcessor()
        self.vector_ops = VectorOperations()
    
    async def upload_class_document(self, db: Session, file: UploadFile, class_id: int, user_id: int) -> DocumentResponse:
        """Upload and process a class-level document"""
        return await self._upload_document(db, file, class_id, user_id, DocumentScope.CLASS)
    
    async def upload_session_document(self, db: Session, file: UploadFile, session_id: int, class_id: int, user_id: int) -> DocumentResponse:
        """Upload and process a session-specific document"""
        return await self._upload_document(db, file, class_id, user_id, DocumentScope.CHAT, session_id)
    
    async def _upload_document(self, db: Session, file: UploadFile, class_id: int, user_id: int, scope: DocumentScope, session_id: int = None) -> DocumentResponse:
        """Internal method to upload and process documents"""
        try:
            # Validate file
            await self._validate_file(file)
            
            # Generate unique filename
            file_ext = file.filename.split('.')[-1].lower()
            unique_filename = f"{uuid.uuid4()}.{file_ext}"
            file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
            
            # Ensure upload directory exists
            os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
            
            # Save file
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            # Create document record
            document = Document(
                filename=unique_filename,
                original_filename=file.filename,
                file_path=file_path,
                file_type=file_ext,
                file_size=len(content),
                scope=scope,
                class_id=class_id,
                session_id=session_id,
                uploaded_by=user_id,
                processing_status=ProcessingStatus.PENDING
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Process document asynchronously (in real implementation, use background task)
            await self._process_document(db, document)
            
            return DocumentResponse.model_validate(document)
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            # Clean up file if it was created
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail="Error uploading document")
    
    async def _validate_file(self, file: UploadFile):
        """Validate uploaded file"""
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Check file extension
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in settings.ALLOWED_FILE_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"File type '{file_ext}' not allowed. Allowed types: {settings.ALLOWED_FILE_TYPES}"
            )
        
        # Check file size
        content = await file.read()
        if len(content) > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed size of {settings.MAX_FILE_SIZE} bytes"
            )
        
        # Reset file position for later reading
        await file.seek(0)
    
    async def _process_document(self, db: Session, document: Document):
        """Process document: extract text, chunk, and vectorize"""
        try:
            # Update status to processing
            document.processing_status = ProcessingStatus.PROCESSING
            db.commit()
            
            # Extract text from document
            text_content = self.file_processor.extract_text(document.file_path)
            
            if not text_content.strip():
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = "No text content found in document"
                db.commit()
                return
            
            # Chunk the text
            chunks = self.file_processor.chunk_text(text_content)
            
            # Create document chunks
            for i, chunk_text in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=chunk_text,
                    chunk_index=i,
                    char_start=i * 1000,  # Approximate
                    char_end=min((i + 1) * 1000, len(text_content))
                )
                db.add(chunk)
            
            # Update status to completed
            document.processing_status = ProcessingStatus.COMPLETED
            db.commit()
            
            logger.info(f"Document processed successfully: {document.original_filename}")
            
        except Exception as e:
            logger.error(f"Error processing document {document.id}: {e}")
            document.processing_status = ProcessingStatus.FAILED
            document.processing_error = str(e)
            db.commit()
    
    async def delete_document(self, db: Session, document_id: int):
        """Delete document and its chunks"""
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise ValueError("Document not found")
            
            # Delete file from filesystem
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete chunks first (foreign key constraint)
            db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
            
            # Delete document record
            db.delete(document)
            db.commit()
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            db.rollback()
            raise

### NEW FILE: app/services/openai_service.py
import openai
from app.config import get_settings
import logging
import tiktoken

settings = get_settings()
logger = logging.getLogger(__name__)

# Set OpenAI API key
openai.api_key = settings.OPENAI_API_KEY

class OpenAIService:
    def __init__(self):
        self.model = settings.OPENAI_MODEL
        self.encoding = tiktoken.encoding_for_model(self.model)
    
    async def generate_response(self, user_message: str, context: str = None) -> tuple[str, int]:
        """Generate AI response using GPT-4o-mini"""
        try:
            # Prepare messages
            messages = []
            
            # System message
            system_message = """You are StudHelper, an AI assistant designed to help students learn from uploaded course materials. 
            You provide clear, educational explanations and help students understand complex topics.
            
            When answering:
            1. Be educational and supportive
            2. Reference the provided context when relevant
            3. Break down complex concepts into understandable parts
            4. Encourage further learning and questions
            """
            
            if context:
                system_message += f"\n\nRelevant course materials:\n{context}"
            
            messages.append({"role": "system", "content": system_message})
            messages.append({"role": "user", "content": user_message})
            
            # Make API call
            response = await openai.ChatCompletion.acreate(
                model=self.model,
                messages=messages,
                max_tokens=1000,
                temperature=0.7,
                top_p=0.9,
                frequency_penalty=0.1,
                presence_penalty=0.1
            )
            
            # Extract response
            ai_response = response.choices[0].message.content.strip()
            
            # Calculate tokens used
            total_tokens = response.usage.total_tokens
            
            return ai_response, total_tokens
            
        except Exception as e:
            logger.error(f"Error generating AI response: {e}")
            # Return fallback response
            return "I apologize, but I'm experiencing technical difficulties. Please try again later.", 50
    
    async def generate_embeddings(self, texts: list[str]) -> list[list[float]]:
        """Generate embeddings for text chunks"""
        try:
            response = await openai.Embedding.acreate(
                model=settings.EMBEDDING_MODEL,
                input=texts
            )
            
            embeddings = [item.embedding for item in response.data]
            return embeddings
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            return []
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        try:
            return len(self.encoding.encode(text))
        except Exception as e:
            logger.error(f"Error counting tokens: {e}")
            return 0

### NEW FILE: app/services/usage_service.py
from sqlalchemy.orm import Session
from sqlalchemy import func
from datetime import datetime, timedelta
from typing import List, Dict, Any
from app.models import ClassUsageTracker, ClassMembership, User, UsageRecord, ChatMessage
from app.schemas import UsageStats, ClassUsageOverview
from app.services.permission_service import PermissionService
import logging

logger = logging.getLogger(__name__)

class UsageService:
    def __init__(self):
        self.permission_service = PermissionService()
    
    async def get_user_usage_by_class(self, db: Session, user_id: int) -> List[UsageStats]:
        """Get user's usage statistics for all classes they're in"""
        try:
            # Get all user's memberships
            memberships = db.query(ClassMembership).filter(
                ClassMembership.user_id == user_id
            ).all()
            
            usage_stats = []
            
            for membership in memberships:
                stats = await self.get_user_class_usage(db, user_id, membership.class_id)
                usage_stats.append(stats)
            
            return usage_stats
            
        except Exception as e:
            logger.error(f"Error getting user usage by class: {e}")
            return []
    
    async def get_user_class_usage(self, db: Session, user_id: int, class_id: int) -> UsageStats:
        """Get user's usage statistics for a specific class"""
        try:
            # Get or create usage tracker
            tracker = await self.permission_service.get_usage_tracker(db, user_id, class_id)
            await self.permission_service.reset_usage_if_needed(db, tracker)
            
            # Get membership for limits
            membership = await self.permission_service.get_user_membership(db, user_id, class_id)
            
            if not membership:
                raise ValueError("User not found in class")
            
            return UsageStats(
                daily_tokens_used=tracker.daily_tokens_used,
                weekly_tokens_used=tracker.weekly_tokens_used,
                monthly_tokens_used=tracker.monthly_tokens_used,
                daily_limit=membership.daily_token_limit,
                weekly_limit=membership.weekly_token_limit,
                monthly_limit=membership.monthly_token_limit,
                daily_remaining=max(0, membership.daily_token_limit - tracker.daily_tokens_used),
                weekly_remaining=max(0, membership.weekly_token_limit - tracker.weekly_tokens_used),
                monthly_remaining=max(0, membership.monthly_token_limit - tracker.monthly_tokens_used)
            )
            
        except Exception as e:
            logger.error(f"Error getting user class usage: {e}")
            # Return default stats if error
            return UsageStats(
                daily_tokens_used=0,
                weekly_tokens_used=0,
                monthly_tokens_used=0,
                daily_limit=1_000_000,
                weekly_limit=5_000_000,
                monthly_limit=15_000_000,
                daily_remaining=1_000_000,
                weekly_remaining=5_000_000,
                monthly_remaining=15_000_000
            )
    
    async def get_class_usage_overview(self, db: Session, class_id: int) -> List[ClassUsageOverview]:
        """Get usage overview for all members of a class"""
        try:
            # Get all class memberships
            memberships = db.query(ClassMembership).join(User).filter(
                ClassMembership.class_id == class_id
            ).all()
            
            overview = []
            
            for membership in memberships:
                # Get usage stats for this member
                usage_stats = await self.get_user_class_usage(
                    db, membership.user_id, class_id
                )
                
                # Get last activity
                last_activity = db.query(ChatMessage.timestamp).join(
                    ChatMessage.session
                ).filter(
                    ChatMessage.session.has(user_id=membership.user_id, class_id=class_id)
                ).order_by(ChatMessage.timestamp.desc()).first()
                
                member_overview = ClassUsageOverview(
                    user_id=membership.user_id,
                    username=membership.user.username,
                    usage_stats=usage_stats,
                    is_sponsored=membership.is_sponsored,
                    last_activity=last_activity[0] if last_activity else None
                )
                
                overview.append(member_overview)
            
            return overview
            
        except Exception as e:
            logger.error(f"Error getting class usage overview: {e}")
            return []

### NEW FILE: app/utils/security.py
from passlib.context import CryptContext
from jose import JWTError, jwt
from fastapi import HTTPException, status, Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.orm import Session
from datetime import datetime, timedelta
from app.database import get_db
from app.config import get_settings
from app.schemas import UserResponse
import logging

settings = get_settings()
logger = logging.getLogger(__name__)

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Security scheme
security = HTTPBearer()

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: timedelta = None) -> str:
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, settings.SECRET_KEY, algorithm=settings.ALGORITHM)
    return encoded_jwt

async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
    db: Session = Depends(get_db)
) -> UserResponse:
    """Get current authenticated user"""
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    
    try:
        # Extract token from credentials
        token = credentials.credentials
        
        # Decode JWT token
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
            
    except JWTError:
        raise credentials_exception
    
    # Get user from database
    from app.models import User
    user = db.query(User).filter(User.username == username).first()
    if user is None:
        raise credentials_exception
    
    if not user.is_active:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Inactive user"
        )
    
    return UserResponse.model_validate(user)

### NEW FILE: app/utils/file_processing.py
import PyPDF2
import docx
from typing import List
import logging
import os

logger = logging.getLogger(__name__)

class FileProcessor:
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from various file types"""
        try:
            file_ext = file_path.split('.')[-1].lower()
            
            if file_ext == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == 'txt':
                return self._extract_from_txt(file_path)
            elif file_ext == 'docx':
                return self._extract_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e2:
                logger.error(f"Error reading TXT {file_path}: {e2}")
                return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if not text or not text.strip():
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            
            # If this isn't the last chunk, try to end at a sentence boundary
            if end < text_length:
                # Look for sentence endings near the chunk boundary
                sentence_endings = ['. ', '! ', '? ', '\n']
                best_end = end
                
                # Search backwards from the end position for a sentence ending
                for i in range(end, max(start + chunk_size // 2, start), -1):
                    if text[i:i+2] in sentence_endings:
                        best_end = i + 1
                        break
                
                end = best_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position (with overlap)
            start = max(start + 1, end - overlap)
            
            # Avoid infinite loop
            if start >= text_length:
                break
        
        return chunks

### NEW FILE: app/utils/vector_operations.py
import numpy as np
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class VectorOperations:
    """Simple vector operations for document similarity search"""
    
    def cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            vec1 = np.array(vec1)
            vec2 = np.array(vec2)
            
            dot_product = np.dot(vec1, vec2)
            norm1 = np.linalg.norm(vec1)
            norm2 = np.linalg.norm(vec2)
            
            if norm1 == 0 or norm2 == 0:
                return 0.0
            
            similarity = dot_product / (norm1 * norm2)
            return float(similarity)
            
        except Exception as e:
            logger.error(f"Error calculating cosine similarity: {e}")
            return 0.0
    
    def find_most_similar(self, query_vector: List[float], vectors: Dict[str, List[float]], top_k: int = 5) -> List[str]:
        """Find the most similar vectors to the query vector"""
        try:
            similarities = []
            
            for key, vector in vectors.items():
                similarity = self.cosine_similarity(query_vector, vector)
                similarities.append((key, similarity))
            
            # Sort by similarity (descending) and return top_k
            similarities.sort(key=lambda x: x[1], reverse=True)
            return [key for key, _ in similarities[:top_k]]
            
        except Exception as e:
            logger.error(f"Error finding similar vectors: {e}")
            return []
    
    def average_vectors(self, vectors: List[List[float]]) -> List[float]:
        """Calculate the average of multiple vectors"""
        try:
            if not vectors:
                return []
            
            vectors_array = np.array(vectors)
            average = np.mean(vectors_array, axis=0)
            return average.tolist()
            
        except Exception as e:
            logger.error(f"Error averaging vectors: {e}")
            return []

### NEW FILE: requirements.txt
fastapi==0.104.1
uvicorn[standard]==0.24.0
python-multipart==0.0.6
sqlalchemy==2.0.23
psycopg2-binary==2.9.9
alembic==1.13.0
pydantic[email]==2.5.0
pydantic-settings==2.1.0
python-jose[cryptography]==3.3.0
passlib[bcrypt]==1.7.4
PyPDF2==3.0.1
python-docx==1.1.0
openai==1.3.7
tiktoken==0.5.2
pytz==2023.3
numpy==1.24.4
python-decouple==3.8

# Development and testing
pytest==7.4.3
pytest-asyncio==0.21.1
httpx==0.25.2

### NEW FILE: .env.example
# Database
DATABASE_URL=postgresql://postgres:password@localhost/studhelper

# Security
SECRET_KEY=your-super-secret-key-change-this-in-production-at-least-32-characters-long
ALGORITHM=HS256
ACCESS_TOKEN_EXPIRE_MINUTES=30

# OpenAI
OPENAI_API_KEY=your-openai-api-key-here

# File Upload
UPLOAD_DIR=uploads
MAX_FILE_SIZE=10485760
ALLOWED_FILE_TYPES=pdf,txt,docx

# CORS (comma separated)
CORS_ORIGINS=http://localhost:3000,http://localhost:8080

# Application
DEBUG=false

### NEW FILE: Dockerfile
FROM python:3.11-slim

WORKDIR /app

# Install system dependencies
RUN apt-get update && apt-get install -y \
    gcc \
    postgresql-client \
    && rm -rf /var/lib/apt/lists/*

# Copy requirements and install Python dependencies
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

# Copy application code
COPY . .

# Create upload directory
RUN mkdir -p uploads logs

# Create non-root user
RUN useradd --create-home --shell /bin/bash studhelper && \
    chown -R studhelper:studhelper /app
USER studhelper

# Expose port
EXPOSE 8000

# Health check
HEALTHCHECK --interval=30s --timeout=30s --start-period=5s --retries=3 \
    CMD python -c "import requests; requests.get('http://localhost:8000/health')"

# Run the application
CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]

### NEW FILE: docker-compose.yml
version: '3.8'

services:
  db:
    image: postgres:15
    environment:
      POSTGRES_DB: studhelper
      POSTGRES_USER: postgres
      POSTGRES_PASSWORD: password
    ports:
      - "5432:5432"
    volumes:
      - postgres_data:/var/lib/postgresql/data
    healthcheck:
      test: ["CMD-SHELL", "pg_isready -U postgres"]
      interval: 30s
      timeout: 10s
      retries: 5

  backend:
    build: .
    ports:
      - "8000:8000"
    environment:
      - DATABASE_URL=postgresql://postgres:password@db:5432/studhelper
      - SECRET_KEY=development-secret-key-change-in-production
      - OPENAI_API_KEY=${OPENAI_API_KEY}
    depends_on:
      db:
        condition: service_healthy
    volumes:
      - ./uploads:/app/uploads
      - ./logs:/app/logs
    restart: unless-stopped

volumes:
  postgres_data:

### NEW FILE: alembic.ini
# A generic, single database configuration.

[alembic]
# path to migration scripts
script_location = app/migrations

# template used to generate migration file names; The default value is %%(rev)s_%%(slug)s
file_template = %%(year)d_%%(month).2d_%%(day).2d_%%(hour).2d%%(minute).2d-%%(rev)s_%%(slug)s

# sys.path path, will be prepended to sys.path if present.
prepend_sys_path = .

# timezone to use when rendering the date within the migration file
# as well as the filename.
timezone = Europe/Madrid

# max length of characters to apply to the
# "slug" field
truncate_slug_length = 40

# set to 'true' to run the environment during
# the 'revision' command, regardless of autogenerate
revision_environment = false

# set to 'true' to allow .pyc and .pyo files without
# a source .py file to be detected as revisions in the
# versions/ directory
sourceless = false

# version path separator; As mentioned above, this is the character used to split
version_locations = %(here)s/app/migrations/versions

# version path separator; default: os.pathsep
version_path_separator = :

# set to 'true' to search source files recursively
# in each "version_locations" directory
recursive_version_locations = false

# the output encoding used when revision files
# are written from script.py.mako
output_encoding = utf-8

sqlalchemy.url = postgresql://postgres:password@localhost/studhelper

[post_write_hooks]
# post_write_hooks defines scripts or Python functions that are run
# on newly generated revision scripts.  See the documentation for further
# detail and examples

# format using "black" - use the console_scripts runner, against the "black" entrypoint
# hooks = black
# black.type = console_scripts
# black.entrypoint = black
# black.options = -l 79 REVISION_SCRIPT_FILENAME

[loggers]
keys = root,sqlalchemy,alembic

[handlers]
keys = console

[formatters]
keys = generic

[logger_root]
level = WARN
handlers = console
qualname =

[logger_sqlalchemy]
level = WARN
handlers =
qualname = sqlalchemy.engine

[logger_alembic]
level = INFO
handlers =
qualname = alembic

[handler_console]
class = StreamHandler
args = (sys.stderr,)
level = NOTSET
formatter = generic

[formatter_generic]
format = %(levelname)-5.5s [%(name)s] %(message)s
datefmt = %H:%M:%S

### NEW FILE: app/migrations/env.py
from logging.config import fileConfig
from sqlalchemy import engine_from_config, pool
from alembic import context
from app.models import Base
from app.config import get_settings

# Alembic Config object
config = context.config

# Interpret the config file for Python logging
if config.config_file_name is not None:
    fileConfig(config.config_file_name)

# Add your model's MetaData object here for 'autogenerate' support
target_metadata = Base.metadata

# Get database URL from settings
settings = get_settings()
config.set_main_option("sqlalchemy.url", settings.DATABASE_URL)

def run_migrations_offline() -> None:
    """Run migrations in 'offline' mode."""
    url = config.get_main_option("sqlalchemy.url")
    context.configure(
        url=url,
        target_metadata=target_metadata,
        literal_binds=True,
        dialect_opts={"paramstyle": "named"},
    )

    with context.begin_transaction():
        context.run_migrations()

def run_migrations_online() -> None:
    """Run migrations in 'online' mode."""
    connectable = engine_from_config(
        config.get_section(config.config_ini_section, {}),
        prefix="sqlalchemy.",
        poolclass=pool.NullPool,
    )

    with connectable.connect() as connection:
        context.configure(
            connection=connection, target_metadata=target_metadata
        )

        with context.begin_transaction():
            context.run_migrations()

if context.is_offline_mode():
    run_migrations_offline()
else:
    run_migrations_online()

### NEW FILE: app/migrations/script.py.mako
"""${message}

Revision ID: ${up_revision}
Revises: ${down_revision | comma,n}
Create Date: ${create_date}

"""
from alembic import op
import sqlalchemy as sa
${imports if imports else ""}

# revision identifiers, used by Alembic.
revision = ${repr(up_revision)}
down_revision = ${repr(down_revision)}
branch_labels = ${repr(branch_labels)}
depends_on = ${repr(depends_on)}

def upgrade() -> None:
    ${upgrades if upgrades else "pass"}

def downgrade() -> None:
    ${downgrades if downgrades else "pass"}

### NEW FILE: README.md
# StudHelper Backend

A class-based AI learning platform with flexible permissions and granular usage tracking.

## Features

- **User Management**: Registration, authentication, profile management
- **Class System**: Create/join classes with unique codes
- **Permission Management**: Granular class-based permissions
- **Dual-Document System**: Class-wide + chat-specific documents
- **AI Chat**: Context-aware conversations with document integration
- **Usage Tracking**: Per-class, per-user token limits and monitoring

## Quick Start

### Prerequisites

- Python 3.11+
- PostgreSQL 12+
- OpenAI API Key

### Local Development

1. **Clone the repository**
```bash
git clone <repository-url>
cd studhelper-backend
```

2. **Set up environment**
```bash
python -m venv venv
source venv/bin/activate  # On Windows: venv\Scripts\activate
pip install -r requirements.txt
```

3. **Configure environment variables**
```bash
cp .env.example .env
# Edit .env with your settings
```

4. **Set up database**
```bash
# Create PostgreSQL database
createdb studhelper

# Run migrations
alembic upgrade head
```

5. **Run the application**
```bash
uvicorn app.main:app --reload --host 0.0.0.0 --port 8000
```

### Docker Development

1. **Using Docker Compose**
```bash
# Copy environment file
cp .env.example .env
# Add your OPENAI_API_KEY to .env

# Start services
docker-compose up --build
```

The API will be available at `http://localhost:8000`

## API Documentation

Once running, visit:
- **Interactive API docs**: `http://localhost:8000/docs`
- **Alternative docs**: `http://localhost:8000/redoc`

## Project Structure

```
app/
├── main.py              # FastAPI application
├── database.py          # Database configuration
├── config.py           # Settings and configuration
├── models/             # Database models
├── schemas/            # Pydantic schemas
├── routes/             # API endpoints
├── services/           # Business logic
├── utils/              # Helper functions
└── migrations/         # Database migrations
```

## Testing

Run the test suite:
```bash
pytest -v
```

Run with coverage:
```bash
pytest --cov=app tests/
```

## Database Migrations

Create a new migration:
```bash
alembic revision --autogenerate -m "Description of changes"
```

Apply migrations:
```bash
alembic upgrade head
```

## Environment Variables

Required environment variables:

```bash
DATABASE_URL=postgresql://user:password@host:port/database
SECRET_KEY=your-secret-key-at-least-32-characters
OPENAI_API_KEY=your-openai-api-key
```

Optional variables:
```bash
DEBUG=false
UPLOAD_DIR=uploads
MAX_FILE_SIZE=10485760
CORS_ORIGINS=http://localhost:3000
```

## Security

- JWT-based authentication
- Password hashing with bcrypt
- CORS configuration
- File type and size validation
- SQL injection protection via SQLAlchemy ORM

## Usage Tracking

The system tracks token usage per user per class with:
- Daily, weekly, and monthly limits
- Automatic reset at 00:00 Madrid time
- Sponsorship support (managers can pay for students)
- Overflow billing detection

## Document Processing

Supports multiple file types:
- PDF files (text extraction)
- TXT files (plain text)
- DOCX files (Microsoft Word)

Documents are:
- Chunked for efficient processing
- Available at class or chat session level
- Processed asynchronously
- Used for context-aware AI responses

## Production Deployment

1. **Use production environment variables**
2. **Set up SSL/TLS termination**
3. **Configure proper database connection pooling**
4. **Set up logging and monitoring**
5. **Use a production WSGI server** (included: uvicorn)

### Docker Production

```bash
docker-compose -f docker-compose.production.yml up -d
```

## Architecture

- **FastAPI**: Modern, fast web framework
- **SQLAlchemy**: ORM with PostgreSQL
- **Alembic**: Database migrations
- **JWT**: Stateless authentication
- **OpenAI GPT-4o-mini**: AI responses with cost optimization
- **Pydantic**: Data validation and serialization

## Contributing

1. Fork the repository
2. Create a feature branch
3. Add tests for new functionality
4. Ensure all tests pass
5. Submit a pull request

## License

[Add your license here]