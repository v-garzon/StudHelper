"""Word document processor - Fixed version."""

import logging
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document as DocxDocument

from .base import DocumentProcessor

logger = logging.getLogger(__name__)


class DocxProcessor(DocumentProcessor):
    """Word document processor."""
    
    def __init__(self):
        super().__init__()
        self.supported_types = [
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
    
    async def process(self, file_path: Path, doc_metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process Word document and extract text content."""
        try:
            result = await self._extract_content(file_path)
            result.update({
                "file_type": "docx",
                "processor": "DocxProcessor",
                "file_path": str(file_path),
            })
            return result
            
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "chunks": [],
                "doc_metadata": {"file_type": "docx", "error": str(e)},
            }
    
    async def _extract_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Word document."""
        def extract():
            chunks = []
            doc_metadata = {
                "extraction_method": "python-docx",
                "paragraphs": 0,
                "tables": 0,
                "word_count": 0,
            }
            
            doc = DocxDocument(file_path)
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    chunks.append({
                        "content": text,
                        "doc_metadata": self._create_chunk_metadata(
                            source=str(file_path),
                            chunk_index=len(chunks),
                            content_type="paragraph",
                            extraction_method="python-docx",
                        ),
                    })
                    doc_metadata["word_count"] += len(text.split())
                    doc_metadata["paragraphs"] += 1
            
            # Process tables
            for table in doc.tables:
                table_text = self._table_to_text(table)
                if table_text:
                    chunks.append({
                        "content": table_text,
                        "doc_metadata": self._create_chunk_metadata(
                            source=str(file_path),
                            chunk_index=len(chunks),
                            content_type="table",
                            extraction_method="python-docx",
                        ),
                    })
                    doc_metadata["word_count"] += len(table_text.split())
                    doc_metadata["tables"] += 1
            
            return {
                "success": True,
                "chunks": chunks,
                "doc_metadata": doc_metadata,
            }
        
        return await self._run_in_thread(extract)
    
    def _table_to_text(self, table) -> str:
        """Convert Word table to text."""
        text_parts = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            
            if any(text for text in row_text):
                text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)